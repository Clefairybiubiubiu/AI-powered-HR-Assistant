"""
Resume-JD Matching Dashboard
A Streamlit app that matches resumes with job descriptions using TF-IDF and cosine similarity.
"""

import streamlit as st
import pandas as pd
import numpy as np
import os
import re
import time
import hashlib
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any
import plotly.express as px
import plotly.graph_objects as go
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import PyPDF2
import docx
from docx import Document
import warnings
warnings.filterwarnings('ignore')

from resume_matcher.skill_taxonomy_onet import ONETSkillTaxonomy

# Semantic matching imports
try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

# Enhanced resume parser imports
try:
    from enhanced_resume_parser import parse_resume, get_parser
    ENHANCED_PARSER_AVAILABLE = True
except ImportError:
    ENHANCED_PARSER_AVAILABLE = False

# LLM API client imports
try:
    from resume_matcher.utils.llm_client import get_llm_client, is_llm_available
    LLM_AVAILABLE = True
except ImportError:
    LLM_AVAILABLE = False
    def get_llm_client():
        return None
    def is_llm_available():
        return False

# Skills database imports
try:
    from resume_matcher.utils.skills_database import (
        extract_skills_from_text, categorize_skills, ALL_TECHNICAL_SKILLS, SKILLS_DICT
    )
    SKILLS_DATABASE_AVAILABLE = True
except ImportError:
    SKILLS_DATABASE_AVAILABLE = False
    def extract_skills_from_text(text: str, min_confidence: float = 0.7) -> list:
        return []
    def categorize_skills(skills: list) -> dict:
        return {}
    ALL_TECHNICAL_SKILLS = []
    SKILLS_DICT = {}

# Page configuration
st.set_page_config(
    page_title="Resume-JD Matcher",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

class DocumentProcessor:
    """Handles document parsing for different file formats."""
    
    @staticmethod
    def normalize_text(text: str) -> str:
        """Normalize text for consistent parsing."""
        
        # 1. First, normalize line breaks and preserve structure
        text = re.sub(r'\r\n', '\n', text)  # Windows line endings
        text = re.sub(r'\r', '\n', text)     # Mac line endings
        
        # 2. Insert line breaks before section headers (before removing spaces)
        section_headers = [
            'PROFESSIONAL SUMMARY', 'PROFILE', 'SUMMARY', 'OBJECTIVE',
            'EDUCATION', 'EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT',
            'SKILLS', 'TECHNICAL SKILLS', 'COMPETENCIES', 'EXPERTISE',
            'CERTIFICATIONS', 'PROJECTS', 'ACHIEVEMENTS', 'AWARDS',
            'PUBLICATIONS', 'LANGUAGES', 'INTERESTS', 'REFERENCES',
            'CONTACT', 'PERSONAL INFORMATION', 'CAREER OBJECTIVE',
            'PROFESSIONAL EXPERIENCE', 'ACADEMIC BACKGROUND'
        ]
        
        for header in section_headers:
            # Insert line break before headers (case insensitive)
            pattern = r'(?<!\n)(' + re.escape(header) + r')(?!\n)'
            text = re.sub(pattern, r'\n\1', text, flags=re.IGNORECASE)
        
        # Also handle mixed case headers
        mixed_case_headers = [
            'Professional Summary', 'Profile', 'Summary', 'Objective',
            'Education', 'Experience', 'Work Experience', 'Employment',
            'Skills', 'Technical Skills', 'Competencies', 'Expertise',
            'Certifications', 'Projects', 'Achievements', 'Awards',
            'Publications', 'Languages', 'Interests', 'References',
            'Contact', 'Personal Information', 'Career Objective',
            'Professional Experience', 'Academic Background'
        ]
        
        for header in mixed_case_headers:
            pattern = r'(?<!\n)(' + re.escape(header) + r')(?!\n)'
            text = re.sub(pattern, r'\n\1', text)
        
        # 3. Remove bullet symbols for consistent parsing
        bullet_symbols = ['•', '–', '○', '▪', '▫', '‣', '⁃', '◦', '‥', '…']
        for bullet in bullet_symbols:
            text = text.replace(bullet, '')
        
        # 4. Normalize spaces within lines (but preserve line breaks)
        lines = text.split('\n')
        normalized_lines = []
        for line in lines:
            # Replace multiple spaces/tabs with single space within each line
            normalized_line = re.sub(r'[ \t]+', ' ', line.strip())
            normalized_lines.append(normalized_line)
        
        text = '\n'.join(normalized_lines)
        
        # 5. Clean up multiple consecutive empty lines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        return text.strip()
    
    @staticmethod
    def detect_encoding(file_path: str) -> str:
        """Detect file encoding."""
        try:
            import chardet
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding']
                confidence = result['confidence']
                
                # If confidence is low, fall back to common encodings
                if confidence < 0.7:
                    return 'utf-8'
                
                return encoding if encoding else 'utf-8'
        except Exception:
            return 'utf-8'
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extract text from TXT file with encoding detection and normalization."""
        try:
            # Detect encoding
            encoding = DocumentProcessor.detect_encoding(file_path)
            
            # Read file with detected encoding
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
            
            # Normalize text
            normalized_text = DocumentProcessor.normalize_text(text)
            
            print(f"DEBUG: TXT file {file_path} - Encoding: {encoding}, Length: {len(normalized_text)} chars")
            return normalized_text
            
        except UnicodeDecodeError:
            # Fallback to latin-1 if encoding detection fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                normalized_text = DocumentProcessor.normalize_text(text)
                print(f"DEBUG: TXT file {file_path} - Fallback encoding: latin-1, Length: {len(normalized_text)} chars")
                return normalized_text
            except Exception as e:
                print(f"ERROR: Failed to read TXT file {file_path}: {e}")
                return ""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file with normalization."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                # Normalize text
                normalized_text = DocumentProcessor.normalize_text(text)
                
                print(f"DEBUG: PDF file {file_path} - Length: {len(normalized_text)} chars")
                return normalized_text
                
        except Exception as e:
            print(f"ERROR: Failed to read PDF file {file_path}: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file with normalization."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Normalize text
            normalized_text = DocumentProcessor.normalize_text(text)
            
            print(f"DEBUG: DOCX file {file_path} - Length: {len(normalized_text)} chars")
            return normalized_text
            
        except Exception as e:
            print(f"ERROR: Failed to read DOCX file {file_path}: {e}")
            return ""
    
    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """Extract text from any supported file format."""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.txt':
            return cls.extract_text_from_txt(file_path)
        elif file_ext == '.pdf':
            return cls.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return cls.extract_text_from_docx(file_path)
        else:
            st.warning(f"Unsupported file format: {file_ext}")
            return ""

class ResumeJDMatcher:
    """Main class for matching resumes with job descriptions."""
    
    def __init__(self, data_dir: str = None):
        self.data_dir = Path(data_dir) if data_dir else None
        self.processor = DocumentProcessor()
        self.resumes = {}
        self.job_descriptions = {}
        self.similarity_matrix = None
        self.candidate_names = []
        self.jd_names = []
        self.original_filenames = {}
        self.extracted_names = {}
    
    def load_documents(self):
        """Load all resumes and job descriptions from the directory."""
        if self.data_dir is None:
            st.error("No directory specified! Please use file upload or specify a directory.")
            return
            
        if not self.data_dir.exists():
            st.error(f"Directory {self.data_dir} does not exist!")
            return
        
        # Clear existing data to prevent duplicates
        self.resumes.clear()
        self.job_descriptions.clear()
        self.candidate_names.clear()
        self.jd_names.clear()
        self.original_filenames.clear()
        self.extracted_names.clear()
        
        # Load job descriptions first (files starting with "JD")
        jd_files = [f for f in self.data_dir.glob("*") if f.is_file() and f.name.lower().startswith("jd")]
        for file_path in jd_files:
            name = file_path.stem
            text = self.processor.extract_text(str(file_path))
            if text.strip():
                self.job_descriptions[name] = text
                self.jd_names.append(name)
        
        # Load ALL other files as candidate resumes (anything that doesn't start with "JD")
        all_files = [f for f in self.data_dir.glob("*") if f.is_file()]
        candidate_files = [f for f in all_files if not f.name.lower().startswith("jd")]
        
        # Sort files by name for consistent ordering
        candidate_files.sort(key=lambda x: x.name.lower())
        
        for i, file_path in enumerate(candidate_files, 1):
            text = self.processor.extract_text(str(file_path))
            if text.strip():
                # Extract candidate name from resume content
                candidate_name = self.extract_candidate_name(text)
                # Create standardized name: Name-resume
                standardized_name = f"{candidate_name}-resume-{i}"
                
                self.resumes[standardized_name] = text
                self.candidate_names.append(standardized_name)
                self.original_filenames[standardized_name] = file_path.name
                self.extracted_names[standardized_name] = candidate_name
        
        st.success(f"Loaded {len(self.resumes)} resumes and {len(self.job_descriptions)} job descriptions")
        
        # Show mapping of standardized names to original filenames
        if self.original_filenames:
            st.info("📋 **File Mapping:**")
            mapping_lines = "\n".join(f"• {std_name} ← {orig_name}" for std_name, orig_name in self.original_filenames.items())
            st.write(mapping_lines)
    
    def load_documents_from_uploads(self, uploaded_resumes=None, uploaded_jds=None):
        """Load resumes and job descriptions from uploaded files."""
        import io
        
        # Clear existing data
        self.resumes.clear()
        self.job_descriptions.clear()
        self.candidate_names.clear()
        self.jd_names.clear()
        self.original_filenames.clear()
        self.extracted_names.clear()
        
        # Load job descriptions from uploads
        if uploaded_jds:
            for idx, uploaded_file in enumerate(uploaded_jds):
                try:
                    # Determine file type and extract text
                    file_ext = Path(uploaded_file.name).suffix.lower()
                    
                    if file_ext == '.pdf':
                        # Read PDF
                        pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                    elif file_ext in ['.docx', '.doc']:
                        # Read DOCX
                        doc = Document(io.BytesIO(uploaded_file.read()))
                        text = "\n".join([para.text for para in doc.paragraphs])
                    elif file_ext == '.txt':
                        # Read TXT
                        text = uploaded_file.read().decode('utf-8', errors='ignore')
                    else:
                        st.warning(f"Unsupported file type for {uploaded_file.name}")
                        continue
                    
                    if text.strip():
                        name = Path(uploaded_file.name).stem
                        self.job_descriptions[name] = text
                        self.jd_names.append(name)
                except Exception as e:
                    st.error(f"Error processing job description {uploaded_file.name}: {e}")
                    continue
        
        # Load resumes from uploads
        if uploaded_resumes:
            for idx, uploaded_file in enumerate(uploaded_resumes, 1):
                try:
                    # Determine file type and extract text
                    file_ext = Path(uploaded_file.name).suffix.lower()
                    
                    if file_ext == '.pdf':
                        # Read PDF
                        pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                    elif file_ext in ['.docx', '.doc']:
                        # Read DOCX
                        doc = Document(io.BytesIO(uploaded_file.read()))
                        text = "\n".join([para.text for para in doc.paragraphs])
                    elif file_ext == '.txt':
                        # Read TXT
                        text = uploaded_file.read().decode('utf-8', errors='ignore')
                    else:
                        st.warning(f"Unsupported file type for {uploaded_file.name}")
                        continue
                    
                    if text.strip():
                        # Extract candidate name from resume content
                        candidate_name = self.extract_candidate_name(text)
                        # Create standardized name: Name-resume
                        standardized_name = f"{candidate_name}-resume-{idx}"
                        
                        self.resumes[standardized_name] = text
                        self.candidate_names.append(standardized_name)
                        self.original_filenames[standardized_name] = uploaded_file.name
                        self.extracted_names[standardized_name] = candidate_name
                except Exception as e:
                    st.error(f"Error processing resume {uploaded_file.name}: {e}")
                    continue
        
        if not self.resumes and not self.job_descriptions:
            st.warning("⚠️ No files uploaded. Please upload at least one resume or job description.")
        else:
            st.success(f"✅ Loaded {len(self.resumes)} resumes and {len(self.job_descriptions)} job descriptions")
            
            # Show mapping of standardized names to original filenames
            if self.original_filenames:
                st.info("📋 **File Mapping:**")
                mapping_lines = "\n".join(f"• {std_name} ← {orig_name}" for std_name, orig_name in self.original_filenames.items())
                st.write(mapping_lines)
    
    def extract_candidate_name(self, text: str) -> str:
        """Extract candidate name from resume text using Gemini API if available, otherwise use rule-based."""
        # Try Gemini API first if available (only if enabled)
        use_llm_enabled = st.session_state.get('use_llm', False)
        if use_llm_enabled and LLM_AVAILABLE and is_llm_available():
            try:
                llm_client = get_llm_client()
                gemini_name = llm_client.extract_candidate_name(text)
                if gemini_name and gemini_name != "Unknown Candidate":
                    return gemini_name
            except Exception as e:
                print(f"DEBUG: Gemini name extraction failed: {e}, falling back to rule-based")
        
        # Fallback to rule-based extraction
        lines = text.split('\n')
        
        # Look for name patterns in the first few lines
        for i, line in enumerate(lines[:10]):  # Check first 10 lines
            line = line.strip()
            if not line:
                continue
            
            # Pattern 1: "name: Gary" or "Name: Gary" format
            name_match = re.search(r'(?:name|Name)\s*:\s*([A-Za-z\s]+)', line)
            if name_match:
                name = name_match.group(1).strip()
                # Clean up the name (remove extra spaces, symbols)
                name = re.sub(r'[^\w\s]', '', name).strip()
                if name and len(name.split()) >= 1:
                    return name
            
            # Pattern 2: "Name (Title)" format
            match = re.match(r'^([A-Za-z\s]+)\s*\([^)]+\)', line)
            if match:
                name = match.group(1).strip()
                # Clean up symbols and special characters
                name = re.sub(r'[^\w\s]', '', name).strip()
                if len(name.split()) >= 1:  # At least first name
                    return name
            
            # Pattern 3: Just name on first line (clean version)
            if i == 0 and len(line.split()) >= 1:
                # Remove common prefixes and suffixes
                clean_line = re.sub(r'^(name|Name|contact|Contact)\s*:?\s*', '', line, flags=re.IGNORECASE)
                clean_line = re.sub(r'[^\w\s]', '', clean_line).strip()
                
                # Check if it looks like a name (no special characters, reasonable length)
                if (clean_line and 
                    re.match(r'^[A-Za-z\s]+$', clean_line) and 
                    2 <= len(clean_line) <= 50 and
                    not any(word.lower() in ['contact', 'email', 'phone', 'address', 'summary', 'objective'] for word in clean_line.split())):
                    return clean_line
            
            # Pattern 4: Name followed by title on next line
            if i < len(lines) - 1:
                next_line = lines[i + 1].strip()
                clean_line = re.sub(r'[^\w\s]', '', line).strip()
                if (len(clean_line.split()) >= 1 and 
                    not any(word.lower() in ['contact', 'email', 'phone', 'address', 'summary', 'objective'] for word in clean_line.split()) and
                    any(word.lower() in ['engineer', 'developer', 'scientist', 'analyst', 'manager', 'specialist'] for word in next_line.split())):
                    return clean_line
        
        # Fallback: return first meaningful line (cleaned)
        for line in lines[:5]:
            line = line.strip()
            if line:
                # Clean the line
                clean_line = re.sub(r'^(name|Name|contact|Contact)\s*:?\s*', '', line, flags=re.IGNORECASE)
                clean_line = re.sub(r'[^\w\s]', '', clean_line).strip()
                if clean_line and len(clean_line.split()) >= 1:
                    return clean_line
        
        return "Unknown Candidate"
    
    def preprocess_text(self, text: str) -> str:
        """Clean and preprocess text for better matching."""
        # Convert to lowercase
        text = text.lower()
        
        # Remove special characters but keep spaces
        text = re.sub(r'[^\w\s]', ' ', text)
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common stop words (basic list)
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should'}
        words = text.split()
        words = [word for word in words if word not in stop_words and len(word) > 2]
        
        return ' '.join(words)
    
    def compute_similarity(self):
        """Compute TF-IDF and cosine similarity between resumes and job descriptions."""
        if not self.resumes or not self.job_descriptions:
            st.error("No documents loaded!")
            return
        
        # Prepare documents
        all_docs = []
        doc_labels = []
        
        # Add resumes
        for name, text in self.resumes.items():
            processed_text = self.preprocess_text(text)
            all_docs.append(processed_text)
            doc_labels.append(f"Resume: {name}")
        
        # Add job descriptions
        for name, text in self.job_descriptions.items():
            processed_text = self.preprocess_text(text)
            all_docs.append(processed_text)
            doc_labels.append(f"JD: {name}")
        
        # Compute TF-IDF
        vectorizer = TfidfVectorizer(
            max_features=1000,
            ngram_range=(1, 2),
            min_df=1,
            max_df=0.8
        )
        
        tfidf_matrix = vectorizer.fit_transform(all_docs)
        
        # Compute cosine similarity
        similarity_matrix = cosine_similarity(tfidf_matrix)
        
        # Extract resume-JD similarities
        num_resumes = len(self.resumes)
        num_jds = len(self.job_descriptions)
        
        self.similarity_matrix = similarity_matrix[:num_resumes, num_resumes:]
        
        return self.similarity_matrix
    
    def get_top_matches(self, top_k: int = 3) -> pd.DataFrame:
        """Get top matches for each job description."""
        if self.similarity_matrix is None:
            return pd.DataFrame()
        
        results = []
        
        for jd_idx, jd_name in enumerate(self.jd_names):
            similarities = self.similarity_matrix[:, jd_idx]
            top_indices = np.argsort(similarities)[::-1][:top_k]
            
            for rank, resume_idx in enumerate(top_indices):
                results.append({
                    'Job Description': jd_name,
                    'Rank': rank + 1,
                    'Candidate': self.candidate_names[resume_idx],
                    'Similarity Score': similarities[resume_idx],
                    'Match Percentage': f"{similarities[resume_idx] * 100:.1f}%"
                })
        
        return pd.DataFrame(results)
    
    def get_directory_info(self) -> Dict:
        """Get information about files in the directory or uploaded files."""
        # If using file uploads (data_dir is None), return info from loaded documents
        if self.data_dir is None:
            return {
                "resume_files": [{"name": name} for name in self.candidate_names],
                "jd_files": [{"name": name} for name in self.jd_names],
                "total_files": len(self.candidate_names) + len(self.jd_names)
            }
        
        # If directory doesn't exist, return empty
        if not self.data_dir.exists():
            return {"resume_files": [], "jd_files": [], "total_files": 0}
        
        resume_files = []
        jd_files = []
        
        for file_path in self.data_dir.glob("*"):
            if file_path.is_file():
                if file_path.name.lower().startswith("jd"):
                    jd_files.append({
                        "name": file_path.name,
                        "size": file_path.stat().st_size,
                        "modified": file_path.stat().st_mtime
                    })
                else:
                    # All non-JD files are treated as candidate resumes
                    resume_files.append({
                        "name": file_path.name,
                        "size": file_path.stat().st_size,
                        "modified": file_path.stat().st_mtime
                    })
        
        return {
            "resume_files": resume_files,
            "jd_files": jd_files,
            "total_files": len(resume_files) + len(jd_files)
        }
    
    def extract_resume_summary(self, resume_text: str) -> Dict:
        """Extract key information from resume."""
        lines = resume_text.split('\n')
        summary = {
            'name': self.extract_candidate_name(resume_text),
            'email': '',
            'phone': '',
            'location': '',
            'skills': [],
            'experience': [],
            'education': [],
            'summary': ''
        }
        
        # Extract contact information
        for line in lines:
            line_lower = line.lower()
            if '@' in line and 'email' in line_lower:
                summary['email'] = line.strip()
            elif 'phone' in line_lower or re.search(r'\(\d{3}\)', line):
                summary['phone'] = line.strip()
            elif any(city in line_lower for city in ['san francisco', 'new york', 'seattle', 'chicago', 'boston', 'austin']):
                summary['location'] = line.strip()
        
        # Extract skills (look for skills sections)
        in_skills_section = False
        for line in lines:
            line_lower = line.lower()
            if 'skill' in line_lower and ('technical' in line_lower or 'core' in line_lower):
                in_skills_section = True
                continue
            elif in_skills_section and line.strip():
                if any(word in line_lower for word in ['experience', 'education', 'work', 'project']):
                    in_skills_section = False
                    break
                # Extract skills from this line
                skills = re.findall(r'\b[A-Za-z][A-Za-z0-9\s]+\b', line)
                summary['skills'].extend([skill.strip() for skill in skills if len(skill.strip()) > 2])
        
        # Extract experience (look for job titles and companies)
        for line in lines:
            if any(word in line.lower() for word in ['engineer', 'developer', 'scientist', 'analyst', 'manager']) and '|' in line:
                summary['experience'].append(line.strip())
        
        # Extract education
        for line in lines:
            if any(word in line.lower() for word in ['university', 'college', 'bachelor', 'master', 'phd', 'degree']):
                summary['education'].append(line.strip())
        
        # Extract summary/objective with enhanced detection
        in_summary = False
        summary_keywords = [
            'summary', 'objective', 'profile', 'about', 'overview', 
            'professional summary', 'professional overview', 'about me',
            'executive summary', 'career summary', 'personal summary',
            'professional profile', 'career objective', 'personal statement'
        ]
        
        for line in lines:
            line_lower = line.lower()
            # Check for summary section headers
            if any(keyword in line_lower for keyword in summary_keywords):
                in_summary = True
                continue
            elif in_summary and line.strip():
                # Stop if we hit another major section
                if any(word in line_lower for word in ['experience', 'education', 'skills', 'contact', 'work history', 'employment']):
                    break
                summary['summary'] += line.strip() + ' '
        
        # Add fallback: if summary is missing, concatenate experience + education
        if not summary['summary'].strip():
            fallback_parts = []
            if summary['experience']:
                fallback_parts.extend(summary['experience'][:2])  # First 2 experience entries
            if summary['education']:
                fallback_parts.extend(summary['education'][:1])   # First education entry
            if fallback_parts:
                summary['summary'] = ' '.join(fallback_parts)
        
        return summary
    
    def extract_jd_requirements(self, jd_text: str) -> Dict:
        """Extract requirements from job description."""
        lines = jd_text.split('\n')
        requirements = {
            'title': '',
            'company': '',
            'location': '',
            'salary': '',
            'requirements': [],
            'responsibilities': [],
            'skills_required': []
        }
        
        # Extract title and company
        for line in lines[:5]:
            if any(word in line.lower() for word in ['engineer', 'developer', 'scientist', 'analyst', 'manager']):
                requirements['title'] = line.strip()
                break
        
        # Extract company
        for line in lines:
            if 'company:' in line.lower() or 'at ' in line.lower():
                requirements['company'] = line.strip()
                break
        
        # Extract requirements
        in_requirements = False
        for line in lines:
            line_lower = line.lower()
            if any(word in line_lower for word in ['requirement', 'qualification', 'must have', 'should have']):
                in_requirements = True
                continue
            elif in_requirements and line.strip():
                if any(word in line_lower for word in ['responsibilit', 'benefit', 'compensation', 'salary']):
                    in_requirements = False
                    break
                if line.strip().startswith('-') or line.strip().startswith('•'):
                    requirements['requirements'].append(line.strip())
        
        # Extract skills
        for line in lines:
            if any(skill in line.lower() for skill in ['python', 'java', 'sql', 'javascript', 'react', 'aws', 'docker']):
                requirements['skills_required'].append(line.strip())
        
        return requirements

@st.cache_resource
def load_sentence_transformer():
    """Load and cache the SentenceTransformer model."""
    if not SENTENCE_TRANSFORMERS_AVAILABLE:
        return None
    try:
        model = SentenceTransformer('all-MiniLM-L6-v2')
        return model
    except Exception as e:
        st.error(f"Failed to load SentenceTransformer: {e}")
        return None

def extract_importance_level(text: str) -> float:
    """Extract importance level from text."""
    text_lower = text.lower()
    if "required" in text_lower or "must have" in text_lower or "mandatory" in text_lower:
        return 1.0
    elif "preferred" in text_lower or "plus" in text_lower or "nice to have" in text_lower:
        return 0.75
    elif "optional" in text_lower or "bonus" in text_lower:
        return 0.5
    return 1.0

def extract_structured_requirements(jd_text: str) -> Dict:
    """
    Extract and structure requirements from job description.
    Returns summaries for Required and Preferred sections.
    """
    
    structured = {
        'required': {
            'summary': '',
            'yoe': None,
            'skills_summary': '',
            'education_summary': ''
        },
        'preferred': {
            'summary': '',
            'yoe': None,
            'skills_summary': '',
            'education_summary': ''
        }
    }
    
    lines = jd_text.split('\n')
    in_requirements_section = False
    current_category = 'required'  # Default to required
    
    # YOE patterns
    yoe_patterns = [
        r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*(?:of\s*)?(?:experience|exp)',
        r'(\d+(?:\.\d+)?)\s*\+?\s*yrs?\s*(?:of\s*)?(?:experience|exp)',
        r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*in',
        r'minimum\s+(\d+(?:\.\d+)?)\s*years?',
        r'at\s+least\s+(\d+(?:\.\d+)?)\s*years?'
    ]
    
    required_items = []
    preferred_items = []
    
    for i, line in enumerate(lines):
        line_lower = line.lower().strip()
        original_line = line.strip()
        
        # Detect requirements section headers
        if any(keyword in line_lower for keyword in ['requirement', 'qualification', 'must have', 'should have']):
            in_requirements_section = True
            # Check if it's preferred section
            if any(word in line_lower for word in ['preferred', 'nice to have', 'nice-to-have', 'plus', 'bonus']):
                current_category = 'preferred'
            else:
                current_category = 'required'
            continue
        
        # Detect preferred section explicitly
        if any(keyword in line_lower for keyword in ['preferred', 'nice to have', 'nice-to-have', 'bonus']):
            if 'requirement' in line_lower or 'qualification' in line_lower:
                current_category = 'preferred'
                in_requirements_section = True
                continue
        
        # Stop at other major sections
        if in_requirements_section:
            if any(keyword in line_lower for keyword in ['responsibilit', 'benefit', 'compensation', 'salary', 'about us', 'company']):
                if not any(keyword in line_lower for keyword in ['requirement', 'qualification', 'preferred']):
                    in_requirements_section = False
                    continue
        
        if not in_requirements_section or not original_line:
            continue
        
        # Remove bullet points and clean
        clean_line = re.sub(r'^[•\-\*\d+\.\)]\s*', '', original_line).strip()
        if not clean_line or len(clean_line) < 3:
            continue
        
        # Skip section headers
        if any(word in clean_line.lower() for word in ['requirement', 'qualification', 'must have', 'should have', 'preferred', 'nice to have']):
            if len(clean_line.split()) <= 5:  # Likely a header
                continue
        
        # Determine importance based on keywords in the line
        importance = extract_importance_level(clean_line)
        category = 'required' if importance >= 1.0 else 'preferred'
        
        # Also check if line contains preferred indicators
        if any(word in line_lower for word in ['preferred', 'nice to have', 'nice-to-have', 'plus', 'bonus', 'optional']):
            category = 'preferred'
        
        # Clean up the line
        clean_line = re.sub(r'^(required|preferred|must have|should have|nice to have|nice-to-have):?\s*', '', clean_line, flags=re.IGNORECASE)
        clean_line = clean_line.strip()
        
        if not clean_line or len(clean_line) < 3:
            continue
        
        # Extract YOE if present
        yoe_found = None
        for pattern in yoe_patterns:
            match = re.search(pattern, line_lower)
            if match:
                try:
                    years = float(match.group(1))
                    if 0.5 <= years <= 20:
                        yoe_found = f"{years:.1f} years"
                        structured[category]['yoe'] = yoe_found
                        # Remove YOE from the line text
                        clean_line = re.sub(pattern, '', clean_line, flags=re.IGNORECASE).strip()
                        clean_line = re.sub(r'\s+', ' ', clean_line)
                        break
                except:
                    continue
        
        # Add to appropriate category
        if category == 'required':
            required_items.append(clean_line)
        else:
            preferred_items.append(clean_line)
    
    # Build summaries as readable paragraphs
    # Required summary
    if required_items:
        # Extract YOE, Skills, and Education from required items
        yoe_text = None
        skills_items = []
        education_items = []
        other_items = []
        
        education_keywords = ['bachelor', 'master', 'phd', 'ph.d', 'mba', 'degree', 'diploma', 'university', 'college']
        
        for item in required_items:
            item_lower = item.lower()
            
            # Check for YOE
            for pattern in yoe_patterns:
                match = re.search(pattern, item_lower)
                if match:
                    try:
                        years = float(match.group(1))
                        if 0.5 <= years <= 20:
                            yoe_text = f"{years:.1f} years"
                            structured['required']['yoe'] = yoe_text
                            break
                    except:
                        continue
            
            # Check for education
            if any(keyword in item_lower for keyword in education_keywords):
                education_items.append(item)
            # Check for skills/experience (most items will be skills/experience)
            elif any(phrase in item_lower for phrase in ['experience', 'skill', 'proficient', 'familiar', 'knowledge', 'with', 'strong', 'good']):
                skills_items.append(item)
            else:
                other_items.append(item)
        
        # Build summary as a single flowing paragraph
        # Combine all items into one coherent summary
        all_items = []
        
        # Add YOE if available
        if yoe_text:
            all_items.append(f"{yoe_text} of experience")
        
        # Add all other items
        all_items.extend(skills_items)
        all_items.extend(education_items)
        all_items.extend(other_items)
        
        # Create a flowing summary paragraph
        if all_items:
            # Join items with commas and proper sentence structure
            summary_text = ", ".join(all_items)
            # Clean up punctuation
            summary_text = re.sub(r'\.\s*\.', '.', summary_text)
            summary_text = re.sub(r',\s*,', ',', summary_text)
            summary_text = re.sub(r'\s+', ' ', summary_text)
            # Ensure it ends with a period
            if not summary_text.endswith('.'):
                summary_text += '.'
            # Capitalize first letter
            summary_text = summary_text[0].upper() + summary_text[1:] if summary_text else summary_text
            
            structured['required']['summary'] = summary_text
            # Store individual summaries for reference
            if skills_items:
                structured['required']['skills_summary'] = ", ".join(skills_items)
            if education_items:
                structured['required']['education_summary'] = ", ".join(education_items)
        else:
            # Fallback: just join all required items
            fallback_text = ", ".join(required_items)
            if not fallback_text.endswith('.'):
                fallback_text += '.'
            structured['required']['summary'] = fallback_text
    
    # Preferred summary
    if preferred_items:
        # Extract YOE, Skills, and Education from preferred items
        yoe_text = None
        skills_items = []
        education_items = []
        other_items = []
        
        education_keywords = ['bachelor', 'master', 'phd', 'ph.d', 'mba', 'degree', 'diploma', 'university', 'college']
        
        for item in preferred_items:
            item_lower = item.lower()
            
            # Check for YOE
            for pattern in yoe_patterns:
                match = re.search(pattern, item_lower)
                if match:
                    try:
                        years = float(match.group(1))
                        if 0.5 <= years <= 20:
                            yoe_text = f"{years:.1f} years"
                            structured['preferred']['yoe'] = yoe_text
                            break
                    except:
                        continue
            
            # Check for education
            if any(keyword in item_lower for keyword in education_keywords):
                education_items.append(item)
            # Check for skills/experience
            elif any(phrase in item_lower for phrase in ['experience', 'skill', 'proficient', 'familiar', 'knowledge', 'with', 'strong', 'good']):
                skills_items.append(item)
            else:
                other_items.append(item)
        
        # Build summary as a single flowing paragraph
        # Combine all items into one coherent summary
        all_items = []
        
        # Add YOE if available
        if yoe_text:
            all_items.append(f"{yoe_text} of experience")
        
        # Add all other items
        all_items.extend(skills_items)
        all_items.extend(education_items)
        all_items.extend(other_items)
        
        # Create a flowing summary paragraph
        if all_items:
            # Join items with commas and proper sentence structure
            summary_text = ", ".join(all_items)
            # Clean up punctuation
            summary_text = re.sub(r'\.\s*\.', '.', summary_text)
            summary_text = re.sub(r',\s*,', ',', summary_text)
            summary_text = re.sub(r'\s+', ' ', summary_text)
            # Ensure it ends with a period
            if not summary_text.endswith('.'):
                summary_text += '.'
            # Capitalize first letter
            summary_text = summary_text[0].upper() + summary_text[1:] if summary_text else summary_text
            
            structured['preferred']['summary'] = summary_text
            # Store individual summaries for reference
            if skills_items:
                structured['preferred']['skills_summary'] = ", ".join(skills_items)
            if education_items:
                structured['preferred']['education_summary'] = ", ".join(education_items)
        else:
            # Fallback: just join all preferred items
            fallback_text = ", ".join(preferred_items)
            if not fallback_text.endswith('.'):
                fallback_text += '.'
            structured['preferred']['summary'] = fallback_text
    
    return structured

class ResumeSemanticMatcher:
    """Semantic matching using Sentence-BERT."""
    
    def __init__(self, data_dir: str = None, use_onet_taxonomy: bool = False):
        self.data_dir = Path(data_dir) if data_dir else None
        self.processor = DocumentProcessor()
        self.resumes = {}
        self.job_descriptions = {}
        self.candidate_names = []
        self.jd_names = []
        self.model = load_sentence_transformer()  # Load model during initialization
        self.embeddings_cache = {}
        self.similarity_matrix = None
        self.original_filenames = {}
        self.extracted_names = {}
        self.use_onet_taxonomy = use_onet_taxonomy
        self.onet_taxonomy = ONETSkillTaxonomy() if use_onet_taxonomy else None
        
    def load_documents(self):
        """Load all resumes and job descriptions from the directory."""
        if self.data_dir is None:
            st.error("No directory specified! Please use file upload or specify a directory.")
            return
            
        if not self.data_dir.exists():
            st.error(f"Directory {self.data_dir} does not exist!")
            return
        
        # Clear existing data to prevent duplicates
        self.resumes.clear()
        self.job_descriptions.clear()
        self.candidate_names.clear()
        self.jd_names.clear()
        self.original_filenames.clear()
        self.extracted_names.clear()
        
        # Load job descriptions first (files starting with "JD")
        jd_files = [f for f in self.data_dir.glob("*") if f.is_file() and f.name.lower().startswith("jd")]
        for file_path in jd_files:
            name = file_path.stem
            text = self.processor.extract_text(str(file_path))
            if text.strip():
                self.job_descriptions[name] = text
                self.jd_names.append(name)
        
        # Load ALL other files as candidate resumes (anything that doesn't start with "JD")
        all_files = [f for f in self.data_dir.glob("*") if f.is_file()]
        candidate_files = [f for f in all_files if not f.name.lower().startswith("jd")]
        
        # Sort files by name for consistent ordering
        candidate_files.sort(key=lambda x: x.name.lower())
        
        for i, file_path in enumerate(candidate_files, 1):
            text = self.processor.extract_text(str(file_path))
            if text.strip():
                # Extract candidate name from resume content
                candidate_name = self.extract_candidate_name(text)
                # Create standardized name: Name-resume
                standardized_name = f"{candidate_name}-resume-{i}"
                
                self.resumes[standardized_name] = text
                self.candidate_names.append(standardized_name)
                self.original_filenames[standardized_name] = file_path.name
                self.extracted_names[standardized_name] = candidate_name
        
        st.success(f"Loaded {len(self.resumes)} resumes and {len(self.job_descriptions)} job descriptions")
    
    def load_documents_from_uploads(self, uploaded_resumes=None, uploaded_jds=None):
        """Load resumes and job descriptions from uploaded files."""
        import io
        
        # Clear existing data
        self.resumes.clear()
        self.job_descriptions.clear()
        self.candidate_names.clear()
        self.jd_names.clear()
        self.original_filenames.clear()
        self.extracted_names.clear()
        
        # Load job descriptions from uploads
        if uploaded_jds:
            for idx, uploaded_file in enumerate(uploaded_jds):
                try:
                    # Determine file type and extract text
                    file_ext = Path(uploaded_file.name).suffix.lower()
                    
                    if file_ext == '.pdf':
                        # Read PDF
                        pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                    elif file_ext in ['.docx', '.doc']:
                        # Read DOCX
                        doc = Document(io.BytesIO(uploaded_file.read()))
                        text = "\n".join([para.text for para in doc.paragraphs])
                    elif file_ext == '.txt':
                        # Read TXT
                        text = uploaded_file.read().decode('utf-8', errors='ignore')
                    else:
                        st.warning(f"Unsupported file type for {uploaded_file.name}")
                        continue
                    
                    if text.strip():
                        name = Path(uploaded_file.name).stem
                        self.job_descriptions[name] = text
                        self.jd_names.append(name)
                except Exception as e:
                    st.error(f"Error processing job description {uploaded_file.name}: {e}")
                    continue
        
        # Load resumes from uploads
        if uploaded_resumes:
            for idx, uploaded_file in enumerate(uploaded_resumes, 1):
                try:
                    # Determine file type and extract text
                    file_ext = Path(uploaded_file.name).suffix.lower()
                    
                    if file_ext == '.pdf':
                        # Read PDF
                        pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                    elif file_ext in ['.docx', '.doc']:
                        # Read DOCX
                        doc = Document(io.BytesIO(uploaded_file.read()))
                        text = "\n".join([para.text for para in doc.paragraphs])
                    elif file_ext == '.txt':
                        # Read TXT
                        text = uploaded_file.read().decode('utf-8', errors='ignore')
                    else:
                        st.warning(f"Unsupported file type for {uploaded_file.name}")
                        continue
                    
                    if text.strip():
                        # Extract candidate name from resume content
                        candidate_name = self.extract_candidate_name(text)
                        # Create standardized name: Name-resume
                        standardized_name = f"{candidate_name}-resume-{idx}"
                        
                        self.resumes[standardized_name] = text
                        self.candidate_names.append(standardized_name)
                        self.original_filenames[standardized_name] = uploaded_file.name
                        self.extracted_names[standardized_name] = candidate_name
                except Exception as e:
                    st.error(f"Error processing resume {uploaded_file.name}: {e}")
                    continue
        
        if not self.resumes and not self.job_descriptions:
            st.warning("⚠️ No files uploaded. Please upload at least one resume or job description.")
        else:
            st.success(f"✅ Loaded {len(self.resumes)} resumes and {len(self.job_descriptions)} job descriptions")
    
    def extract_candidate_name(self, text: str) -> str:
        """Extract candidate name from resume text."""
        lines = text.split('\n')
        
        # Look for name patterns in the first few lines
        for i, line in enumerate(lines[:10]):  # Check first 10 lines
            line = line.strip()
            if not line:
                continue
            
            # Pattern 1: "name: Gary" or "Name: Gary" format
            name_match = re.search(r'(?:name|Name)\s*:\s*([A-Za-z\s]+)', line)
            if name_match:
                name = name_match.group(1).strip()
                # Clean up the name (remove extra spaces, symbols)
                name = re.sub(r'[^\w\s]', '', name).strip()
                if name and len(name.split()) >= 1:
                    return name
            
            # Pattern 2: "Name (Title)" format
            match = re.match(r'^([A-Za-z\s]+)\s*\([^)]+\)', line)
            if match:
                name = match.group(1).strip()
                # Clean up symbols and special characters
                name = re.sub(r'[^\w\s]', '', name).strip()
                if len(name.split()) >= 1:  # At least first name
                    return name
            
            # Pattern 3: Just name on first line (clean version)
            if i == 0 and len(line.split()) >= 1:
                # Remove common prefixes and suffixes
                clean_line = re.sub(r'^(name|Name|contact|Contact)\s*:?\s*', '', line, flags=re.IGNORECASE)
                clean_line = re.sub(r'[^\w\s]', '', clean_line).strip()
                
                # Check if it looks like a name (no special characters, reasonable length)
                if (clean_line and 
                    re.match(r'^[A-Za-z\s]+$', clean_line) and 
                    2 <= len(clean_line) <= 50 and
                    not any(word.lower() in ['contact', 'email', 'phone', 'address', 'summary', 'objective'] for word in clean_line.split())):
                    return clean_line
        
        # Fallback: return first meaningful line (cleaned)
        for line in lines[:5]:
            line = line.strip()
            if line:
                # Clean the line
                clean_line = re.sub(r'^(name|Name|contact|Contact)\s*:?\s*', '', line, flags=re.IGNORECASE)
                clean_line = re.sub(r'[^\w\s]', '', clean_line).strip()
                if clean_line and len(clean_line.split()) >= 1:
                    return clean_line
        
        return "Unknown Candidate"
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """Extract Education, Skills, and Experience sections from text using Gemini API if available."""
        sections = {
            'education': '',
            'skills': '',
            'experience': '',
            'summary': ''
        }
        section_char_limits = {
            'summary': 1200,
            'experience': 2600,
            'education': 1200,
            'skills': 1000
        }
        
        # Try Gemini API first if available (only if enabled)
        use_llm_enabled = st.session_state.get('use_llm', False)
        if use_llm_enabled and LLM_AVAILABLE and is_llm_available():
            try:
                llm_client = get_llm_client()
                gemini_sections = llm_client.enhance_resume_parsing(text, target_section="all")
                if gemini_sections:
                    # Use Gemini results if they're substantial
                    for section in ['education', 'skills', 'experience', 'summary']:
                        if section in gemini_sections and len(gemini_sections[section].strip()) > 20:
                            sections[section] = gemini_sections[section]
                            print(f"DEBUG: Using Gemini-extracted {section} section ({len(gemini_sections[section])} chars)")
            except Exception as e:
                print(f"DEBUG: Gemini section extraction failed: {e}, falling back to rule-based")
        
        # Fallback to rule-based extraction (or supplement if Gemini didn't extract everything)
        # Clean and normalize the text
        text = self._clean_resume_text(text)
        
        # Split into lines and process
        lines = text.split('\n')
        current_section = None
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            original_line = line.strip()
            
            # Skip empty lines
            if not original_line:
                continue
            
            # Detect section headers
            section_detected = self._detect_section_header(line_lower)
            if section_detected:
                current_section = section_detected
                print(f"DEBUG: Detected section '{section_detected}' from line: '{original_line}'")
                continue
            
            # Skip contact information sections
            if self._is_contact_info(line_lower):
                current_section = None
                continue
            
            # Add content to current section (only if Gemini didn't extract it or it's too short)
            if current_section and original_line:
                if self._should_include_content(current_section, line_lower, original_line):
                    existing = sections.get(current_section, '')
                    char_limit = section_char_limits.get(current_section, 1200)
                    if len(existing) < char_limit:
                        addition = original_line.strip() + " "
                        available = char_limit - len(existing)
                        sections[current_section] = (existing + addition)[:char_limit]

        # If summary is still weak, pull first descriptive paragraph before experience section
        if len(sections['summary'].strip()) < 80:
            fallback_summary = self._extract_summary_fallback(text)
            if fallback_summary:
                sections['summary'] = fallback_summary

        # If education missing, try heuristic extraction without explicit headers
        if len(sections['education'].strip()) < 80:
            fallback_education = self._extract_education_fallback(text)
            if fallback_education:
                sections['education'] = fallback_education
        
        # Cache structured education entries for downstream displays
        structured_entries = self._structure_education_entries(sections.get('education', ''))
        if structured_entries:
            sections['_education_entries'] = structured_entries
        
        # Enhanced skill extraction using skills database (if skills section is empty or short)
        if not sections.get('skills') or len(sections.get('skills', '').strip()) < 20:
            # Extract skills from entire resume text using comprehensive skills database
            if SKILLS_DATABASE_AVAILABLE:
                extracted_skills = extract_skills_from_text(text)
                if extracted_skills:
                    # Format skills as comma-separated string
                    skills_text = ', '.join(extracted_skills[:30])  # Limit to top 30 skills
                    if not sections.get('skills') or len(sections.get('skills', '').strip()) < len(skills_text):
                        sections['skills'] = skills_text
                        print(f"DEBUG: Enhanced skill extraction found {len(extracted_skills)} skills using skills database")
        
        # Final Gemini enhancement pass for any weak sections
        if use_llm_enabled and LLM_AVAILABLE and is_llm_available():
            weak_sections = []
            for section_name, min_len in [('summary', 80), ('education', 80), ('skills', 60), ('experience', 120)]:
                section_value = sections.get(section_name, '')
                if not section_value or len(section_value.strip()) < min_len:
                    weak_sections.append(section_name)
            if weak_sections:
                try:
                    llm_client = get_llm_client()
                    for target_section in weak_sections:
                        enhanced = llm_client.enhance_resume_parsing(text, target_section=target_section)
                        enhanced_text = enhanced.get(target_section)
                        if enhanced_text and len(enhanced_text.strip()) > len(sections.get(target_section, '').strip()):
                            sections[target_section] = enhanced_text.strip()
                            print(f"DEBUG: Targeted Gemini enhancement improved {target_section} section")
                except Exception as e:
                    print(f"DEBUG: Targeted Gemini enhancement failed: {e}")
        
        # Debug: Print parsed sections keys and content lengths
        print(f"DEBUG: Parsed sections keys: {list(sections.keys())}")
        for key, content in sections.items():
            if isinstance(content, list):
                str_items = []
                for item in content:
                    if isinstance(item, dict):
                        dict_str = ' '.join(f"{k}: {v}" for k, v in item.items())
                        str_items.append(dict_str)
                    else:
                        str_items.append(str(item))
                joined_content = ' '.join(str_items)
            else:
                joined_content = content or ''
            print(f"DEBUG: {key}: {len(joined_content)} characters")
            if joined_content.strip():
                print(f"DEBUG: {key} preview: {joined_content[:100]}...")
        
        # Always include raw text for fallback extraction
        sections['raw_text'] = text
        
        # Store placeholder O*NET expansion data for downstream consumers
        if self.use_onet_taxonomy:
            sections['_onet_skill_clusters'] = self._expand_skills_with_onet(sections.get('skills'))
        
        return sections
    
    # ------------------------------------------------------------------
    # O*NET scaffolding helpers
    # ------------------------------------------------------------------
    def _coerce_skill_list(self, skills_entry: Any) -> List[str]:
        """Normalize skill content into a list so expansion can iterate safely."""
        if not skills_entry:
            return []
        if isinstance(skills_entry, list):
            return [str(item).strip() for item in skills_entry if str(item).strip()]
        if isinstance(skills_entry, str):
            parts = re.split(r"[,;/\n]", skills_entry)
            return [part.strip() for part in parts if part.strip()]
        return [str(skills_entry).strip()]

    def _expand_skills_with_onet(self, skills_entry: Any) -> Dict[str, Dict[str, List[str]]]:
        """
        Expand a collection of skills via O*NET when enabled.
        Returns a mapping of original skill -> expanded cluster data.
        """
        if not (self.use_onet_taxonomy and self.onet_taxonomy):
            return {}
        skill_list = self._coerce_skill_list(skills_entry)
        expanded: Dict[str, Dict[str, List[str]]] = {}
        for skill in skill_list:
            if not skill:
                continue
            expanded[skill] = self.onet_taxonomy.expand_skill(skill)
        return expanded
    
    def extract_sections_enhanced(self, file_path: str) -> Dict[str, str]:
        """Extract sections using enhanced parser with better format detection."""
        if not ENHANCED_PARSER_AVAILABLE:
            # Fallback to original method
            text = self.processor.extract_text(file_path)
            return self.extract_sections(text)
        
        try:
            # Use enhanced parser
            parsed_data = parse_resume(file_path)
            
            # Use the improved section extraction method
            parser = get_parser()
            sections = parser._extract_sections_from_parsed_data(parsed_data)
            
            # Also try the new parse_resume_sections method for better summary detection
            text = self.processor.extract_text(file_path)
            enhanced_sections = parser.parse_resume_sections(text, str(file_path))
            
            # Merge results, prioritizing enhanced_sections for summary
            if enhanced_sections.get('summary'):
                sections['summary'] = enhanced_sections['summary']
                print(f"DEBUG: Enhanced summary detection found: {enhanced_sections['summary'][:100]}...")
            
            # Store parsed data for later use
            self._cached_parsed_data = parsed_data
            
            # Add raw text for fallback purposes
            sections['raw_text'] = text
            
            return sections
            
        except Exception as e:
            print(f"Enhanced parser failed: {e}")
            # Fallback to original method
            text = self.processor.extract_text(file_path)
            sections = self.extract_sections(text)
            # Add raw text for fallback purposes
            sections['raw_text'] = text
            return sections
    
    def _clean_resume_text(self, text: str) -> str:
        """Clean and normalize resume text for better section detection."""
        # Handle case where text is all on one line (common with PDF extraction)
        if '\n' not in text or len(text.split('\n')) < 3:
            print("DEBUG: PDF text appears to be single line, applying enhanced splitting")
            
            # More comprehensive splitting for PDF text
            separators = [
                'Experience', 'Skills', 'Education', 'Profile', 'Contact',
                'Professional Summary', 'Work Experience', 'Technical Skills',
                'Academic Background', 'Career Objective', 'Summary',
                'Employment', 'Work History', 'Professional Experience',
                '•', 'o', '-', '–', '▪', '▫'
            ]
            
            for sep in separators:
                text = text.replace(sep, f'\n{sep}\n')
            
            # Also split on common patterns
            text = re.sub(r'([A-Z][a-z]+ [A-Z][a-z]+)', r'\n\1\n', text)  # Names like "Daniel Park"
            text = re.sub(r'(\d{4}–\d{4}|\d{4}-\d{4})', r'\n\1\n', text)  # Date ranges
            text = re.sub(r'([A-Z]{2,})', r'\n\1\n', text)  # All caps words
        
        return text
    
    def _detect_section_header(self, line_lower: str) -> str:
        """Detect if a line is a section header - must be a clear header, not content."""
        # Skip if line looks like contact info (email, phone, URL, etc.)
        if '@' in line_lower or re.search(r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', line_lower):
            return None
        
        # Skip if line looks like a name (starts with capital, has parentheses, short)
        if len(line_lower.split()) <= 4 and '(' in line_lower and line_lower[0].isupper():
            return None
        
        # Summary section - must be a clear header
        summary_patterns = [
            'professional summary', 'summary of qualifications',
            'professional overview', 'about me', 'objective', 'career objective',
            'executive summary', 'personal summary', 'career summary',
            'professional profile', 'personal statement'
        ]
        # "Profile" and "Summary" alone are too ambiguous - require more context
        if line_lower.strip() in ['profile', 'summary']:
            return 'summary'  # Allow these as they're common headers
        elif any(pattern in line_lower for pattern in summary_patterns):
            return 'summary'
        
        # Education section - must be a clear header or contain education keywords
        # Be more strict: require "Education" as a standalone word or with education degree keywords
        if line_lower.strip() == 'education' or line_lower.strip() == 'academic':
            return 'education'
        # Also match if it contains education degree keywords (but not just any word)
        education_degree_patterns = [
            r'\b(bachelor|master|phd|ph\.d|mba|b\.s|b\.a|m\.s|m\.a|bs|ba|ms|ma|degree|diploma)\b',
            r'\b(university|college)\b.*\b(graduated|degree)\b',
            r'\b(graduated|degree)\b.*\b(university|college)\b'
        ]
        if any(re.search(pattern, line_lower) for pattern in education_degree_patterns):
            return 'education'
        
        # Skills section - must be a clear header
        if line_lower.strip() in ['skills', 'technical skills', 'core competencies', 'expertise']:
            return 'skills'
        # Also match if it's clearly a skills header
        skills_header_patterns = [
            'technical skills', 'core competencies', 'technical expertise',
            'programming languages', 'technologies', 'tools & technologies'
        ]
        if any(pattern in line_lower for pattern in skills_header_patterns):
            return 'skills'
        
        # Experience section - must be a clear header
        experience_patterns = [
            'experience', 'work', 'employment', 'career', 'professional', 'work history', 
            'employment history', 'position', 'role', 'job', 'company', 'engineer', 'analyst',
            'manager', 'director', 'senior', 'junior', 'lead', 'principal'
        ]
        if any(pattern in line_lower for pattern in experience_patterns):
            return 'experience'
        
        return None
    
    def _is_contact_info(self, line_lower: str) -> bool:
        """Check if a line contains contact information."""
        contact_keywords = [
            'contact', 'email', 'phone', 'github', 'linkedin', 'website', 
            'address', 'location', '@', 'http', 'www', 'linkedin.com', 'github.com'
        ]
        
        # Only flag as contact info if it contains actual contact patterns
        has_contact_pattern = any(keyword in line_lower for keyword in contact_keywords)
        
        # Also check for email patterns
        has_email = '@' in line_lower and '.' in line_lower
        
        # Check for phone patterns
        has_phone = re.search(r'\+?\d{1,3}[-.\s]?\d{3,4}[-.\s]?\d{3,4}', line_lower)
        
        return has_contact_pattern or has_email or bool(has_phone)
    
    def _should_include_content(self, section: str, line_lower: str, original_line: str) -> bool:
        """Determine if content should be included in a section."""
        # Skip contact information
        if self._is_contact_info(line_lower):
            return False
        
        # Skip very short lines (less than 3 characters)
        if len(original_line.strip()) < 3:
            return False
        
        # Skip lines that are just section headers
        if line_lower in ['experience', 'skills', 'education', 'summary', 'profile', 'contact']:
            return False
        
        # For PDFs, be more lenient - include most content unless it's clearly contact info
        # This helps with PDFs where content might not match exact patterns
        
        # For experience section, be more flexible
        if section == 'experience':
            # Include if it has job-related content OR if it's substantial content (for PDFs)
            job_indicators = [
                'engineer', 'analyst', 'scientist', 'developer', 'manager', 'director',
                'implemented', 'built', 'migrated', 'decreased', 'increased', 'created',
                'designed', 'developed', 'managed', 'led', 'improved', 'optimized',
                'data', 'software', 'system', 'platform', 'pipeline', 'model'
            ]
            has_job_content = any(indicator in line_lower for indicator in job_indicators)
            is_substantial = len(original_line.strip()) > 20  # Substantial content for PDFs
            return has_job_content or is_substantial
        
        # For skills section, be more flexible
        elif section == 'skills':
            skill_indicators = [
                'python', 'java', 'sql', 'javascript', 'react', 'aws', 'docker',
                'kubernetes', 'spark', 'kafka', 'airflow', 'snowflake', 'bigquery',
                'machine learning', 'data science', 'etl', 'streaming', 'tool', 'technology'
            ]
            has_skill_content = any(indicator in line_lower for indicator in skill_indicators)
            is_substantial = len(original_line.strip()) > 15  # Substantial content for PDFs
            return has_skill_content or is_substantial
        
        # For education section, be very flexible (often missing, so cast a wide net)
        elif section == 'education':
            education_indicators = [
                'bachelor', 'master', 'phd', 'ph.d', 'university', 'college', 'degree',
                'diploma', 'ms', 'bs', 'mba', 'computer science', 'data science', 'graduated',
                'b.s', 'b.a', 'm.s', 'm.a', 'ba', 'ma', 'academic', 'education', 'school',
                'institute', 'major', 'minor', 'gpa', 'honors', 'cum laude', 'certification'
            ]
            has_education_content = any(indicator in line_lower for indicator in education_indicators)
            # Also include lines with year patterns (likely graduation dates)
            has_year_pattern = bool(re.search(r'\b(19|20)\d{2}\b', original_line))
            # Include lines that are short and look like degree info
            is_degree_like = len(original_line.split()) <= 15 and any(
                word in line_lower for word in ['in', 'of', 'from', 'university', 'college']
            )
            is_substantial = len(original_line.strip()) > 8  # Lower threshold for education
            return has_education_content or has_year_pattern or (is_degree_like and is_substantial)
        
        # For summary section, include most content
        elif section == 'summary':
            return len(original_line.strip()) > 10
        
        return True

    def _extract_summary_fallback(self, text: str) -> str:
        """Heuristic summary extraction when explicit section not found."""
        lines = text.split('\n')
        summary_lines = []
        for line in lines:
            clean_line = line.strip()
            if not clean_line:
                if summary_lines:
                    summary_lines.append('')
                continue
            line_lower = clean_line.lower()
            # Stop when we hit another major section header
            detected_section = self._detect_section_header(line_lower)
            if detected_section in ['experience', 'education', 'skills']:
                break
            if self._is_contact_info(line_lower):
                continue
            if len(clean_line) < 5:
                continue
            summary_lines.append(clean_line)
            if len(" ".join(summary_lines)) > 1000:
                break
        summary_text = " ".join(summary_lines).strip()
        if len(summary_text) >= 80:
            return summary_text[:1200]
        return ''

    def _extract_education_fallback(self, text: str) -> str:
        """Heuristic education extraction without explicit headers."""
        lines = text.split('\n')
        education_blocks = []
        current_block = []
        education_keywords = [
            'bachelor', 'master', 'phd', 'ph.d', 'mba', 'degree', 'diploma',
            'university', 'college', 'school', 'academy', 'institute', 'certification',
            'certified', 'coursework', 'major', 'minor', 'gpa', 'honors'
        ]
        school_keywords = ['university', 'college', 'school', 'academy', 'institute']
        
        for i, line in enumerate(lines):
            clean_line = line.strip()
            if not clean_line:
                if current_block:
                    education_blocks.append(" ".join(current_block))
                    current_block = []
                continue
            lower = clean_line.lower()
            if self._detect_section_header(lower) == 'experience':
                if education_blocks or current_block:
                    break
            keyword_hit = any(keyword in lower for keyword in education_keywords) or re.search(r'\b(19|20)\d{2}\b', clean_line)
            if keyword_hit:
                # Include preceding line if it looks like a school name
                if i > 0:
                    prev_line = lines[i-1].strip()
                    if prev_line and any(k in prev_line.lower() for k in school_keywords):
                        if not current_block or current_block[-1] != prev_line:
                            current_block.append(prev_line)
                current_block.append(clean_line)
                continue
            if current_block:
                education_blocks.append(" ".join(current_block))
                current_block = []
        if current_block:
            education_blocks.append(" ".join(current_block))
        education_text = " ".join(block for block in education_blocks if len(block) > 20)
        education_text = education_text.strip()
        if education_text:
            return education_text[:1200]
        return ''
    

    def _fallback_section_from_raw_text(self, section: str, raw_text: Optional[str], resume_name: str = "") -> str:
        """Generate fallback content for a section directly from raw resume text."""
        if not raw_text or len(raw_text.strip()) < 20:
            print(f"ERROR: Raw text is too short for {resume_name or 'resume'} ({len(raw_text) if raw_text else 0} chars), cannot extract {section}")
            return ''
        
        keywords = {
            'education': [
                'university', 'college', 'degree', 'bachelor', 'master', 'phd', 'ph.d', 'education',
                'b.s', 'b.a', 'm.s', 'm.a', 'mba', 'ms', 'bs', 'ba', 'ma', 'graduated', 'graduate',
                'diploma', 'certificate', 'certification', 'school', 'institute', 'academic',
                'major', 'minor', 'gpa', 'honors', 'cum laude', 'magna cum laude'
            ],
            'skills': ['skill', 'proficient', 'expert', 'experience with', 'knowledge of'],
            'experience': ['experience', 'worked', 'engineer', 'developer', 'manager', 'role']
        }
        
        section_keywords = keywords.get(section, [])
        lines = raw_text.split('\n')
        relevant_lines = [line for line in lines if any(kw in line.lower() for kw in section_keywords)]
        
        if relevant_lines:
            if section == 'education':
                enhanced_lines = []
                for i, line in enumerate(lines):
                    if any(kw in line.lower() for kw in section_keywords):
                        enhanced_lines.append(line)
                        for j in range(1, 3):
                            if i + j < len(lines) and lines[i + j].strip():
                                enhanced_lines.append(lines[i + j])
                return ' '.join(enhanced_lines[:8])
            return ' '.join(relevant_lines[:5])
        
        # If no keywords found, use first portion of raw text as absolute last resort
        limit = 1000 if section == 'experience' else 500
        fallback_text = raw_text[:limit] if len(raw_text) > limit else raw_text
        print(f"WARNING: No keywords found for {section} in {resume_name or 'resume'}, using raw text fallback ({len(fallback_text)} chars)")
        return fallback_text

    def _structure_education_entries(self, education_text: str) -> List[Dict[str, str]]:
        """Normalize education text into structured entries with school, degree, major, and year."""
        if not education_text:
            return []
        
        split_pattern = r'[\n•\-\u2022;]+'
        chunks = [chunk.strip(" .,\t") for chunk in re.split(split_pattern, education_text) if chunk.strip()]
        
        degree_pattern = re.compile(
            r'(Associate|Bachelor|Master|MBA|Doctor|Ph\.?D|B\.?\s?S\.?|B\.?\s?A\.?|M\.?\s?S\.?|M\.?\s?A\.?|BSc|MSc|BEng|MEng|Diploma|Certificate)[^,;\n]*',
            re.IGNORECASE
        )
        major_pattern = re.compile(r'\b(?:in|of)\s+([A-Za-z0-9&\/\-\s]+)')
        school_pattern = re.compile(
            r'([A-Z][A-Za-z0-9&.,\'\s]+(?:University|College|Institute|School|Academy|Polytechnic))',
            re.IGNORECASE
        )
        year_pattern = re.compile(r'(19|20)\d{2}')
        
        # Combine chunks so each entry contains degree + school context
        combined_entries = []
        buffer = ""
        for chunk in chunks:
            if degree_pattern.search(chunk) or school_pattern.search(chunk):
                if buffer:
                    combined_entries.append(buffer.strip(" ,"))
                buffer = chunk
            else:
                buffer = f"{buffer} {chunk}".strip()
        if buffer:
            combined_entries.append(buffer.strip(" ,"))
        
        structured_entries = []
        for entry_text in combined_entries:
            entry = {'school': '', 'degree': '', 'major': '', 'year': ''}
            
            degree_match = degree_pattern.search(entry_text)
            if degree_match:
                entry['degree'] = degree_match.group(0).strip()
            
            major_match = major_pattern.search(entry_text)
            if major_match:
                potential_major = major_match.group(1).strip().rstrip('.,)')
                if len(potential_major) > 3 and not re.search(r'(university|college|school|academy)', potential_major, re.IGNORECASE):
                    entry['major'] = potential_major
            
            school_match = school_pattern.search(entry_text)
            if school_match:
                entry['school'] = school_match.group(0).strip()
            
            year_match = year_pattern.search(entry_text)
            if year_match:
                entry['year'] = year_match.group(0)
            
            if any(entry.values()):
                structured_entries.append(entry)
        
        return structured_entries
    
    def extract_jd_requirements_with_importance(self, jd_text: str) -> Dict[str, List[Tuple[str, float]]]:
        """Extract requirements with importance levels from job description using Gemini API if available."""
        requirements = {
            'education': [],
            'skills': [],
            'experience': []
        }
        
        # Try Gemini API first if available (only if enabled)
        use_llm_enabled = st.session_state.get('use_llm', False)
        if use_llm_enabled and LLM_AVAILABLE and is_llm_available():
            try:
                llm_client = get_llm_client()
                gemini_reqs = llm_client.extract_jd_requirements_enhanced(jd_text)
                if gemini_reqs:
                    # Convert Gemini format to our format with importance levels
                    for req_type in ['education', 'skills', 'experience']:
                        if req_type in gemini_reqs and gemini_reqs[req_type]:
                            for req in gemini_reqs[req_type]:
                                # Determine importance (default to 1.0, can be enhanced)
                                importance = extract_importance_level(req)
                                requirements[req_type].append((req, importance))
                    # If we got good results from Gemini, return them
                    if any(requirements.values()):
                        print(f"DEBUG: Using Gemini-extracted JD requirements")
                        return requirements
            except Exception as e:
                print(f"DEBUG: Gemini JD extraction failed: {e}, falling back to rule-based")
        
        # Fallback to rule-based extraction
        lines = jd_text.split('\n')
        current_section = None
        
        # Education keywords for content detection (not just headers)
        education_keywords = ['bachelor', 'master', 'phd', 'ph.d', 'mba', 'degree', 'diploma', 
                             'university', 'college', 'education', 'graduate', 'undergraduate',
                             'bs', 'ba', 'ms', 'ma', 'b.s', 'b.a', 'm.s', 'm.a']
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Detect section headers
            if any(word in line_lower for word in ['education', 'degree', 'qualification']):
                current_section = 'education'
                continue
            elif any(word in line_lower for word in ['skill', 'technical', 'requirement']):
                current_section = 'skills'
                continue
            elif any(word in line_lower for word in ['experience', 'work', 'employment']):
                current_section = 'experience'
                continue
            
            # Extract requirements with importance
            if current_section and line.strip():
                importance = extract_importance_level(line)
                requirements[current_section].append((line.strip(), importance))
            elif not current_section and line.strip():
                # If no current section, check if line contains education keywords
                # This helps capture education requirements that aren't under explicit headers
                if any(keyword in line_lower for keyword in education_keywords):
                    # Check if it's actually about education (not just mentioning a degree in passing)
                    # Look for patterns like "Bachelor's degree", "Master's in", "PhD required", etc.
                    education_patterns = [
                        r'\b(bachelor|master|phd|ph\.d|mba)\s+(degree|in|of)',
                        r'\b(degree|diploma)\s+(in|of|from)',
                        r'\b(university|college)\s+(degree|education)',
                        r'education.*required',
                        r'degree.*required'
                    ]
                    if any(re.search(pattern, line_lower) for pattern in education_patterns):
                        importance = extract_importance_level(line)
                        requirements['education'].append((line.strip(), importance))
        
        if self.use_onet_taxonomy:
            skills_only = [req[0] for req in requirements.get('skills', [])]
            requirements['_onet_skill_clusters'] = self._expand_skills_with_onet(skills_only)
        
        return requirements
    
    def compute_semantic_similarity(self, education_weight: float = 0.1, skills_weight: float = 0.4, experience_weight: float = 0.4, summary_weight: float = 0.2):
        """Compute semantic similarity using improved weighted embedding algorithm."""
        if not self.model:
            st.error("SentenceTransformer not available. Please install sentence-transformers.")
            return None
        
        start_time = time.time()
        
        # Check if we have cached sections to ensure consistency across reruns
        resume_keys_hash = hashlib.sha256(str(sorted(self.resumes.keys())).encode()).hexdigest()
        sections_cache_key = f"sections_cache_{resume_keys_hash}"
        
        resume_sections = {}
        if sections_cache_key in st.session_state:
            cached_sections = st.session_state[sections_cache_key]
            # Verify cache is still valid (same number of resumes)
            if len(cached_sections) == len(self.resumes) and all(name in cached_sections for name in self.resumes.keys()):
                st.info("📋 Using cached section extraction for consistent results")
                resume_sections = cached_sections
            else:
                # Cache invalid, clear it
                del st.session_state[sections_cache_key]
        
        jd_requirements = {}
        
        # Extract sections from all documents with API enhancement and visual feedback
        # Only if not cached
        if not resume_sections:
            resume_sections = {}
        
        # Progress tracking for resume extraction
        total_resumes = len(self.resumes)
        if total_resumes > 0 and not resume_sections:
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for idx, (name, text) in enumerate(self.resumes.items()):
                status_text.text(f"📄 Extracting sections from Resume {idx + 1}/{total_resumes}: {name}")
                progress_bar.progress((idx + 1) / total_resumes)
                
                # Validate text extraction first
                if not text or len(text.strip()) < 20:
                    st.warning(f"⚠️ Resume '{name}' has very little extracted text ({len(text) if text else 0} chars). This may cause low similarity scores.")
                    print(f"WARNING: Resume '{name}' text extraction may have failed. Text length: {len(text) if text else 0}")
                
            # Try enhanced parser first, fallback to original
                sections = {}
                api_used = False
                original_path = None  # Initialize before conditional block
                
                if (
                    ENHANCED_PARSER_AVAILABLE
                    and self.data_dir is not None
                    and hasattr(self, 'original_filenames')
                ):
                    original_path = self.original_filenames.get(name)
                
                if original_path:
                    sections = self.extract_sections_enhanced(str(self.data_dir / original_path))
                    # Store enhanced data for later use
                    if hasattr(self, '_cached_parsed_data'):
                        self._enhanced_data_cache = getattr(self, '_enhanced_data_cache', {})
                        self._enhanced_data_cache[name] = self._cached_parsed_data
            else:
                sections = self.extract_sections(text)
                
                # Enhance with Gemini API if available and sections are weak (only if enabled)
                use_llm_enabled = st.session_state.get('use_llm', False)
                if use_llm_enabled and LLM_AVAILABLE and is_llm_available():
                    # Check if sections need enhancement (empty or very short)
                    # Prioritize education extraction since it's often missing
                    needs_enhancement = any(
                        not sections.get(section) or len(sections.get(section, '').strip()) < 20
                        for section in ['education', 'skills', 'experience']
                    )
                    
                    # Specifically check education - it's critical and often missing
                    education_missing = not sections.get('education') or len(sections.get('education', '').strip()) < 15
                    
                    if needs_enhancement:
                        try:
                            llm_client = get_llm_client()
                            with st.spinner(f"🤖 Using AI to enhance extraction for {name}..."):
                                # If education is missing, prioritize it
                                if education_missing:
                                    # First try to extract just education
                                    edu_result = llm_client.enhance_resume_parsing(text, target_section="education")
                                    if edu_result.get('education') and len(edu_result['education'].strip()) > 15:
                                        sections['education'] = edu_result['education']
                                        api_used = True
                                        print(f"DEBUG: Enhanced education using Gemini API ({len(edu_result['education'])} chars)")
                                
                                # Then do full extraction for other sections
                                gemini_sections = llm_client.enhance_resume_parsing(text, target_section="all")
                                
                                # Merge Gemini results with existing (prefer Gemini if better)
                                for section in ['education', 'skills', 'experience', 'summary']:
                                    if section in gemini_sections and len(gemini_sections[section].strip()) > 20:
                                        # Use Gemini result if it's better
                                        if not sections.get(section) or len(sections.get(section, '').strip()) < len(gemini_sections[section].strip()):
                                            sections[section] = gemini_sections[section]
                                            api_used = True
                        except Exception as e:
                            print(f"DEBUG: Gemini enhancement failed for {name}: {e}")
                
                # Ensure all required sections exist and have content
                if 'raw_text' not in sections:
                    sections['raw_text'] = text
                if 'summary' not in sections:
                    sections['summary'] = ''
                
                # CRITICAL: Ensure raw_text is always available and substantial
                raw_text = sections.get('raw_text', text)
                if not raw_text or len(raw_text.strip()) < 20:
                    # If raw_text is missing or too short, use the original text
                    raw_text = text
                    sections['raw_text'] = text
                    print(f"WARNING: Raw text was missing/empty for {name}, using original text ({len(text)} chars)")
                
                # Final fallback: if section is still empty, use raw text heuristics
                for section in ['education', 'skills', 'experience']:
                    if not sections.get(section) or len(sections.get(section, '').strip()) < 10:
                        fallback_value = self._fallback_section_from_raw_text(
                            section=section,
                            raw_text=raw_text,
                            resume_name=name
                        )
                        if fallback_value:
                            sections[section] = fallback_value
                
                # Final validation: ensure at least one section has substantial content
                has_content = any(
                    sections.get(section) and len(sections.get(section, '').strip()) >= 20
                    for section in ['education', 'skills', 'experience', 'summary']
                )
                if not has_content and raw_text and len(raw_text.strip()) >= 20:
                    # Emergency fallback: use raw text for all sections if nothing else worked
                    print(f"CRITICAL: No sections extracted for {name}, using raw text for all sections")
                    sections['experience'] = raw_text[:1500] if len(raw_text) > 1500 else raw_text
                    sections['skills'] = raw_text[:800] if len(raw_text) > 800 else raw_text
                    sections['education'] = raw_text[:500] if len(raw_text) > 500 else raw_text
                
                # Store API usage info
                if api_used:
                    sections['_api_enhanced'] = True
                
                # DIAGNOSTIC: Log section extraction results
                print(f"\n{'='*60}")
                print(f"EXTRACTION DIAGNOSTIC for {name}:")
                print(f"{'='*60}")
                print(f"Raw text: {len(sections.get('raw_text', ''))} chars")
                print(f"Education: {len(sections.get('education', ''))} chars - '{sections.get('education', '')[:50]}...'")
                print(f"Skills: {len(sections.get('skills', ''))} chars - '{sections.get('skills', '')[:50]}...'")
                print(f"Experience: {len(sections.get('experience', ''))} chars - '{sections.get('experience', '')[:50]}...'")
                print(f"Summary: {len(sections.get('summary', ''))} chars - '{sections.get('summary', '')[:50]}...'")
                
                # Check if extraction was successful
                sections_with_content = sum(1 for section in ['education', 'skills', 'experience', 'summary'] 
                                          if sections.get(section) and len(sections.get(section, '').strip()) >= 10)
                print(f"Sections with content (>=10 chars): {sections_with_content}/4")
                
                if sections_with_content == 0 and len(sections.get('raw_text', '').strip()) < 20:
                    st.error(f"❌ **EXTRACTION FAILED**: Resume '{name}' - No sections extracted and raw text is too short!")
                    print(f"ERROR: Extraction completely failed for {name}")
                elif sections_with_content == 0:
                    st.warning(f"⚠️ Resume '{name}' - No sections extracted, will use raw text fallback")
                    print(f"WARNING: No sections extracted for {name}, will rely on raw text")
                print(f"{'='*60}\n")
            
            resume_sections[name] = sections
            
            # Cache sections in session state for consistency
            resume_keys_hash = hashlib.sha256(str(sorted(self.resumes.keys())).encode()).hexdigest()
            sections_cache_key = f"sections_cache_{resume_keys_hash}"
            st.session_state[sections_cache_key] = resume_sections
            
            progress_bar.empty()
            status_text.empty()
        
        # Process JD requirements with API enhancement (cache for consistency)
        jd_requirements = {}
        jd_cache_key = None
        status_text = None
        
        if len(self.job_descriptions) > 0:
            jd_keys_hash = hashlib.sha256(str(sorted(self.job_descriptions.keys())).encode()).hexdigest()
            jd_cache_key = f"jd_requirements_cache_{jd_keys_hash}"
            
            if jd_cache_key in st.session_state:
                st.info("📋 Using cached JD requirements for consistent results")
                jd_requirements = st.session_state[jd_cache_key]
            else:
                status_text = st.empty()
                status_text.text("📋 Processing Job Descriptions with AI enhancement...")
        
        for name, text in self.job_descriptions.items():
            jd_requirements[name] = self.extract_jd_requirements_with_importance(text)
        
        # Cache JD requirements (only if we have job descriptions)
        if jd_cache_key and len(self.job_descriptions) > 0:
            st.session_state[jd_cache_key] = jd_requirements
        
        # Clear status text if it was created
        if status_text is not None:
            status_text.empty()
        
        if self.use_onet_taxonomy:
            for sections in resume_sections.values():
                if '_onet_skill_clusters' not in sections:
                    sections['_onet_skill_clusters'] = self._expand_skills_with_onet(sections.get('skills'))
            for jd_data in jd_requirements.values():
                if isinstance(jd_data, dict) and '_onet_skill_clusters' not in jd_data:
                    skills_only = [req[0] for req in jd_data.get('skills', [])]
                    jd_data['_onet_skill_clusters'] = self._expand_skills_with_onet(skills_only)
        
        # Compute improved similarities using weighted embedding approach
        similarities = {}
        
        # Progress tracking for similarity computation
        # Use candidate_names length to ensure all resumes are included
        total_comparisons = len(jd_requirements) * len(self.candidate_names)
        current_comparison = 0
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for jd_name, jd_reqs in jd_requirements.items():
            similarities[jd_name] = {}
            
            # Combine JD requirements into single text (deterministic order)
            jd_text_parts = []
            for req_type in ['education', 'skills', 'experience']:  # Fixed order for consistency
                if req_type in jd_reqs and jd_reqs[req_type]:
                    jd_text_parts.extend([req[0] for req in jd_reqs[req_type]])
            jd_text = " ".join(jd_text_parts) if jd_text_parts else self.job_descriptions.get(jd_name, "")
            
            # Process ALL resumes from candidate_names, not just those in resume_sections
            # This ensures all resumes are processed even if section extraction failed
            for resume_name in sorted(self.candidate_names):
                # Ensure resume_sections has an entry for this resume (create empty if missing)
                if resume_name not in resume_sections:
                    print(f"WARNING: Resume '{resume_name}' not in resume_sections, creating empty entry")
                    # Try to get the resume text and create basic sections
                    resume_text = self.resumes.get(resume_name, '')
                    if resume_text:
                        resume_sections[resume_name] = {
                            'raw_text': resume_text,
                            'education': '',
                            'skills': '',
                            'experience': '',
                            'summary': ''
                        }
                    else:
                        # If no text available, create empty sections
                        resume_sections[resume_name] = {
                            'raw_text': '',
                            'education': '',
                            'skills': '',
                            'experience': '',
                            'summary': ''
                        }
                resume_sections_data = resume_sections[resume_name]
                current_comparison += 1
                status_text.text(f"🔍 Computing similarity: {resume_name} vs {jd_name} ({current_comparison}/{total_comparisons})")
                progress_bar.progress(current_comparison / total_comparisons)
                
                print(f"\nDEBUG: Computing semantic similarity for {resume_name} vs {jd_name}")
                
                # Ensure sections have content - use raw text as fallback
                # Only enhance if not already cached (to maintain consistency)
                for section in ['education', 'skills', 'experience']:
                    if not resume_sections_data.get(section) or len(resume_sections_data.get(section, '').strip()) < 10:
                        # Try to extract from raw text if section is empty (only if enabled)
                        raw_text = resume_sections_data.get('raw_text', '')
                        use_llm_enabled = st.session_state.get('use_llm', False)
                        if raw_text and use_llm_enabled and LLM_AVAILABLE and is_llm_available():
                            try:
                                llm_client = get_llm_client()
                                # Use cached API results if available
                                enhanced = llm_client.enhance_resume_parsing(raw_text, target_section=section)
                                if section in enhanced and len(enhanced[section].strip()) > 10:
                                    resume_sections_data[section] = enhanced[section]
                                    print(f"DEBUG: Enhanced {section} using Gemini API")
                            except Exception as e:
                                print(f"DEBUG: Failed to enhance {section}: {e}")
                        
                        # Fallback: Use skills database for skills extraction if still empty
                        if section == 'skills' and (not resume_sections_data.get('skills') or len(resume_sections_data.get('skills', '').strip()) < 10):
                            raw_text = resume_sections_data.get('raw_text', '')
                            if raw_text and SKILLS_DATABASE_AVAILABLE:
                                extracted_skills = extract_skills_from_text(raw_text)
                                if extracted_skills:
                                    skills_text = ', '.join(extracted_skills[:30])
                                    resume_sections_data['skills'] = skills_text
                                    print(f"DEBUG: Enhanced skills extraction using skills database: {len(extracted_skills)} skills found")
                
                # Generate professional summary from experience (cached by LLM client)
                generated_summary = self.generate_professional_summary(resume_sections_data)
                print(f"DEBUG: Generated professional summary for {resume_name}: {generated_summary[:100]}...")
                
                # Update resume sections with generated summary
                resume_sections_data['summary'] = generated_summary
                
                # DIAGNOSTIC: Check resume sections before similarity computation
                print(f"\n{'='*80}")
                print(f"DIAGNOSTIC: {resume_name} vs {jd_name}")
                print(f"{'='*80}")
                print(f"Raw text length: {len(resume_sections_data.get('raw_text', ''))} chars")
                print(f"Education length: {len(resume_sections_data.get('education', ''))} chars")
                print(f"Skills length: {len(resume_sections_data.get('skills', ''))} chars")
                print(f"Experience length: {len(resume_sections_data.get('experience', ''))} chars")
                print(f"Summary length: {len(resume_sections_data.get('summary', ''))} chars")
                
                # Check if we have any usable content
                has_content = any(
                    len(resume_sections_data.get(section, '').strip()) >= 10
                    for section in ['education', 'skills', 'experience', 'summary', 'raw_text']
                )
                print(f"Has usable content: {has_content}")

                if self.use_onet_taxonomy and self.onet_taxonomy:
                    resume_expanded = resume_sections_data.get('_onet_skill_clusters', {})
                    jd_expanded = jd_reqs.get('_onet_skill_clusters', {})
                    # Placeholder for future integration:
                    # use resume_expanded and jd_expanded to award partial credit
                    # for related O*NET cluster matches.
                
                if not has_content:
                    st.warning(f"⚠️ **Warning**: Resume '{resume_name}' has no extractable content! This will result in 0 similarity scores.")
                    print(f"ERROR: No usable content found for {resume_name}")
                
                # Compute weighted similarity using improved algorithm
                raw_sim, display_score = self.compute_weighted_similarity_with_custom_weights(
                    resume_sections_data, 
                    jd_text,
                    education_weight=education_weight,
                    skills_weight=skills_weight,
                    experience_weight=experience_weight,
                    summary_weight=summary_weight
                )
                
                # Get debug breakdown with section-specific JD requirements
                breakdown = self.debug_similarity_breakdown(resume_sections_data, jd_text, jd_reqs)
                
                # Store detailed results
                similarities[jd_name][resume_name] = {
                    'raw_similarity': raw_sim,
                    'display_score': display_score,
                    'total': display_score / 100,  # Convert to 0-1 range for compatibility
                    'breakdown': breakdown
                }
                
                print(f"DEBUG: Final result - Raw: {raw_sim:.4f}, Display: {display_score}%")
                
                # DIAGNOSTIC: Warn if score is 0
                if display_score == 0.0 or raw_sim == 0.0:
                    print(f"⚠️ WARNING: Zero similarity detected for {resume_name} vs {jd_name}")
                    print(f"   Raw similarity: {raw_sim}")
                    print(f"   Display score: {display_score}")
                    print(f"   This may indicate:")
                    print(f"   1. Text extraction failed (check raw_text length)")
                    print(f"   2. All sections are empty")
                    print(f"   3. Embedding generation failed")
                    print(f"   4. No matching content between resume and JD")
                print(f"{'='*80}\n")
        
        # Convert to matrix format using display scores
        similarity_matrix = np.zeros((len(self.candidate_names), len(self.jd_names)))
        section_matrices = {
            'education': np.zeros((len(self.candidate_names), len(self.jd_names))),
            'skills': np.zeros((len(self.candidate_names), len(self.jd_names))),
            'experience': np.zeros((len(self.candidate_names), len(self.jd_names))),
            'summary': np.zeros((len(self.candidate_names), len(self.jd_names)))
        }
        
        for jd_idx, jd_name in enumerate(self.jd_names):
            for resume_idx, resume_name in enumerate(self.candidate_names):
                if jd_name in similarities and resume_name in similarities[jd_name]:
                    # Use display score (0-100) converted to 0-1 range
                    similarity_matrix[resume_idx, jd_idx] = similarities[jd_name][resume_name]['display_score'] / 100
                    
                    # Store section breakdowns
                    breakdown = similarities[jd_name][resume_name]['breakdown']
                    for section in ['education', 'skills', 'experience', 'summary']:
                        # Always store the value, even if it's 0.0 (ensures all sections are populated)
                        if section in breakdown:
                            section_matrices[section][resume_idx, jd_idx] = breakdown[section]
                        else:
                            # If section is missing from breakdown, set to 0.0
                            section_matrices[section][resume_idx, jd_idx] = 0.0
                else:
                    # DIAGNOSTIC: Log when a resume-JD pair is missing from similarities
                    print(f"WARNING: Missing similarity data for {resume_name} vs {jd_name}")
                    print(f"  Available JDs in similarities: {list(similarities.keys())}")
                    if jd_name in similarities:
                        print(f"  Available resumes for {jd_name}: {list(similarities[jd_name].keys())}")
                    # Set to 0.0 if missing (this will show as 0 in the heatmap)
                    similarity_matrix[resume_idx, jd_idx] = 0.0
                    for section in ['education', 'skills', 'experience', 'summary']:
                            section_matrices[section][resume_idx, jd_idx] = 0.0
        
        progress_bar.empty()
        status_text.empty()
        
        end_time = time.time()
        st.success(f"✅ Semantic matching completed in {end_time - start_time:.2f} seconds")
        
        # Show extraction statistics
        api_enhanced_count = sum(1 for sections in resume_sections.values() if sections.get('_api_enhanced', False))
        if api_enhanced_count > 0:
            st.info(f"🤖 AI-enhanced extraction used for {api_enhanced_count}/{len(resume_sections)} resumes")
        
        # Store similarity matrix for get_top_matches method
        self.similarity_matrix = similarity_matrix
        
        return similarity_matrix, section_matrices, similarities
    
    def compute_weighted_embedding_similarity(self, summary_weight: float = 0.5, skills_weight: float = 0.3, experience_weight: float = 0.2):
        """Compute similarity using weighted embedding averaging for high-level alignment."""
        if not self.model:
            self.model = load_sentence_transformer()
            if not self.model:
                st.error("SentenceTransformer not available. Please install sentence-transformers.")
                return None
        
        start_time = time.time()
        
        # Extract sections from all documents
        resume_sections = {}
        jd_requirements = {}
        
        for name, text in self.resumes.items():
            # Try enhanced parser first, fallback to original
            if ENHANCED_PARSER_AVAILABLE and self.data_dir is not None:
                # Find the original file path for enhanced parsing (only works with directory mode)
                original_path = None
                if hasattr(self, 'original_filenames'):
                    original_path = self.original_filenames.get(name)
                
                if original_path:
                    sections = self.extract_sections_enhanced(str(self.data_dir / original_path))
                    # Store enhanced data for later use
                    if hasattr(self, '_cached_parsed_data'):
                        self._enhanced_data_cache = getattr(self, '_enhanced_data_cache', {})
                        self._enhanced_data_cache[name] = self._cached_parsed_data
                else:
                    sections = self.extract_sections(text)
            else:
                # Use regular extraction for file uploads or when enhanced parser not available
                sections = self.extract_sections(text)
            
            resume_sections[name] = sections
        
        for name, text in self.job_descriptions.items():
            jd_requirements[name] = self.extract_jd_requirements_with_importance(text)
        
        # Compute weighted embedding similarities
        similarities = {}
        
        for jd_name, jd_reqs in jd_requirements.items():
            similarities[jd_name] = {}
            
            for resume_name, resume_sections_data in resume_sections.items():
                # Create weighted embeddings for resume
                resume_embeddings = []
                resume_weights = []
                
                # Generate professional summary from experience
                generated_summary = self.generate_professional_summary(resume_sections_data)
                print(f"DEBUG: Generated professional summary for {resume_name}: {generated_summary[:100]}...")
                
                # Summary embedding (highest weight for high-level alignment)
                summary_text = generated_summary
                
                if summary_text:
                    summary_embedding = self._get_embedding(summary_text, 'summary', resume_sections_data)
                    resume_embeddings.append(summary_embedding)
                    resume_weights.append(summary_weight)
                    print(f"DEBUG: Added summary embedding for {resume_name} (weight: {summary_weight})")
                
                # Skills embedding (medium weight for technical alignment)
                skills_text = resume_sections_data.get('skills', '').strip()
                if skills_text:
                    skills_embedding = self._get_embedding(skills_text, 'skills', resume_sections_data)
                    resume_embeddings.append(skills_embedding)
                    resume_weights.append(skills_weight)
                    print(f"DEBUG: Added skills embedding for {resume_name} (weight: {skills_weight})")
                
                # Experience embedding (lower weight for detailed alignment)
                experience_text = resume_sections_data.get('experience', '').strip()
                if experience_text:
                    experience_embedding = self._get_embedding(experience_text, 'experience', resume_sections_data)
                    resume_embeddings.append(experience_embedding)
                    resume_weights.append(experience_weight)
                    print(f"DEBUG: Added experience embedding for {resume_name} (weight: {experience_weight})")
                
                if not resume_embeddings:
                    similarities[jd_name][resume_name] = {'total': 0.0, 'sections': {}}
                    continue
                
                # Normalize weights
                total_weight = sum(resume_weights)
                normalized_weights = [w / total_weight for w in resume_weights]
                
                # Compute weighted average embedding
                weighted_resume_embedding = np.average(resume_embeddings, axis=0, weights=normalized_weights)
                
                # Create weighted embeddings for JD
                jd_embeddings = []
                jd_weights = []
                
                # JD Summary embedding
                if jd_reqs.get('summary'):
                    jd_summary_text = " ".join([req[0] for req in jd_reqs['summary']])
                    jd_summary_embedding = self._get_embedding(jd_summary_text, 'summary_jd')
                    jd_embeddings.append(jd_summary_embedding)
                    jd_weights.append(summary_weight)
                
                # JD Skills embedding
                if jd_reqs.get('skills'):
                    jd_skills_text = " ".join([req[0] for req in jd_reqs['skills']])
                    jd_skills_embedding = self._get_embedding(jd_skills_text, 'skills_jd')
                    jd_embeddings.append(jd_skills_embedding)
                    jd_weights.append(skills_weight)
                
                # JD Experience embedding
                if jd_reqs.get('experience'):
                    jd_experience_text = " ".join([req[0] for req in jd_reqs['experience']])
                    jd_experience_embedding = self._get_embedding(jd_experience_text, 'experience_jd')
                    jd_embeddings.append(jd_experience_embedding)
                    jd_weights.append(experience_weight)
                
                if not jd_embeddings:
                    similarities[jd_name][resume_name] = {'total': 0.0, 'sections': {}}
                    continue
                
                # Normalize JD weights
                jd_total_weight = sum(jd_weights)
                jd_normalized_weights = [w / jd_total_weight for w in jd_weights]
                
                # Compute weighted average JD embedding
                weighted_jd_embedding = np.average(jd_embeddings, axis=0, weights=jd_normalized_weights)
                
                # Compute cosine similarity between weighted embeddings
                similarity = cosine_similarity([weighted_resume_embedding], [weighted_jd_embedding])[0][0]
                
                # Store results
                similarities[jd_name][resume_name] = {
                    'total': float(similarity),
                    'sections': {
                        'summary': summary_weight,
                        'skills': skills_weight,
                        'experience': experience_weight
                    }
                }
                
                print(f"DEBUG: Weighted embedding similarity for {resume_name} vs {jd_name}: {similarity:.4f}")
        
        # Convert to matrix format
        similarity_matrix = np.zeros((len(self.candidate_names), len(self.jd_names)))
        
        for jd_idx, jd_name in enumerate(self.jd_names):
            for resume_idx, resume_name in enumerate(self.candidate_names):
                if jd_name in similarities and resume_name in similarities[jd_name]:
                    similarity_matrix[resume_idx, jd_idx] = similarities[jd_name][resume_name]['total']
        
        end_time = time.time()
        st.info(f"🎯 Weighted embedding matching completed in {end_time - start_time:.2f} seconds")
        
        return similarity_matrix, similarities
    
    def scale_yoe(self, candidate_years: float, required_years: float) -> float:
        """Scale years of experience with linear interpolation."""
        if required_years <= 0:
            return 1.0
        
        ratio = min(candidate_years / required_years, 1.0)
        # Use bounded scaling: minimum 0.1, maximum 1.0
        return max(0.1, ratio)
    
    def extract_years_of_experience(self, text: str) -> float:
        """Extract years of experience from text using regex patterns."""
        
        # Common patterns for years of experience
        patterns = [
            r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*(?:of\s*)?(?:experience|exp)',
            r'(\d+(?:\.\d+)?)\s*\+?\s*yrs?\s*(?:of\s*)?(?:experience|exp)',
            r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*in',
            r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*working',
            r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*with',
            r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*of\s*professional',
            r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*in\s*the\s*field',
            r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*of\s*industry'
        ]
        
        years_found = []
        for pattern in patterns:
            matches = re.findall(pattern, text.lower())
            for match in matches:
                try:
                    years = float(match)
                    if 0.1 <= years <= 50:  # Reasonable range
                        years_found.append(years)
                except ValueError:
                    continue
        
        # Return the maximum years found (most relevant)
        return max(years_found) if years_found else 0.0
    
    def get_skill_similarity(self, skill1: str, skill2: str) -> float:
        """Calculate similarity between two skills using predefined mappings and embeddings."""
        skill1_lower = skill1.lower().strip()
        skill2_lower = skill2.lower().strip()
        
        # Exact match
        if skill1_lower == skill2_lower:
            return 1.0
        
        # Predefined skill similarity mappings
        skill_similarities = {
            # Data Processing & Streaming
            'kafka': {'flink': 0.8, 'kinesis': 0.8, 'pulsar': 0.7, 'rabbitmq': 0.6},
            'spark': {'pyspark': 0.9, 'hadoop': 0.7, 'flink': 0.8, 'storm': 0.6},
            'hadoop': {'spark': 0.7, 'hive': 0.8, 'pig': 0.6, 'mapreduce': 0.9},
            
            # Machine Learning Frameworks
            'pytorch': {'tensorflow': 0.8, 'keras': 0.7, 'scikit-learn': 0.6, 'mxnet': 0.6},
            'tensorflow': {'pytorch': 0.8, 'keras': 0.9, 'scikit-learn': 0.6, 'mxnet': 0.6},
            'keras': {'tensorflow': 0.9, 'pytorch': 0.7, 'scikit-learn': 0.6},
            
            # ML Platforms & Tools
            'mlflow': {'weights & biases': 0.7, 'neptune': 0.6, 'comet': 0.6, 'tensorboard': 0.5},
            'weights & biases': {'mlflow': 0.7, 'neptune': 0.8, 'comet': 0.7},
            'neptune': {'mlflow': 0.6, 'weights & biases': 0.8, 'comet': 0.7},
            
            # Cloud Platforms
            'aws': {'azure': 0.6, 'gcp': 0.6, 'google cloud': 0.6},
            'azure': {'aws': 0.6, 'gcp': 0.6, 'google cloud': 0.6},
            'gcp': {'aws': 0.6, 'azure': 0.6, 'google cloud': 0.9},
            'google cloud': {'aws': 0.6, 'azure': 0.6, 'gcp': 0.9},
            
            # Databases
            'postgresql': {'postgres': 0.9, 'mysql': 0.7, 'oracle': 0.6, 'sql server': 0.6},
            'mysql': {'postgresql': 0.7, 'postgres': 0.7, 'mariadb': 0.8},
            'mongodb': {'cassandra': 0.6, 'couchdb': 0.5, 'dynamodb': 0.6},
            'redis': {'memcached': 0.7, 'hazelcast': 0.6},
            
            # Programming Languages
            'python': {'r': 0.6, 'julia': 0.5, 'matlab': 0.5},
            'java': {'scala': 0.7, 'kotlin': 0.6, 'groovy': 0.6},
            'javascript': {'typescript': 0.8, 'node.js': 0.7, 'react': 0.6},
            'typescript': {'javascript': 0.8, 'node.js': 0.7},
            
            # Container & Orchestration
            'docker': {'podman': 0.8, 'containerd': 0.7},
            'kubernetes': {'docker swarm': 0.6, 'nomad': 0.5, 'rancher': 0.7},
            'helm': {'kustomize': 0.6, 'skaffold': 0.5},
            
            # CI/CD & DevOps
            'jenkins': {'gitlab ci': 0.7, 'github actions': 0.7, 'azure devops': 0.6},
            'gitlab ci': {'jenkins': 0.7, 'github actions': 0.8, 'azure devops': 0.7},
            'github actions': {'jenkins': 0.7, 'gitlab ci': 0.8, 'azure devops': 0.7},
            
            # Data Warehouses
            'snowflake': {'bigquery': 0.7, 'redshift': 0.7, 'databricks': 0.6},
            'bigquery': {'snowflake': 0.7, 'redshift': 0.6, 'databricks': 0.6},
            'redshift': {'snowflake': 0.7, 'bigquery': 0.6, 'databricks': 0.6},
            
            # Visualization
            'tableau': {'power bi': 0.7, 'looker': 0.6, 'qlik': 0.6},
            'power bi': {'tableau': 0.7, 'looker': 0.6, 'qlik': 0.6},
            'looker': {'tableau': 0.6, 'power bi': 0.6, 'qlik': 0.7},
            
            # Version Control
            'git': {'svn': 0.5, 'mercurial': 0.4},
            'github': {'gitlab': 0.8, 'bitbucket': 0.7},
            'gitlab': {'github': 0.8, 'bitbucket': 0.7}
        }
        
        # Check predefined mappings
        if skill1_lower in skill_similarities:
            if skill2_lower in skill_similarities[skill1_lower]:
                return skill_similarities[skill1_lower][skill2_lower]
        
        # Reverse check
        if skill2_lower in skill_similarities:
            if skill1_lower in skill_similarities[skill2_lower]:
                return skill_similarities[skill2_lower][skill1_lower]
        
        # If no predefined mapping, use embedding similarity as fallback
        if self.model:
            try:
                embedding1 = self._get_embedding(skill1, 'skill')
                embedding2 = self._get_embedding(skill2, 'skill')
                similarity = cosine_similarity([embedding1], [embedding2])[0][0]
                # Only return high similarity (>0.7) for embedding-based matches
                return similarity if similarity > 0.7 else 0.0
            except Exception:
                pass
        
        return 0.0
    
    def compute_graded_similarity(self, resume_sections: dict, jd_requirements: dict) -> dict:
        """Compute graded similarity with partial credit for preferred skills and YOE scaling."""
        scores = {
            'required_skills': 0.0,
            'preferred_skills': 0.0,
            'yoe_score': 0.0,
            'total_score': 0.0,
            'breakdown': []
        }
        
        # Extract years of experience from resume
        experience_text = resume_sections.get('experience', '') + ' ' + resume_sections.get('summary', '')
        candidate_yoe = self.extract_years_of_experience(experience_text)
        
        # Process required skills (binary gates)
        required_skills = jd_requirements.get('skills', [])
        required_matches = 0
        required_total = 0
        
        for skill_req, importance in required_skills:
            if importance >= 1.0:  # Required skill
                required_total += 1
                skill_matched = False
                
                # Check against resume skills
                resume_skills_text = resume_sections.get('skills', '') + ' ' + resume_sections.get('experience', '')
                resume_skills = resume_skills_text.lower()
                
                # Check for exact match or high similarity
                if skill_req.lower() in resume_skills:
                    required_matches += 1
                    skill_matched = True
                    scores['breakdown'].append(f"✅ {skill_req} (required): Exact match")
                else:
                    # Check for similar skills
                    for word in resume_skills.split():
                        word = word.strip('.,!?()[]{}')
                        similarity = self.get_skill_similarity(skill_req, word)
                        if similarity >= 0.8:  # High similarity threshold for required skills
                            required_matches += 1
                            skill_matched = True
                            scores['breakdown'].append(f"✅ {skill_req} (required): Similar to {word} (similarity: {similarity:.2f})")
                            break
                
                if not skill_matched:
                    scores['breakdown'].append(f"❌ {skill_req} (required): Not found")
        
        # Calculate required skills score
        if required_total > 0:
            scores['required_skills'] = required_matches / required_total
        else:
            scores['required_skills'] = 1.0  # No required skills
        
        # Process preferred skills (weighted bonuses)
        preferred_bonus = 0.0
        preferred_total = 0
        
        for skill_req, importance in required_skills:
            if importance < 1.0:  # Preferred/nice-to-have skill
                preferred_total += 1
                best_match_score = 0.0
                best_match_skill = ""
                
                # Check against resume skills
                resume_skills_text = resume_sections.get('skills', '') + ' ' + resume_sections.get('experience', '')
                resume_skills = resume_skills_text.lower()
                
                # Check for exact match
                if skill_req.lower() in resume_skills:
                    best_match_score = 1.0
                    best_match_skill = skill_req
                else:
                    # Check for similar skills
                    for word in resume_skills.split():
                        word = word.strip('.,!?()[]{}')
                        similarity = self.get_skill_similarity(skill_req, word)
                        if similarity > best_match_score:
                            best_match_score = similarity
                            best_match_skill = word
                
                # Add bonus based on match quality
                if best_match_score > 0.0:
                    bonus = best_match_score * importance  # Weight by importance level
                    preferred_bonus += bonus
                    scores['breakdown'].append(f"🎯 {skill_req} (preferred): Matched with {best_match_skill} (score: {bonus:.2f})")
                else:
                    scores['breakdown'].append(f"⚪ {skill_req} (preferred): Not found")
        
        # Calculate preferred skills score
        if preferred_total > 0:
            scores['preferred_skills'] = preferred_bonus / preferred_total
        else:
            scores['preferred_skills'] = 1.0  # No preferred skills
        
        # Process Years of Experience scaling
        yoe_requirements = []
        jd_text = ' '.join([req[0] for req in jd_requirements.get('experience', [])])
        yoe_patterns = [
            r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*(?:of\s*)?(?:experience|exp)',
            r'(\d+(?:\.\d+)?)\s*\+?\s*yrs?\s*(?:of\s*)?(?:experience|exp)',
            r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*preferred',
            r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s*required'
        ]
        
        for pattern in yoe_patterns:
            matches = re.findall(pattern, jd_text.lower())
            for match in matches:
                try:
                    required_yoe = float(match)
                    if 1 <= required_yoe <= 20:  # Reasonable range
                        yoe_requirements.append(required_yoe)
                except ValueError:
                    continue
        
        if yoe_requirements:
            # Use the maximum required YOE
            max_required_yoe = max(yoe_requirements)
            scores['yoe_score'] = self.scale_yoe(candidate_yoe, max_required_yoe)
            scores['breakdown'].append(f"📅 YOE: Candidate has {candidate_yoe:.1f} years, required {max_required_yoe:.1f} years (score: {scores['yoe_score']:.2f})")
        else:
            scores['yoe_score'] = 1.0  # No YOE requirement
            scores['breakdown'].append(f"📅 YOE: No specific requirement (candidate has {candidate_yoe:.1f} years)")
        
        # Calculate composite score
        # Formula: (required_match * 0.6) + (preferred_match * 0.3) + (yoe_scaled * 0.1)
        scores['total_score'] = (
            scores['required_skills'] * 0.6 +
            scores['preferred_skills'] * 0.3 +
            scores['yoe_score'] * 0.1
        )
        
        return scores
    
    def extract_experience_elements(self, experience_text: str) -> list:
        """Extract core elements from experience section."""
        
        experience_entries = []
        
        # Split by common separators (bullet points, line breaks, etc.)
        lines = re.split(r'[•\n\r]+', experience_text)
        
        for line in lines:
            line = line.strip()
            if not line or len(line) < 10:  # Skip very short lines
                continue
            
            # Extract job title and company using common patterns
            job_title = ""
            company = ""
            duration = ""
            achievements = []
            
            # Pattern 1: "Job Title, Company (2020-2023)" or "Job Title, Company (2020–Present)"
            title_company_pattern = r'([^,]+?),\s*([^(]+?)\s*\(([^)]+)\)'
            match = re.search(title_company_pattern, line, re.IGNORECASE)
            if match:
                job_title = match.group(1).strip()
                company = match.group(2).strip()
                duration = match.group(3).strip()
            else:
                # Pattern 2: "Job Title at Company (2020-2023)"
                title_company_pattern2 = r'([^,]+?)\s+(?:at|@)\s+([^(]+?)\s*\(([^)]+)\)'
                match = re.search(title_company_pattern2, line, re.IGNORECASE)
                if match:
                    job_title = match.group(1).strip()
                    company = match.group(2).strip()
                    duration = match.group(3).strip()
                else:
                    # Pattern 3: "Job Title, Company, Duration" (comma-separated)
                    parts = [p.strip() for p in line.split(',')]
                    if len(parts) >= 2:
                        job_title = parts[0]
                        company = parts[1]
                        if len(parts) >= 3:
                            duration = parts[2]
                        else:
                            # Pattern 4: Try to extract from line with year patterns
                            year_match = re.search(r'(\d{4})\s*[-–]\s*(\d{4}|present|current)', line, re.IGNORECASE)
                            if year_match:
                                duration = f"{year_match.group(1)}-{year_match.group(2)}"
                                # Try to extract job title and company
                                parts_before_year = line[:year_match.start()].strip()
                                if ',' in parts_before_year:
                                    title_company = [p.strip() for p in parts_before_year.split(',')]
                                    if len(title_company) >= 2:
                                        job_title = title_company[0]
                                        company = title_company[1]
                                    elif len(title_company) == 1:
                                        job_title = title_company[0]
            
            # Extract achievements (lines starting with action verbs)
            achievement_verbs = [
                'built', 'developed', 'led', 'implemented', 'deployed', 'created', 'designed',
                'managed', 'optimized', 'improved', 'increased', 'reduced', 'delivered',
                'launched', 'established', 'coordinated', 'supervised', 'mentored',
                'architected', 'scaled', 'automated', 'integrated', 'migrated', 'transformed'
            ]
            
            # Look for achievement statements in the same line or following lines
            achievement_pattern = r'(?:^|\s)(?:' + '|'.join(achievement_verbs) + r')\s+[^.!?]*(?:[.!?]|$)'
            achievement_matches = re.findall(achievement_pattern, line, re.IGNORECASE)
            
            for match in achievement_matches:
                achievement = match.strip()
                if len(achievement) > 15:  # Only keep substantial achievements
                    achievements.append(achievement)
            
            # Only add if we have meaningful information
            if job_title or company or achievements:
                experience_entries.append({
                    'job_title': job_title,
                    'company': company,
                    'duration': duration,
                    'achievements': achievements[:2]  # Limit to top 2 achievements
                })
        
        return experience_entries
    
    def calculate_total_experience(self, experience_entries: list) -> float:
        """Calculate total years of experience from entries."""
        from datetime import datetime
        
        total_years = 0.0
        current_year = datetime.now().year
        
        for entry in experience_entries:
            duration = entry.get('duration', '')
            if not duration:
                # Try to extract from experience text or achievements
                exp_text = str(entry.get('achievements', [])) + ' ' + str(entry.get('job_title', ''))
                # Look for year patterns in the text
                year_matches = re.findall(r'(\d{4})', exp_text)
                if year_matches:
                    try:
                        years = [int(y) for y in year_matches if 1980 <= int(y) <= current_year]
                        if len(years) >= 2:
                            duration = f"{min(years)}-{max(years)}"
                        elif len(years) == 1:
                            duration = f"{years[0]}-present"
                    except:
                        pass
            
            if not duration:
                continue
            
            # Extract years from duration strings
            year_patterns = [
                r'(\d{4})\s*[-–]\s*(\d{4})',  # 2020-2023
                r'(\d{4})\s*[-–]\s*(?:present|current|now)',  # 2020-present
                r'(\d+(?:\.\d+)?)\s*years?',  # 2.5 years
                r'(\d+)\s*\+?\s*years?'  # 3+ years
            ]
            
            found_match = False
            for pattern in year_patterns:
                matches = re.findall(pattern, duration.lower())
                if matches:
                    found_match = True
                    if len(matches[0]) == 2:  # Date range
                        try:
                            start_year = int(matches[0][0])
                            end_str = matches[0][1].lower()
                            if end_str in ['present', 'current', 'now']:
                                end_year = current_year
                            else:
                                end_year = int(end_str) if end_str.isdigit() else current_year
                            years = end_year - start_year
                            # Add partial year if current year
                            if end_str in ['present', 'current', 'now']:
                                years += 0.5  # Assume mid-year
                            total_years += max(0, years)
                        except (ValueError, IndexError):
                            pass
                    else:  # Single year value
                        try:
                            years = float(matches[0])
                            total_years += years
                        except ValueError:
                            pass
                    break
        
        return total_years
    
    def identify_recurring_themes(self, experience_entries: list) -> list:
        """Identify recurring themes across experience entries."""
        theme_keywords = {
            'machine_learning': ['ml', 'machine learning', 'ai', 'artificial intelligence', 'model', 'prediction', 'algorithm'],
            'data_science': ['data science', 'analytics', 'statistics', 'data analysis', 'insights'],
            'software_development': ['software', 'development', 'programming', 'coding', 'engineering'],
            'cloud_computing': ['aws', 'azure', 'gcp', 'cloud', 'infrastructure'],
            'data_engineering': ['etl', 'pipeline', 'data warehouse', 'big data', 'spark', 'hadoop'],
            'devops': ['devops', 'deployment', 'ci/cd', 'docker', 'kubernetes', 'automation'],
            'product_management': ['product', 'management', 'strategy', 'roadmap', 'stakeholder'],
            'leadership': ['lead', 'managed', 'team', 'supervised', 'mentored', 'directed']
        }
        
        theme_counts = {theme: 0 for theme in theme_keywords.keys()}
        
        for entry in experience_entries:
            text = f"{entry.get('job_title', '')} {entry.get('company', '')} {' '.join(entry.get('achievements', []))}".lower()
            
            for theme, keywords in theme_keywords.items():
                for keyword in keywords:
                    if keyword in text:
                        theme_counts[theme] += 1
                        break
        
        # Return themes that appear in multiple entries
        recurring_themes = [theme for theme, count in theme_counts.items() if count >= 2]
        return recurring_themes
    
    def generate_professional_summary(self, resume_sections: dict) -> str:
        """Generate a comprehensive professional summary from experience and skills using LLM if available."""
        
        # Check if there's already a summary
        existing_summary = resume_sections.get('summary', '').strip()
        experience_text = resume_sections.get('experience', '').strip()
        skills_text = resume_sections.get('skills', '').strip()
        education_text = resume_sections.get('education', '').strip()
        raw_text = resume_sections.get('raw_text', '').strip()
        
        # Try to use LLM for enhanced summary generation (only if enabled)
        use_llm_enabled = st.session_state.get('use_llm', False)
        if use_llm_enabled and LLM_AVAILABLE and is_llm_available():
            try:
                llm_client = get_llm_client()
                llm_summary = llm_client.generate_professional_summary(
                    experience=experience_text or raw_text[:800],
                    skills=skills_text or '',
                    education=education_text or '',
                    raw_text=raw_text[:1000] if raw_text else None
                )
                if llm_summary and len(llm_summary.strip()) > 50:
                    return llm_summary
            except Exception as e:
                print(f"DEBUG: LLM summary generation failed: {e}")
                # Fall through to rule-based generation
        
        # If no experience in sections, try to extract from raw text (common for PDFs)
        if not experience_text and raw_text:
            # Try to find experience section in raw text
            # Look for experience-related patterns
            experience_patterns = [
                r'experience[^.]*?(?:senior|data engineer|engineer|developer|analyst|manager)[^.]*?(?:\d{4}[-–]\d{4}|\d{4}[-–]present|present)',
                r'(?:senior|data engineer|engineer|developer|analyst|manager)[^.]*?(?:\d{4}[-–]\d{4}|\d{4}[-–]present|present)[^.]*?implemented|built|developed|created',
                r'(?:senior|data engineer|engineer|developer|analyst|manager).*?(?:flowmetrics|retailops|company|corp|inc)',
            ]
            
            for pattern in experience_patterns:
                matches = re.finditer(pattern, raw_text, re.IGNORECASE | re.DOTALL)
                for match in matches:
                    exp_section = match.group(0)
                    if len(exp_section) > 50:  # Substantial match
                        experience_text = exp_section
                        break
                if experience_text:
                    break
            
            # If still no experience, try to extract from lines containing job titles
        if not experience_text:
                lines = raw_text.split('\n')
                experience_lines = []
                in_experience = False
                
                for i, line in enumerate(lines):
                    line_lower = line.lower()
                    # Detect experience section
                    if any(keyword in line_lower for keyword in ['experience', 'work experience', 'employment', 'professional experience']):
                        in_experience = True
                        continue
                    # Stop at other sections
                    if in_experience and any(keyword in line_lower for keyword in ['education', 'skills', 'projects', 'certifications']):
                        break
                    # Collect experience lines
                    if in_experience and line.strip():
                        # Look for job titles or companies
                        if any(word in line_lower for word in ['engineer', 'developer', 'analyst', 'manager', 'senior', 'data', 'software']) or \
                           any(word in line_lower for word in ['implemented', 'built', 'developed', 'created', 'migrated', 'decreased', 'increased']):
                            experience_lines.append(line.strip())
                
                if experience_lines:
                    experience_text = ' '.join(experience_lines[:10])  # First 10 experience lines
        
        # If still no experience, try to extract from the entire raw text
        if not experience_text and raw_text:
            # Extract sentences with experience-related content
            # First, try to find the Experience section explicitly
            exp_section_match = re.search(r'experience[^.]*?(?=education|skills|projects|certifications|$)', raw_text, re.IGNORECASE | re.DOTALL)
            if exp_section_match:
                exp_section = exp_section_match.group(0)
                # Extract meaningful lines from experience section
                lines = exp_section.split('\n')
                exp_lines = []
                for line in lines:
                    line = line.strip()
                    if len(line) > 20 and any(word in line.lower() for word in ['engineer', 'developer', 'analyst', 'manager', 'senior', 'implemented', 'built', 'developed', 'created', 'migrated']):
                        exp_lines.append(line)
                if exp_lines:
                    experience_text = ' '.join(exp_lines[:15])  # First 15 experience lines
            
            # If still no experience, extract sentences with experience-related content
            if not experience_text:
                sentences = re.split(r'[.!?]\s+', raw_text)
                experience_sentences = []
                for sentence in sentences:
                    sentence_lower = sentence.lower()
                    # Look for experience indicators
                    if (any(word in sentence_lower for word in ['engineer', 'developer', 'analyst', 'manager', 'senior', 'data engineer']) and
                        any(word in sentence_lower for word in ['implemented', 'built', 'developed', 'created', 'migrated', 'decreased', 'increased', 'years', 'experience'])):
                        experience_sentences.append(sentence.strip())
                        if len(experience_sentences) >= 5:
                            break
                
                if experience_sentences:
                    experience_text = ' '.join(experience_sentences)
        
        # If no experience found at all, return existing summary or try to generate from available data
        if not experience_text:
            if existing_summary and len(existing_summary) > 50:
                return existing_summary
            # Try to generate a basic summary from skills and education
            if skills_text or education_text:
                summary_parts = []
                if skills_text:
                    # Extract key skills
                    skills_list = skills_text.split(',')[:5]
                    skills_display = ", ".join([s.strip() for s in skills_list if s.strip()])
                    summary_parts.append(f"Professional with expertise in {skills_display}")
                if education_text:
                    # Extract degree
                    degree_match = re.search(r'(b\.?s\.?|m\.?s\.?|bachelor|master|phd|ph\.?d\.?|mba)', education_text.lower())
                    if degree_match:
                        summary_parts.append(f"holding a {degree_match.group(0).upper()} degree")
                if summary_parts:
                    return ". ".join(summary_parts) + "."
            
            # Last resort: generate from raw text if available
            if raw_text and len(raw_text.strip()) > 100:
                summary_parts = []
                
                # Extract role from raw text
                role_patterns = [
                    r'(?:senior\s+)?(?:data\s+)?(?:engineer|scientist|analyst|developer|manager)',
                    r'(?:software|data|machine learning|ml)\s+(?:engineer|scientist|analyst)',
                    r'(?:senior|lead|principal)\s+(?:engineer|developer|analyst)'
                ]
                for pattern in role_patterns:
                    role_match = re.search(pattern, raw_text.lower())
                    if role_match:
                        summary_parts.append(role_match.group(0).title())
                        break
                
                # Extract years of experience
                years_match = re.search(r'(\d+)\s*years?\s*(?:of\s*)?(?:experience|exp)', raw_text.lower())
                if years_match:
                    summary_parts.append(f"with {years_match.group(1)} years of experience")
                
                # Extract key technologies
                tech_keywords = ['python', 'sql', 'spark', 'kafka', 'airflow', 'snowflake', 'bigquery', 'aws', 'docker', 'kubernetes', 'tableau', 'pandas', 'scikit-learn']
                found_techs = [tech for tech in tech_keywords if tech in raw_text.lower()]
                if found_techs:
                    summary_parts.append(f"proficient in {', '.join(found_techs[:5])}")
                
                if summary_parts:
                    return ". ".join(summary_parts) + "."
            
            return "Professional with relevant experience and technical skills."
        
        # Extract experience elements
        experience_entries = self.extract_experience_elements(experience_text)
        
        # If extraction failed but we have experience text, try a more lenient extraction
        if not experience_entries and experience_text and len(experience_text) > 50:
            # Try to extract basic info from the text directly
            # Look for job titles and companies
            job_patterns = [
                r'(?:senior\s+)?(?:data\s+)?engineer[^.]*?(?:flowmetrics|retailops|company|corp|inc)',
                r'(?:senior\s+)?(?:data\s+)?engineer[^.]*?\d{4}[-–]\d{4}',
                r'(?:data\s+)?engineer[^.]*?implemented|built|developed|created|migrated',
            ]
            
            for pattern in job_patterns:
                matches = re.finditer(pattern, experience_text, re.IGNORECASE)
                for match in matches:
                    exp_text = match.group(0)
                    if len(exp_text) > 30:
                        experience_entries.append({
                            'job_title': 'Data Engineer',  # Default title
                            'company': '',
                            'duration': '',
                            'achievements': [exp_text]
                        })
                        break
                if experience_entries:
                    break
        
        if not experience_entries:
            if existing_summary and len(existing_summary) > 50:
                return existing_summary
            # If we have raw text with experience keywords, create a basic summary
            if raw_text:
                # Extract key information from raw text
                summary_parts = []
                # Look for role
                role_match = re.search(r'(?:senior\s+)?(?:data\s+)?engineer', raw_text.lower())
                if role_match:
                    summary_parts.append("Data Engineer")
                # Look for years
                years_match = re.search(r'(\d+)\s*years?', raw_text.lower())
                if years_match:
                    summary_parts.append(f"with {years_match.group(1)} years of experience")
                # Look for key technologies
                tech_keywords = ['spark', 'kafka', 'airflow', 'snowflake', 'bigquery', 'python', 'sql']
                found_techs = [tech for tech in tech_keywords if tech in raw_text.lower()]
                if found_techs:
                    summary_parts.append(f"proficient in {', '.join(found_techs[:5])}")
                # Look for companies
                company_match = re.search(r'(flowmetrics|retailops)', raw_text, re.IGNORECASE)
                if company_match:
                    summary_parts.append(f"at {company_match.group(0)}")
                
                if summary_parts:
                    return ". ".join(summary_parts) + "."
            
            # Last resort: generate a basic summary from raw text
            if raw_text and len(raw_text.strip()) > 50:
                # Extract key information from raw text for a basic summary
                summary_parts = []
                
                # Look for any role/title
                role_patterns = [
                    r'(?:senior\s+)?(?:data\s+)?(?:engineer|scientist|analyst|developer|manager)',
                    r'(?:software|data|machine learning|ml)\s+(?:engineer|scientist|analyst)',
                    r'(?:senior|lead|principal)\s+(?:engineer|developer|analyst)'
                ]
                for pattern in role_patterns:
                    role_match = re.search(pattern, raw_text.lower())
                    if role_match:
                        summary_parts.append(role_match.group(0).title())
                        break
                
                # Look for years of experience
                years_match = re.search(r'(\d+)\s*years?\s*(?:of\s*)?(?:experience|exp)', raw_text.lower())
                if years_match:
                    summary_parts.append(f"with {years_match.group(1)} years of experience")
                
                # Look for key technologies/skills
                tech_keywords = ['python', 'sql', 'spark', 'kafka', 'airflow', 'snowflake', 'bigquery', 'aws', 'docker', 'kubernetes']
                found_techs = [tech for tech in tech_keywords if tech in raw_text.lower()]
                if found_techs:
                    summary_parts.append(f"proficient in {', '.join(found_techs[:5])}")
                
                if summary_parts:
                    return ". ".join(summary_parts) + "."
            
            return "Professional with relevant experience and technical skills."
        
        # Calculate total experience
        total_years = self.calculate_total_experience(experience_entries)
        
        # Identify recurring themes
        recurring_themes = self.identify_recurring_themes(experience_entries)
        
        # Extract key skills
        key_skills = []
        if skills_text:
            # Extract skills (handle comma, semicolon, or newline separated)
            skill_separators = [',', ';', '\n', '|']
            skill_words = [skills_text]
            for sep in skill_separators:
                if sep in skills_text:
                    skill_words = skills_text.split(sep)
                    break
            
            key_skills = [
                skill.strip().title() 
                for skill in skill_words 
                if len(skill.strip()) > 2 and not skill.strip().lower() in ['and', 'or', 'the']
            ][:15]  # Get up to 15 skills
        
        # Build comprehensive summary
        summary_paragraphs = []
        
        # Paragraph 1: Introduction with role, experience, and companies
        intro_parts = []
        if experience_entries:
            primary_role = experience_entries[0].get('job_title', 'Professional')
            if not primary_role or primary_role == 'Professional':
                # Try to extract role from experience text or raw text
                role_keywords = ['senior data engineer', 'data engineer', 'software engineer', 'engineer', 'developer', 'analyst', 'manager', 'scientist', 'specialist', 'consultant']
                search_text = (raw_text or experience_text or '').lower()
                for keyword in role_keywords:
                    if keyword in search_text:
                        primary_role = keyword.title()
                        # If it's "Senior Data Engineer", keep the full title
                        if 'senior' in keyword:
                            primary_role = keyword.replace('senior', 'Senior').replace('data', 'Data').replace('engineer', 'Engineer')
                        break
            
            # Determine experience level - also try to extract from raw text if calculation failed
            if total_years < 0.5:
                # Try to extract years from raw text or experience text
                year_pattern = r'(\d+)\s*years?\s*(?:of\s*)?(?:experience|exp)'
                year_matches = re.findall(year_pattern, (raw_text or experience_text or '').lower())
                if year_matches:
                    try:
                        total_years = float(year_matches[0])
                    except:
                        pass
                # If still no years, try to calculate from dates in text
                if total_years < 0.5:
                    date_pattern = r'(\d{4})\s*[-–]\s*(\d{4}|present|current)'
                    date_matches = re.findall(date_pattern, (raw_text or experience_text or ''), re.IGNORECASE)
                    if date_matches:
                        from datetime import datetime
                        current_year = datetime.now().year
                        for start, end in date_matches:
                            try:
                                start_year = int(start)
                                end_year = current_year if end.lower() in ['present', 'current'] else int(end)
                                years = end_year - start_year
                                if end.lower() in ['present', 'current']:
                                    years += 0.5
                                total_years += max(0, years)
                            except:
                                pass
            
            # Format experience level
            if total_years >= 10:
                exp_level = f"{total_years:.0f}+ years"
            elif total_years >= 5:
                exp_level = f"{total_years:.0f}+ years"
            elif total_years >= 1:
                exp_level = f"{total_years:.1f} years"
            else:
                exp_level = "entry-level"
            
            intro_parts.append(f"{primary_role} with {exp_level} of professional experience")
        
        # Add company context
        companies = [entry.get('company', '') for entry in experience_entries if entry.get('company')]
        if companies:
            unique_companies = list(dict.fromkeys(companies))
            if len(unique_companies) == 1:
                intro_parts.append(f"at {unique_companies[0]}")
            elif len(unique_companies) <= 3:
                company_context = ", ".join(unique_companies[:-1]) + f", and {unique_companies[-1]}"
                intro_parts.append(f"across {company_context}")
            else:
                intro_parts.append(f"across {unique_companies[0]} and {len(unique_companies)-1} other organizations")
        
        if intro_parts:
            summary_paragraphs.append(". ".join(intro_parts) + ".")
        
        # Paragraph 2: Core expertise and themes
        if recurring_themes or key_skills:
            expertise_parts = []
            
            if recurring_themes:
                theme_descriptions = {
                    'machine_learning': 'machine learning and artificial intelligence',
                    'data_science': 'data science and analytics',
                    'software_development': 'software development',
                    'cloud_computing': 'cloud computing and infrastructure',
                    'data_engineering': 'data engineering and ETL pipelines',
                    'devops': 'DevOps, CI/CD, and automation',
                    'product_management': 'product management and strategy',
                    'leadership': 'team leadership and management'
                }
                theme_texts = [theme_descriptions.get(theme, theme.replace('_', ' ')) for theme in recurring_themes[:4]]
                if theme_texts:
                    expertise_parts.append(f"Specialized expertise in {', '.join(theme_texts)}")
            
            if key_skills and len(key_skills) > 0:
                # Group skills into categories if possible
                tech_skills = [s for s in key_skills[:10] if any(tech in s.lower() for tech in ['python', 'java', 'sql', 'javascript', 'react', 'aws', 'docker', 'kubernetes', 'spark', 'kafka', 'airflow', 'snowflake', 'bigquery', 'terraform', 'gcp', 'azure'])]
                if tech_skills:
                    skills_display = ", ".join(tech_skills[:10])  # Show more skills
                    expertise_parts.append(f"Proficient in technologies including {skills_display}")
                elif key_skills:
                    # If no tech skills matched, show top skills anyway
                    skills_display = ", ".join(key_skills[:8])
                    expertise_parts.append(f"Skilled in {skills_display}")
            
            if expertise_parts:
                summary_paragraphs.append(" ".join(expertise_parts) + ".")
        
        # Paragraph 3: Key achievements and responsibilities
        achievement_sentences = []
        # Define action verbs at the start so they're always available
        action_verbs = ['developed', 'implemented', 'built', 'created', 'designed', 'led', 'managed', 'improved', 'increased', 'reduced', 'optimized', 'migrated', 'deployed', 'established', 'architected', 'scaled', 'automated']
        
        for entry in experience_entries[:3]:  # Top 3 roles
            achievements = entry.get('achievements', [])
            job_title = entry.get('job_title', '')
            company = entry.get('company', '')
            duration = entry.get('duration', '')
            
            if achievements:
                for achievement in achievements[:3]:  # Top 3 achievements per role for more detail
                    if len(achievement) > 20:  # Only substantial achievements
                        # Clean up the achievement text
                        achievement_clean = achievement.strip()
                        if not achievement_clean.endswith('.'):
                            achievement_clean += '.'
                        # Add context if we have company info
                        if company and company not in achievement_clean:
                            achievement_sentences.append(f"At {company}, {achievement_clean.lower()}")
                        else:
                            achievement_sentences.append(achievement_clean)
        
        # If no structured achievements, extract from experience text
        if not achievement_sentences and experience_text:
            # Extract sentences with action verbs
            sentences = re.split(r'[.!?]\s+', experience_text)
            for sentence in sentences:
                if any(verb in sentence.lower() for verb in action_verbs) and len(sentence) > 30:
                    achievement_sentences.append(sentence.strip() + '.')
                    if len(achievement_sentences) >= 6:  # Get more achievements
                        break
        
        # Also extract bullet points that might be achievements
        if len(achievement_sentences) < 6 and experience_text:
            bullet_pattern = r'[•\-\*o]\s*([^•\-\*\n]+)'
            bullet_matches = re.findall(bullet_pattern, experience_text)
            for bullet in bullet_matches:
                bullet = bullet.strip()
                if len(bullet) > 30 and any(verb in bullet.lower() for verb in action_verbs):
                    # Check if this achievement is already captured
                    is_duplicate = any(bullet.lower() in existing.lower() or existing.lower() in bullet.lower() 
                                      for existing in achievement_sentences)
                    if not is_duplicate:
                        achievement_sentences.append(bullet + ('' if bullet.endswith('.') else '.'))
                    if len(achievement_sentences) >= 8:  # Get more achievements
                        break
        
        # Also check raw text for additional achievements if we need more
        if len(achievement_sentences) < 4 and raw_text:
            raw_bullets = re.findall(r'[•\-\*o]\s*([^•\-\*\n]+)', raw_text)
            for bullet in raw_bullets:
                bullet = bullet.strip()
                if len(bullet) > 30 and any(verb in bullet.lower() for verb in action_verbs):
                    is_duplicate = any(bullet.lower() in existing.lower() or existing.lower() in bullet.lower() 
                                      for existing in achievement_sentences)
                    if not is_duplicate:
                        achievement_sentences.append(bullet + ('' if bullet.endswith('.') else '.'))
                    if len(achievement_sentences) >= 8:
                        break
        
        if achievement_sentences:
            achievements_text = " ".join(achievement_sentences[:8])  # Up to 8 key achievements for comprehensive detail
            summary_paragraphs.append(f"Key accomplishments include: {achievements_text}")
        
        # Paragraph 4: Education and additional context
        if education_text:
            # Extract degree information
            degree_keywords = ['bachelor', 'master', 'phd', 'ph.d', 'mba', 'bs', 'ms', 'ba', 'ma']
            education_lines = education_text.split('\n')[:3]  # First 3 education entries
            education_summary = []
            for line in education_lines:
                if any(keyword in line.lower() for keyword in degree_keywords):
                    education_summary.append(line.strip())
            
            if education_summary:
                edu_text = ". ".join(education_summary[:3])
                summary_paragraphs.append(f"Educational background: {edu_text}.")
        
        # Paragraph 5: Additional context from skills and technologies
        if skills_text and len(summary_paragraphs) < 4:
            # Extract key technologies that haven't been mentioned
            mentioned_techs = set()
            for para in summary_paragraphs:
                para_lower = para.lower()
                for tech in ['python', 'sql', 'spark', 'kafka', 'airflow', 'aws', 'gcp', 'azure', 'docker', 'kubernetes']:
                    if tech in para_lower:
                        mentioned_techs.add(tech)
            
            # Get skills that haven't been mentioned
            skills_list = [s.strip() for s in skills_text.replace(',', '|').replace(';', '|').split('|') if s.strip()]
            unmentioned_skills = [s for s in skills_list if s.lower() not in mentioned_techs and len(s) > 2][:5]
            
            if unmentioned_skills:
                additional_skills = ", ".join(unmentioned_skills)
                summary_paragraphs.append(f"Additional technical competencies include {additional_skills}.")
        
        # Combine all paragraphs
        if summary_paragraphs:
            summary = " ".join(summary_paragraphs)
        else:
            summary = f"Professional with {total_years:.1f} years of experience in the field."
        
        # If we have an existing summary, use it as a base and enhance
        if existing_summary and len(existing_summary) > 50:
            # Combine existing summary with generated details
            summary = f"{existing_summary} {summary}"
        
        # Ensure summary is comprehensive but not too long (target: 200-400 words for detailed summaries)
        words = summary.split()
        if len(words) > 450:
            # Trim to ~400 words but keep complete sentences
            sentences = re.split(r'([.!?]\s+)', summary)
            trimmed_summary = ""
            word_count = 0
            for i in range(0, len(sentences), 2):
                sentence = sentences[i] + (sentences[i+1] if i+1 < len(sentences) else '')
                sentence_words = len(sentence.split())
                if word_count + sentence_words > 400:
                    break
                trimmed_summary += sentence
                word_count += sentence_words
            summary = trimmed_summary.strip()
        elif len(words) < 80:
            # If too short, add more detail from experience and skills
            enhancement_parts = []
            if experience_text and len(experience_text) > 100:
                # Extract more detailed experience
                exp_sentences = re.split(r'[.!?]\s+', experience_text)
                detailed_exp = [s.strip() for s in exp_sentences if len(s.strip()) > 40][:3]
                if detailed_exp:
                    enhancement_parts.append(" ".join(detailed_exp))
            
            if skills_text and len(skills_text) > 20:
                skills_list = [s.strip() for s in skills_text.replace(',', '|').replace(';', '|').split('|') if s.strip()][:8]
                if skills_list:
                    enhancement_parts.append(f"Technical skills include {', '.join(skills_list)}.")
            
            if enhancement_parts:
                summary += " " + " ".join(enhancement_parts)
        
        return summary.strip()
    
    def preprocess_resume(self, parsed_resume: dict) -> str:
        """Preprocess and clean resume text for better embedding quality."""
        # Define irrelevant sections to remove
        irrelevant_sections = ["contact", "links", "certifications", "references", "personal"]
        
        # Combine important sections with fallback logic
        summary = parsed_resume.get("summary", parsed_resume.get("profile", ""))
        skills = parsed_resume.get("skills", "")
        experience = parsed_resume.get("experience", "")
        education = parsed_resume.get("education", "")
        
        # Merge the most informative content
        merged_content = f"{summary}\n{skills}\n{experience}\n{education}"
        
        # Remove low-signal sections
        for section in irrelevant_sections:
            merged_content = merged_content.replace(section, "")
        
        # Normalize whitespace and clean text
        merged_content = " ".join(merged_content.split())
        
        # Remove common noise patterns
        noise_patterns = [
            r'https?://\S+',  # URLs
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',  # Phone numbers
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',  # Email addresses
            r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\b',  # Dates
            r'\b\d{4}\b',  # Years
            r'[^\w\s.,!?-]',  # Special characters except basic punctuation
        ]
        
        for pattern in noise_patterns:
            merged_content = re.sub(pattern, ' ', merged_content, flags=re.IGNORECASE)
        
        # Normalize whitespace again after cleaning
        merged_content = " ".join(merged_content.split())
        
        # Truncate overly long text for embedding stability
        return merged_content[:4000]  # roughly ~700–800 tokens
    
    def get_weighted_resume_embedding(self, parsed_resume: dict) -> np.ndarray:
        """Generate weighted resume embedding emphasizing skills and experience."""
        # Use default weights from improved algorithm
        return self.get_weighted_resume_embedding_with_custom_weights(
            parsed_resume,
            education_weight=0.1,
            skills_weight=0.4,
            experience_weight=0.4,
            summary_weight=0.2
        )
    
    def get_weighted_resume_embedding_with_custom_weights(self, parsed_resume: dict, 
                                                        education_weight: float = 0.1,
                                                        skills_weight: float = 0.4,
                                                        experience_weight: float = 0.4,
                                                        summary_weight: float = 0.2) -> np.ndarray:
        """Generate weighted resume embedding with custom weights."""
        summary_text = parsed_resume.get("summary", parsed_resume.get("profile", ""))
        skills_text = parsed_resume.get("skills", "")
        experience_text = parsed_resume.get("experience", "")
        education_text = parsed_resume.get("education", "")
        
        # Encode each section (skip if empty)
        vectors = {}
        weights = {}
        
        if summary_text and len(summary_text.strip()) > 10:
            vectors["summary"] = self._get_embedding(summary_text, 'summary')
            weights["summary"] = summary_weight
            print(f"DEBUG: Added summary embedding (weight: {summary_weight})")
        
        if skills_text and len(skills_text.strip()) > 10:
            vectors["skills"] = self._get_embedding(skills_text, 'skills')
            weights["skills"] = skills_weight
            print(f"DEBUG: Added skills embedding (weight: {skills_weight})")
        
        if experience_text and len(experience_text.strip()) > 20:
            vectors["experience"] = self._get_embedding(experience_text, 'experience')
            weights["experience"] = experience_weight
            print(f"DEBUG: Added experience embedding (weight: {experience_weight})")
        
        if education_text and len(education_text.strip()) > 10:
            vectors["education"] = self._get_embedding(education_text, 'education')
            weights["education"] = education_weight
            print(f"DEBUG: Added education embedding (weight: {education_weight})")
        
        if not vectors:
            print("⚠️ WARNING: No valid sections found for embedding")
            print(f"DEBUG: Resume sections - Summary: '{summary_text[:50] if summary_text else 'EMPTY'}...', Skills: '{skills_text[:50] if skills_text else 'EMPTY'}...', Experience: '{experience_text[:50] if experience_text else 'EMPTY'}...', Education: '{education_text[:50] if education_text else 'EMPTY'}...'")
            print(f"DEBUG: Section lengths - Summary: {len(summary_text) if summary_text else 0}, Skills: {len(skills_text) if skills_text else 0}, Experience: {len(experience_text) if experience_text else 0}, Education: {len(education_text) if education_text else 0}")
            
            # Try to use the raw resume text as fallback
            raw_text = f"{summary_text} {skills_text} {experience_text} {education_text}".strip()
            if len(raw_text) > 20:
                print(f"DEBUG: Using combined sections as fallback for embedding ({len(raw_text)} chars)")
                fallback_vector = self._get_embedding(raw_text, 'fallback')
                return fallback_vector
            else:
                print("ERROR: Resume has insufficient content for meaningful embedding")
                # Last resort: try to get raw text from the parsed_resume dict
                if 'raw_text' in parsed_resume and len(parsed_resume['raw_text'].strip()) > 20:
                    print(f"DEBUG: Using raw_text from parsed_resume as last resort ({len(parsed_resume['raw_text'])} chars)")
                    last_resort_vector = self._get_embedding(parsed_resume['raw_text'], 'last_resort')
                    return last_resort_vector
                else:
                    print("CRITICAL: No usable content found - trying emergency text extraction")
                    # Emergency fallback: try to extract any meaningful text from the file
                    try:
                        # This is a last resort - try to get any text from the file
                        emergency_text = self._emergency_text_extraction(parsed_resume)
                        if emergency_text and len(emergency_text.strip()) > 20:
                            print(f"DEBUG: Emergency text extraction found {len(emergency_text)} characters")
                            emergency_vector = self._get_embedding(emergency_text, 'emergency')
                            return emergency_vector
                        else:
                            print("CRITICAL: Emergency extraction also failed - trying PDF-specific fallback")
                            # PDF-specific fallback: try to extract from raw text with better parsing
                            pdf_fallback_text = self._pdf_specific_fallback(parsed_resume)
                            if pdf_fallback_text and len(pdf_fallback_text.strip()) > 20:
                                print(f"DEBUG: PDF-specific fallback found {len(pdf_fallback_text)} characters")
                                pdf_vector = self._get_embedding(pdf_fallback_text, 'pdf_fallback')
                                return pdf_vector
                            else:
                                print("CRITICAL: All fallback methods failed - using minimal text fallback")
                                # Absolute last resort: use ANY text available, even if short
                                minimal_text = parsed_resume.get('raw_text', '') or summary_text or skills_text or experience_text or education_text
                                if minimal_text and len(minimal_text.strip()) > 5:
                                    print(f"DEBUG: Using minimal text fallback ({len(minimal_text)} chars)")
                                    minimal_vector = self._get_embedding(minimal_text, 'minimal_fallback')
                                    return minimal_vector
                                else:
                                    print("CRITICAL: No text available at all - returning zero vector (this will cause 0 similarity)")
                                return np.zeros(384)  # Default embedding size
                    except Exception as e:
                        print(f"CRITICAL: Emergency extraction failed: {e}")
                        # Try one more time with raw_text if available
                        minimal_text = parsed_resume.get('raw_text', '')
                        if minimal_text and len(minimal_text.strip()) > 5:
                            print(f"DEBUG: Exception fallback - using raw_text ({len(minimal_text)} chars)")
                            return self._get_embedding(minimal_text, 'exception_fallback')
                        return np.zeros(384)  # Default embedding size
        
        # Normalize weights
        total_weight = sum(weights.values())
        normalized_weights = {k: v / total_weight for k, v in weights.items()}
        
        # Weighted combination
        weighted_vec = np.zeros(384)  # Default embedding size
        for section, vector in vectors.items():
            weight = normalized_weights[section]
            weighted_vec += weight * vector
            print(f"DEBUG: Added {section} contribution (normalized weight: {weight:.3f})")
        
        return weighted_vec
    
    def _pdf_specific_fallback(self, parsed_resume: dict) -> str:
        """PDF-specific fallback for when normal section extraction fails."""
        try:
            # Get raw text and apply aggressive PDF parsing
            raw_text = parsed_resume.get('raw_text', '')
            if not raw_text:
                return ""
            
            print("DEBUG: Applying PDF-specific parsing to raw text")
            
            # Split the text more aggressively for PDFs
            # Look for common PDF patterns
            pdf_patterns = [
                r'([A-Z][a-z]+ [A-Z][a-z]+)',  # Names like "Daniel Park"
                r'(\d{4}–\d{4}|\d{4}-\d{4})',  # Date ranges
                r'([A-Z][a-z]+ [A-Z][a-z]+ [A-Z][a-z]+)',  # Full names
                r'(Python|SQL|Java|JavaScript|React|Angular|Docker|Kubernetes)',  # Common tech skills
                r'(Bachelor|Master|PhD|University|College)',  # Education keywords
                r'(Engineer|Analyst|Manager|Director|Developer)',  # Job titles
            ]
            
            # Extract meaningful chunks
            meaningful_chunks = []
            for pattern in pdf_patterns:
                matches = re.findall(pattern, raw_text)
                meaningful_chunks.extend(matches)
            
            # Also try to extract sentences that contain key information
            sentences = re.split(r'[.!?]', raw_text)
            for sentence in sentences:
                sentence = sentence.strip()
                if len(sentence) > 20 and any(keyword in sentence.lower() for keyword in 
                    ['python', 'sql', 'experience', 'engineer', 'data', 'software', 'development', 'analysis']):
                    meaningful_chunks.append(sentence)
            
            if meaningful_chunks:
                pdf_text = ' '.join(meaningful_chunks)
                print(f"DEBUG: PDF-specific fallback extracted {len(meaningful_chunks)} meaningful chunks")
                return pdf_text
            
            return ""
            
        except Exception as e:
            print(f"PDF-specific fallback error: {e}")
            return ""
    
    def _emergency_text_extraction(self, parsed_resume: dict) -> str:
        """Emergency text extraction when all normal methods fail."""
        try:
            # Try to get any text from any possible key in the parsed_resume
            possible_keys = ['raw_text', 'text', 'content', 'full_text', 'document_text']
            
            for key in possible_keys:
                if key in parsed_resume and parsed_resume[key]:
                    text = str(parsed_resume[key]).strip()
                    if len(text) > 20:
                        print(f"DEBUG: Emergency extraction found text in '{key}' key")
                        return text
            
            # If no direct text, try to concatenate all string values
            all_text_parts = []
            for key, value in parsed_resume.items():
                if isinstance(value, str) and len(value.strip()) > 5:
                    all_text_parts.append(value.strip())
                elif isinstance(value, list):
                    for item in value:
                        if isinstance(item, str) and len(item.strip()) > 5:
                            all_text_parts.append(item.strip())
            
            if all_text_parts:
                emergency_text = " ".join(all_text_parts)
                print(f"DEBUG: Emergency extraction concatenated {len(all_text_parts)} text parts")
                return emergency_text
            
            return ""
            
        except Exception as e:
            print(f"Emergency extraction error: {e}")
            return ""
    
    def compute_weighted_similarity(self, parsed_resume: dict, jd_text: str) -> tuple:
        """Compute weighted similarity between resume and JD with improved preprocessing."""
        # Use default weights from improved algorithm
        return self.compute_weighted_similarity_with_custom_weights(
            parsed_resume, jd_text,
            education_weight=0.1,
            skills_weight=0.4,
            experience_weight=0.4,
            summary_weight=0.2
        )
    
    def compute_weighted_similarity_with_custom_weights(self, parsed_resume: dict, jd_text: str, 
                                                      education_weight: float = 0.1, 
                                                      skills_weight: float = 0.4, 
                                                      experience_weight: float = 0.4, 
                                                      summary_weight: float = 0.2) -> tuple:
        """Compute weighted similarity between resume and JD with custom weights."""
        # Preprocess both inputs
        clean_resume = self.preprocess_resume(parsed_resume)
        clean_jd = " ".join(jd_text.split())[:4000]  # normalize JD length
        
        print(f"DEBUG: Cleaned resume length: {len(clean_resume)} chars")
        print(f"DEBUG: Cleaned JD length: {len(clean_jd)} chars")
        
        # Generate embeddings with custom weights
        resume_vec = self.get_weighted_resume_embedding_with_custom_weights(
            parsed_resume,
            education_weight=education_weight,
            skills_weight=skills_weight,
            experience_weight=experience_weight,
            summary_weight=summary_weight
        )
        jd_vec = self._get_embedding(clean_jd, 'jd')
        
        # Compute cosine similarity
        sim = cosine_similarity(
            np.array(resume_vec).reshape(1, -1),
            np.array(jd_vec).reshape(1, -1)
        )[0][0]
        
        # Scale similarity for human-readable display with better discrimination
        scaling_mode = st.session_state.get('scaling_mode', 'Improved (Better Discrimination)')
        
        if scaling_mode == "Linear":
            # Linear scaling: raw cosine similarity
            scaled_sim = sim
        elif scaling_mode == "Original (Sigmoid)":
            # Original sigmoid scaling
            scaled_sim = 1 / (1 + np.exp(-10 * (sim - 0.3)))
        else:  # "Improved (Better Discrimination)"
            # Use improved piecewise linear scaling
            if sim <= 0.1:
                # Very low scores: compress to 0-20%
                scaled_sim = sim * 2.0
            elif sim <= 0.3:
                # Low scores: linear scaling 20-60%
                scaled_sim = 0.2 + (sim - 0.1) * 2.0
            elif sim <= 0.7:
                # Medium scores: linear scaling 60-90%
                scaled_sim = 0.6 + (sim - 0.3) * 0.75
            else:
                # High scores: preserve differences in 90-100% range
                scaled_sim = 0.9 + (sim - 0.7) * 0.33
        
        # Ensure bounds and apply minimum threshold
        scaled_sim = max(0.0, min(1.0, scaled_sim))
        
        # Apply minimum threshold to prevent zero scores (unless truly zero)
        if scaled_sim == 0.0 and sim > 0.0:
            scaled_sim = 0.01  # Minimum 1% for non-zero raw similarities
        
        display_score = round(scaled_sim * 100, 1)  # Convert to percentage
        
        print(f"DEBUG: Raw cosine similarity: {sim:.4f}")
        print(f"DEBUG: Scaled similarity: {scaled_sim:.4f}")
        print(f"DEBUG: Display score: {display_score}%")
        
        return sim, display_score
    
    def debug_similarity_breakdown(self, parsed_resume: dict, jd_text: str, jd_requirements: dict = None) -> dict:
        """Debug similarity breakdown by section for interpretability."""
        breakdown = {}
        
        # Ensure all sections are initialized
        for section in ["summary", "skills", "experience", "education"]:
            breakdown[section] = 0.0
        
        # Map sections to JD requirement keys
        section_to_jd_key = {
            "education": "education",
            "skills": "skills",
            "experience": "experience",
            "summary": None  # Summary uses full JD text
        }
        
        for section in ["summary", "skills", "experience", "education"]:
            text = parsed_resume.get(section, "")
            
            # If section is empty, try to extract from raw text using API (only if enabled)
            if not text or len(text.strip()) < 10:
                raw_text = parsed_resume.get('raw_text', '')
                use_llm_enabled = st.session_state.get('use_llm', False)
                if raw_text and use_llm_enabled and LLM_AVAILABLE and is_llm_available():
                    try:
                        llm_client = get_llm_client()
                        enhanced = llm_client.enhance_resume_parsing(raw_text, target_section=section)
                        if section in enhanced and len(enhanced[section].strip()) > 10:
                            text = enhanced[section]
                            parsed_resume[section] = text  # Update for future use
                            print(f"DEBUG: Enhanced {section} using Gemini API ({len(text)} chars)")
                    except Exception as e:
                        print(f"DEBUG: Failed to enhance {section}: {e}")
            
            # Final fallback: use raw text if still empty
            if not text or len(text.strip()) < 10:
                raw_text = parsed_resume.get('raw_text', '')
                if raw_text:
                    # Extract relevant lines from raw text with better education extraction
                    keywords = {
                        'education': [
                            'university', 'college', 'degree', 'bachelor', 'master', 'phd', 'ph.d', 'education',
                            'b.s', 'b.a', 'm.s', 'm.a', 'mba', 'ms', 'bs', 'ba', 'ma', 'graduated', 'graduate',
                            'diploma', 'certificate', 'certification', 'school', 'institute', 'academic',
                            'major', 'minor', 'gpa', 'honors', 'cum laude', 'magna cum laude'
                        ],
                        'skills': ['skill', 'proficient', 'expert', 'knowledge'],
                        'experience': ['experience', 'worked', 'engineer', 'developer', 'manager'],
                        'summary': ['summary', 'objective', 'profile', 'overview']
                    }
                    section_keywords = keywords.get(section, [])
                    lines = raw_text.split('\n')
                    
                    if section == 'education':
                        # For education, get lines with keywords and surrounding context
                        relevant_lines = []
                        for i, line in enumerate(lines):
                            line_lower = line.lower()
                            if any(kw in line_lower for kw in section_keywords):
                                relevant_lines.append(line)
                                # Add next 2 lines for context (degree details, dates, etc.)
                                for j in range(1, 3):
                                    if i + j < len(lines) and lines[i + j].strip():
                                        next_line = lines[i + j].strip()
                                        # Only add if it looks like education info (has year, degree, or is short)
                                        if (re.search(r'\d{4}', next_line) or 
                                            any(edu_kw in next_line.lower() for edu_kw in ['degree', 'university', 'college', 'major', 'minor']) or
                                            len(next_line.split()) <= 10):
                                            relevant_lines.append(next_line)
                        text = ' '.join(relevant_lines[:10])  # Up to 10 lines for education
                    else:
                        relevant_lines = [line for line in lines if any(kw in line.lower() for kw in section_keywords)]
                        text = ' '.join(relevant_lines[:3])  # First 3 relevant lines
                    
                    if text:
                        print(f"DEBUG: Using fallback extraction for {section} ({len(text)} chars)")
            
            if not text or len(text.strip()) < 10:
                breakdown[section] = 0.0
                print(f"DEBUG: {section.title()} Similarity: 0.000 (empty/insufficient content)")
                continue
            
            # Get section-specific JD text if available
            if section == "summary":
                # Summary compares against full JD text
                jd_section_text = jd_text
            elif jd_requirements and section_to_jd_key[section] in jd_requirements:
                # Extract section-specific requirements from JD
                section_reqs = jd_requirements[section_to_jd_key[section]]
                if section_reqs:
                    # Combine all requirements for this section
                    jd_section_text = " ".join([req[0] for req in section_reqs])
                    print(f"DEBUG: Using {len(section_reqs)} {section} requirements from JD")
                else:
                    # No specific requirements, fall back to full JD text
                    jd_section_text = jd_text
                    print(f"DEBUG: No {section} requirements in JD, using full JD text")
            else:
                # Fall back to full JD text if requirements not available
                jd_section_text = jd_text
                print(f"DEBUG: JD requirements not available, using full JD text for {section}")
            
            # Skip if JD section text is empty
            if not jd_section_text or len(jd_section_text.strip()) < 10:
                breakdown[section] = 0.0
                print(f"DEBUG: {section.title()} Similarity: 0.000 (JD section text empty)")
                continue
            
            # Compute similarity
            resume_vec = self._get_embedding(text, section)
            jd_vec = self._get_embedding(jd_section_text, f'jd_{section}')
            sim = cosine_similarity(
                np.array(resume_vec).reshape(1, -1),
                np.array(jd_vec).reshape(1, -1)
            )[0][0]
            
            breakdown[section] = sim
            print(f"DEBUG: {section.title()} Similarity: {sim:.3f} (resume: {len(text)} chars, JD: {len(jd_section_text)} chars)")
        
        return breakdown
    
    def compute_improved_similarity_matrix(self):
        """Compute similarity matrix using improved preprocessing and weighted embeddings."""
        if not self.model:
            self.model = load_sentence_transformer()
            if not self.model:
                st.error("SentenceTransformer not available. Please install sentence-transformers.")
                return None
        
        start_time = time.time()
        
        # Extract sections from all documents
        resume_sections = {}
        jd_requirements = {}
        
        for name, text in self.resumes.items():
            # Try enhanced parser first, fallback to original
            if ENHANCED_PARSER_AVAILABLE:
                # Find the original file path for enhanced parsing
                original_path = None
                if hasattr(self, 'original_filenames'):
                    original_path = self.original_filenames.get(name)
                
                if original_path and self.data_dir is not None:
                    sections = self.extract_sections_enhanced(str(self.data_dir / original_path))
                else:
                    sections = self.extract_sections(text)
            else:
                sections = self.extract_sections(text)
            
            resume_sections[name] = sections
        
        for name, text in self.job_descriptions.items():
            jd_requirements[name] = self.extract_jd_requirements_with_importance(text)
        
        # Compute improved similarities
        similarities = {}
        
        for jd_name, jd_reqs in jd_requirements.items():
            similarities[jd_name] = {}
            
            # Combine JD requirements into single text
            jd_text_parts = []
            for req_type, reqs in jd_reqs.items():
                if reqs:
                    jd_text_parts.extend([req[0] for req in reqs])
            jd_text = " ".join(jd_text_parts)
            
            for resume_name, resume_sections_data in resume_sections.items():
                print(f"\nDEBUG: Computing improved similarity for {resume_name} vs {jd_name}")
                
                # Debug: Check if resume sections are empty
                section_lengths = {k: len(v) if v else 0 for k, v in resume_sections_data.items()}
                print(f"DEBUG: Resume sections lengths: {section_lengths}")
                
                # Compute weighted similarity
                raw_sim, display_score = self.compute_weighted_similarity(resume_sections_data, jd_text)
                
                # Get debug breakdown with section-specific JD requirements
                breakdown = self.debug_similarity_breakdown(resume_sections_data, jd_text, jd_reqs)
                
                # Store detailed results
                similarities[jd_name][resume_name] = {
                    'raw_similarity': raw_sim,
                    'display_score': display_score,
                    'breakdown': breakdown
                }
                
                print(f"DEBUG: Final result - Raw: {raw_sim:.4f}, Display: {display_score}%")
        
        # Convert to matrix format using display scores
        similarity_matrix = np.zeros((len(self.candidate_names), len(self.jd_names)))
        
        for jd_idx, jd_name in enumerate(self.jd_names):
            for resume_idx, resume_name in enumerate(self.candidate_names):
                if jd_name in similarities and resume_name in similarities[jd_name]:
                    # Use display score (0-100) converted to 0-1 range
                    similarity_matrix[resume_idx, jd_idx] = similarities[jd_name][resume_name]['display_score'] / 100
        
        end_time = time.time()
        st.info(f"🎯 Improved similarity matching completed in {end_time - start_time:.2f} seconds")
        
        # Store similarity matrix for get_top_matches method
        self.similarity_matrix = similarity_matrix
        
        return similarity_matrix, similarities
    
    def compute_graded_similarity_matrix(self):
        """Compute similarity matrix using graded scoring system."""
        if not self.model:
            self.model = load_sentence_transformer()
            if not self.model:
                st.error("SentenceTransformer not available. Please install sentence-transformers.")
                return None
        
        start_time = time.time()
        
        # Extract sections from all documents
        resume_sections = {}
        jd_requirements = {}
        
        for name, text in self.resumes.items():
            # Try enhanced parser first, fallback to original
            if ENHANCED_PARSER_AVAILABLE and self.data_dir is not None:
                # Find the original file path for enhanced parsing (only works with directory mode)
                original_path = None
                if hasattr(self, 'original_filenames'):
                    original_path = self.original_filenames.get(name)
                
                if original_path:
                    sections = self.extract_sections_enhanced(str(self.data_dir / original_path))
                    # Store enhanced data for later use
                    if hasattr(self, '_cached_parsed_data'):
                        self._enhanced_data_cache = getattr(self, '_enhanced_data_cache', {})
                        self._enhanced_data_cache[name] = self._cached_parsed_data
                else:
                    sections = self.extract_sections(text)
            else:
                # Use regular extraction for file uploads or when enhanced parser not available
                sections = self.extract_sections(text)
            
            resume_sections[name] = sections
        
        for name, text in self.job_descriptions.items():
            jd_requirements[name] = self.extract_jd_requirements_with_importance(text)
        
        # Compute graded similarities
        similarities = {}
        
        for jd_name, jd_reqs in jd_requirements.items():
            similarities[jd_name] = {}
            
            for resume_name, resume_sections_data in resume_sections.items():
                # Compute graded similarity
                graded_scores = self.compute_graded_similarity(resume_sections_data, jd_reqs)
                
                # Store detailed results
                similarities[jd_name][resume_name] = {
                    'total': graded_scores['total_score'],
                    'required_skills': graded_scores['required_skills'],
                    'preferred_skills': graded_scores['preferred_skills'],
                    'yoe_score': graded_scores['yoe_score'],
                    'breakdown': graded_scores['breakdown']
                }
                
                print(f"DEBUG: Graded similarity for {resume_name} vs {jd_name}:")
                print(f"  Required skills: {graded_scores['required_skills']:.2f}")
                print(f"  Preferred skills: {graded_scores['preferred_skills']:.2f}")
                print(f"  YOE score: {graded_scores['yoe_score']:.2f}")
                print(f"  Total score: {graded_scores['total_score']:.2f}")
        
        # Convert to matrix format
        similarity_matrix = np.zeros((len(self.candidate_names), len(self.jd_names)))
        
        for jd_idx, jd_name in enumerate(self.jd_names):
            for resume_idx, resume_name in enumerate(self.candidate_names):
                if jd_name in similarities and resume_name in similarities[jd_name]:
                    similarity_matrix[resume_idx, jd_idx] = similarities[jd_name][resume_name]['total']
        
        end_time = time.time()
        st.info(f"🎯 Graded scoring matching completed in {end_time - start_time:.2f} seconds")
        
        # Store similarity matrix for get_top_matches method
        self.similarity_matrix = similarity_matrix
        
        return similarity_matrix, similarities
    
    def diagnose_problematic_resume(self, resume_name: str) -> dict:
        """Diagnose a specific resume that's showing zero scores."""
        print(f"\n🔍 DIAGNOSING PROBLEMATIC RESUME: {resume_name}")
        print("=" * 60)
        
        diagnosis = {
            'resume_name': resume_name,
            'file_exists': False,
            'text_extraction': False,
            'section_extraction': False,
            'sections_content': {},
            'raw_text_length': 0,
            'embedding_generation': False,
            'issues': []
        }
        
        try:
            # Find the resume file
            resume_files = list(Path("/Users/junfeibai/Desktop/5560/test/").glob("*"))
            resume_file = None
            
            for file_path in resume_files:
                if resume_name.lower().replace('-resume-', '').replace('-', ' ') in file_path.name.lower():
                    resume_file = file_path
                    break
            
            if not resume_file:
                diagnosis['issues'].append("Resume file not found")
                return diagnosis
            
            diagnosis['file_exists'] = True
            print(f"✅ Found file: {resume_file}")
            
            # Test text extraction
            try:
                raw_text = self.processor.extract_text(str(resume_file))
                diagnosis['raw_text_length'] = len(raw_text)
                diagnosis['text_extraction'] = True
                print(f"✅ Text extraction: {len(raw_text)} characters")
                print(f"   Preview: {raw_text[:200]}...")
            except Exception as e:
                diagnosis['issues'].append(f"Text extraction failed: {e}")
                print(f"❌ Text extraction failed: {e}")
                return diagnosis
            
            # Test section extraction
            try:
                sections = self.extract_sections_enhanced(str(resume_file))
                diagnosis['section_extraction'] = True
                diagnosis['sections_content'] = sections
                
                print(f"✅ Section extraction successful")
                for section, content in sections.items():
                    if section != 'raw_text':
                        print(f"   {section}: {len(content)} chars - '{content[:50]}...'")
                
            except Exception as e:
                diagnosis['issues'].append(f"Section extraction failed: {e}")
                print(f"❌ Section extraction failed: {e}")
                return diagnosis
            
            # Test embedding generation
            try:
                embedding = self.get_weighted_resume_embedding(sections)
                diagnosis['embedding_generation'] = True
                print(f"✅ Embedding generation: {len(embedding)} dimensions")
                print(f"   Non-zero elements: {np.count_nonzero(embedding)}")
                
                if np.all(embedding == 0):
                    diagnosis['issues'].append("Generated zero embedding vector")
                    print("❌ Generated zero embedding vector!")
                else:
                    print("✅ Non-zero embedding generated")
                    
            except Exception as e:
                diagnosis['issues'].append(f"Embedding generation failed: {e}")
                print(f"❌ Embedding generation failed: {e}")
            
        except Exception as e:
            diagnosis['issues'].append(f"Diagnosis failed: {e}")
            print(f"❌ Diagnosis failed: {e}")
        
        print(f"\n📊 DIAGNOSIS SUMMARY:")
        print(f"   File exists: {diagnosis['file_exists']}")
        print(f"   Text extraction: {diagnosis['text_extraction']}")
        print(f"   Section extraction: {diagnosis['section_extraction']}")
        print(f"   Embedding generation: {diagnosis['embedding_generation']}")
        print(f"   Issues: {len(diagnosis['issues'])}")
        
        if diagnosis['issues']:
            print(f"\n🚨 ISSUES FOUND:")
            for issue in diagnosis['issues']:
                print(f"   - {issue}")
        
        return diagnosis
    
    def _compute_section_similarity(self, resume_text: str, jd_requirements: List[str], section_name: str = "content", parsed_sections: dict = None) -> float:
        """Compute similarity between resume section and JD requirements with validation."""
        if not resume_text.strip() or not jd_requirements:
            return 0.0
        
        # Combine JD requirements
        jd_text = " ".join(jd_requirements)
        
        # Get embeddings with validation and fallback
        resume_embedding = self._get_embedding(resume_text, section_name, parsed_sections)
        jd_embedding = self._get_embedding(jd_text, f"{section_name}_jd")
        
        # Compute cosine similarity
        similarity = cosine_similarity([resume_embedding], [jd_embedding])[0][0]
        return float(similarity)
    
    def _compute_weighted_section_similarity(self, resume_text: str, jd_requirements: List[Tuple[str, float]]) -> float:
        """Compute weighted similarity for skills section."""
        if not resume_text.strip() or not jd_requirements:
            return 0.0
        
        resume_embedding = self._get_embedding(resume_text)
        
        # Compute weighted similarity
        total_weight = 0
        weighted_similarity = 0
        
        for req_text, importance in jd_requirements:
            req_embedding = self._get_embedding(req_text)
            similarity = cosine_similarity([resume_embedding], [req_embedding])[0][0]
            weighted_similarity += similarity * importance
            total_weight += importance
        
        return float(weighted_similarity / total_weight) if total_weight > 0 else 0.0
    
    def _validate_content_for_embedding(self, text: str, section_name: str = "content") -> tuple[str, bool]:
        """Validate content length and quality before embedding."""
        if not text or not text.strip():
            return "", False
        
        text = text.strip()
        
        # Check minimum length requirements
        min_lengths = {
            'summary': 20,
            'education': 15,
            'skills': 10,
            'experience': 30,
            'content': 10
        }
        
        min_length = min_lengths.get(section_name, 10)
        
        if len(text) < min_length:
            print(f"WARNING: {section_name} content too short ({len(text)} chars < {min_length} min)")
            return text, False
        
        # Check for meaningful content (not just repeated characters or symbols)
        unique_chars = len(set(text.lower()))
        if unique_chars < 5:
            print(f"WARNING: {section_name} content lacks diversity ({unique_chars} unique chars)")
            return text, False
        
        # Check for excessive whitespace or special characters
        alpha_ratio = sum(1 for c in text if c.isalpha()) / len(text)
        if alpha_ratio < 0.3:
            print(f"WARNING: {section_name} content has low alphabetic ratio ({alpha_ratio:.2f})")
            return text, False
        
        return text, True
    
    def _create_fallback_content(self, parsed_sections: dict, section_name: str) -> str:
        """Create fallback content when primary section is missing or insufficient."""
        fallback_strategies = {
            'summary': [
                ('experience', 200),  # First 200 chars of experience
                ('education', 100),   # First 100 chars of education
                ('skills', 50)        # First 50 chars of skills
            ],
            'experience': [
                ('summary', 150),     # Use summary if experience is missing
                ('skills', 100)       # Fallback to skills
            ],
            'education': [
                ('summary', 100),     # Use summary if education is missing
                ('experience', 50)    # Fallback to experience
            ],
            'skills': [
                ('summary', 100),     # Use summary if skills are missing
                ('experience', 50)    # Fallback to experience
            ]
        }
        
        strategies = fallback_strategies.get(section_name, [])
        fallback_parts = []
        
        for fallback_section, max_length in strategies:
            if fallback_section in parsed_sections and parsed_sections[fallback_section]:
                content = parsed_sections[fallback_section].strip()
                if content:
                    # Truncate to max_length and add to fallback
                    truncated = content[:max_length]
                    fallback_parts.append(truncated)
                    print(f"DEBUG: Using {fallback_section} as fallback for {section_name}")
        
        if fallback_parts:
            fallback_text = " ".join(fallback_parts)
            print(f"DEBUG: Created fallback {section_name} content: {len(fallback_text)} chars")
            return fallback_text
        
        return ""
    
    def _get_embedding(self, text: str, section_name: str = "content", parsed_sections: dict = None) -> np.ndarray:
        """Get embedding for text with content validation and fallback."""
        # Use stable hash (SHA256) instead of hash() for consistent caching across sessions
        cache_key = hashlib.sha256(text.encode('utf-8')).hexdigest()
        
        # Check instance cache first
        if cache_key in self.embeddings_cache:
            return self.embeddings_cache[cache_key]
        
        # Check session state cache for persistence across reruns
        session_cache_key = f"embedding_cache_{cache_key}"
        if session_cache_key in st.session_state:
            cached_embedding = st.session_state[session_cache_key]
            self.embeddings_cache[cache_key] = cached_embedding
            return cached_embedding
        
        # Validate content length and quality
        validated_text, is_valid = self._validate_content_for_embedding(text, section_name)
        
        # If content is invalid and we have parsed sections, try fallback
        if not is_valid and parsed_sections:
            fallback_text = self._create_fallback_content(parsed_sections, section_name)
            if fallback_text:
                validated_text, is_valid = self._validate_content_for_embedding(fallback_text, section_name)
                if is_valid:
                    print(f"DEBUG: Using fallback content for {section_name}")
                    text = fallback_text
        
        # If still invalid, use original text but log warning
        if not is_valid:
            print(f"WARNING: Proceeding with potentially low-quality {section_name} content")
            validated_text = text
        
        # Generate embedding
        try:
            embedding = self.model.encode(validated_text)
            # Store in both instance cache and session state for persistence
            self.embeddings_cache[cache_key] = embedding
            session_cache_key = f"embedding_cache_{cache_key}"
            st.session_state[session_cache_key] = embedding
            return embedding
        except Exception as e:
            print(f"ERROR: Failed to generate embedding for {section_name}: {e}")
            # Return zero embedding as fallback
            return np.zeros(384)  # Default embedding size for all-MiniLM-L6-v2
    
    def generate_explanation(self, resume_name: str, jd_name: str, similarities: Dict) -> str:
        """Generate natural language explanation for match using LLM if available."""
        if jd_name not in similarities or resume_name not in similarities[jd_name]:
            return "No match data available."
        
        match_data = similarities[jd_name][resume_name]
        
        # Handle different data structures from different modes
        if 'breakdown' in match_data:
            # Improved similarity mode structure
            sections = match_data['breakdown']
            total_score = match_data.get('display_score', 0) / 100  # Convert percentage to decimal
        elif 'sections' in match_data:
            # Semantic mode structure
            sections = match_data['sections']
            total_score = match_data.get('total', 0)
        else:
            return "No section data available for explanation."
        
        # Try to use LLM for enhanced explanation (only if enabled)
        use_llm_enabled = st.session_state.get('use_llm', False)
        if use_llm_enabled and LLM_AVAILABLE and is_llm_available():
            try:
                llm_client = get_llm_client()
                
                # Get resume summary
                resume_summary = ""
                if hasattr(self, 'resumes') and resume_name in self.resumes:
                    resume_text = self.resumes[resume_name]
                    if hasattr(self, 'extract_sections'):
                        resume_sections = self.extract_sections(resume_text)
                        resume_summary = resume_sections.get('summary', '')[:500]
                
                # Get JD requirements
                jd_requirements = ""
                if hasattr(self, 'job_descriptions') and jd_name in self.job_descriptions:
                    jd_requirements = self.job_descriptions[jd_name][:500]
                
                # Generate enhanced explanation
                explanation = llm_client.generate_match_explanation(
                    resume_name=resume_name,
                    jd_name=jd_name,
                    match_score=total_score,
                    section_scores=sections,
                    resume_summary=resume_summary,
                    jd_requirements=jd_requirements
                )
                
                if explanation:
                    return explanation
            except Exception as e:
                print(f"DEBUG: LLM explanation generation failed: {e}")
                # Fall through to fallback
        
        # Fallback to rule-based explanation
        strong_sections = []
        for section, score in sections.items():
            if score > 0.3:  # Threshold for "strong" match
                strong_sections.append(section)
        
        if not strong_sections:
            return f"Limited alignment between {resume_name} and {jd_name} requirements (match score: {total_score:.2%})."
        
        if len(strong_sections) == 1:
            return f"Strong match in {strong_sections[0]} with {sections[strong_sections[0]]:.2%} similarity."
        else:
            avg_score = np.mean([sections[s] for s in strong_sections])
            return f"Strong alignment across {', '.join(strong_sections)} with average similarity of {avg_score:.2%}."
    
    def get_top_matches(self, top_k: int = 3) -> pd.DataFrame:
        """Get top matches for each job description."""
        if self.similarity_matrix is None:
            return pd.DataFrame()
        
        results = []
        
        for jd_idx, jd_name in enumerate(self.jd_names):
            similarities = self.similarity_matrix[:, jd_idx]
            top_indices = np.argsort(similarities)[::-1][:top_k]
            
            for rank, resume_idx in enumerate(top_indices):
                results.append({
                    'Job Description': jd_name,
                    'Rank': rank + 1,
                    'Candidate': self.candidate_names[resume_idx],
                    'Similarity Score': similarities[resume_idx],
                    'Match Percentage': f"{similarities[resume_idx] * 100:.1f}%"
                })
        
        return pd.DataFrame(results)
    
    def get_directory_info(self) -> Dict:
        """Get information about files in the directory or uploaded files."""
        # If using file uploads (data_dir is None), return info from loaded documents
        if self.data_dir is None:
            return {
                "resume_files": [{"name": name} for name in self.candidate_names],
                "jd_files": [{"name": name} for name in self.jd_names],
                "total_files": len(self.candidate_names) + len(self.jd_names)
            }
        
        # If directory doesn't exist, return empty
        if not self.data_dir.exists():
            return {"resume_files": [], "jd_files": [], "total_files": 0}
        
        resume_files = []
        jd_files = []
        
        for file_path in self.data_dir.glob("*"):
            if file_path.is_file():
                if file_path.name.lower().startswith("jd"):
                    jd_files.append({
                        "name": file_path.name,
                        "size": file_path.stat().st_size,
                        "modified": file_path.stat().st_mtime
                    })
                else:
                    # All non-JD files are treated as candidate resumes
                    resume_files.append({
                        "name": file_path.name,
                        "size": file_path.stat().st_size,
                        "modified": file_path.stat().st_mtime
                    })
        
        return {
            "resume_files": resume_files,
            "jd_files": jd_files,
            "total_files": len(resume_files) + len(jd_files)
        }

def create_heatmap(similarity_matrix: np.ndarray, candidate_names: List[str], jd_names: List[str]):
    """Create an interactive heatmap of similarity scores."""
    
    # Ensure unique names to avoid non-unique index error
    unique_candidate_names = []
    unique_jd_names = []
    
    # Handle duplicate candidate names
    candidate_name_counts = {}
    for name in candidate_names:
        if name in candidate_name_counts:
            candidate_name_counts[name] += 1
            unique_candidate_names.append(f"{name} ({candidate_name_counts[name]})")
        else:
            candidate_name_counts[name] = 1
            unique_candidate_names.append(name)
    
    # Handle duplicate JD names
    jd_name_counts = {}
    for name in jd_names:
        if name in jd_name_counts:
            jd_name_counts[name] += 1
            unique_jd_names.append(f"{name} ({jd_name_counts[name]})")
        else:
            jd_name_counts[name] = 1
            unique_jd_names.append(name)
    
    df = pd.DataFrame(
        similarity_matrix,
        index=unique_candidate_names,
        columns=unique_jd_names
    )
    
    # Round to 3 decimal places for display
    df_display = df.round(3)
    
    fig = px.imshow(
        df_display,
        text_auto=True,
        aspect="auto",
        color_continuous_scale="RdYlBu_r",
        title="Resume-JD Similarity Heatmap"
    )
    
    fig.update_layout(
        xaxis_title="Job Descriptions",
        yaxis_title="Candidates",
        font=dict(size=12)
    )
    
    return fig, df

def main():
    """Main Streamlit application."""
    
    # Header
    st.title("🎯 Resume-JD Matching Dashboard")
    st.markdown("**Match candidates with job descriptions using TF-IDF and cosine similarity**")
    
    # Sidebar
    st.sidebar.header("Configuration")
    
    # LLM API Configuration
    st.sidebar.subheader("🤖 AI Enhancement (Optional)")
    
    # Initialize session state for API key (user must enter their own)
    if 'gemini_api_key' not in st.session_state:
        # Don't use environment variables - user must enter their own key
        st.session_state.gemini_api_key = ""
    
    # Always show API key input field - user must enter their own key
    st.sidebar.info("💡 Enter your Google Gemini API key (required for AI features)")
    api_key_input = st.sidebar.text_input(
        "Google Gemini API Key",
        value=st.session_state.gemini_api_key,
        type="password",
        help="Enter your Google Gemini API key. Get one free at https://makersuite.google.com/app/apikey"
    )
    
    # Update session state when user enters/changes API key
    if api_key_input and api_key_input != st.session_state.gemini_api_key:
        # Store in session state
        st.session_state.gemini_api_key = api_key_input
        
        # Reinitialize the LLM client with the new key
        if LLM_AVAILABLE:
            try:
                from resume_matcher.utils.llm_client import LLMClient
                # Force reinitialize by clearing the global client
                import resume_matcher.utils.llm_client as llm_module
                llm_module._llm_client = None
                llm_module._llm_client = LLMClient(api_key=api_key_input)
                
                if llm_module._llm_client.available:
                    st.sidebar.success("✅ API Key configured successfully!")
                    st.rerun()
                else:
                    st.sidebar.warning("⚠️ API key entered but validation failed. Please check your key.")
            except Exception as e:
                st.sidebar.error(f"❌ Error configuring API: {str(e)}")
    
    # Check if API key is available
    api_key_available = is_llm_available() if LLM_AVAILABLE and st.session_state.gemini_api_key else False
    
    # Store use_llm in session state for persistence
    if 'use_llm' not in st.session_state:
        st.session_state.use_llm = api_key_available if LLM_AVAILABLE else False
    
    use_llm = st.sidebar.checkbox(
        "Enable AI-Powered Enhancements",
        value=st.session_state.use_llm,
        help="Use LLM API for better explanations and summaries"
        )
    
    # Update session state when checkbox changes
    st.session_state.use_llm = use_llm

    if 'use_onet_taxonomy' not in st.session_state:
        st.session_state.use_onet_taxonomy = False
    use_onet_taxonomy = st.sidebar.checkbox(
        "Enable O*NET Smart Skill Expansion",
        value=st.session_state.use_onet_taxonomy,
        help="Expand skills with O*NET clusters for richer matching (beta scaffold)"
    )
    st.session_state.use_onet_taxonomy = use_onet_taxonomy
    if use_onet_taxonomy:
        st.sidebar.info("O*NET expansion scaffold active. Add CSVs under data/onet/ to enable full functionality.")
    
    if use_llm and LLM_AVAILABLE:
        if is_llm_available():
            st.sidebar.success("✅ AI Enhancement Active (Google Gemini)")
        else:
            st.sidebar.warning("⚠️ Google Gemini API not configured")
            st.sidebar.info("""
            **To enable AI enhancements:**
            1. Install: `pip install google-generativeai`
            2. Enter your API key in the field above
            3. Get free API key from: https://makersuite.google.com/app/apikey
            """)
    elif not LLM_AVAILABLE:
        st.sidebar.info("💡 Install google-generativeai to enable AI enhancements")
    
    # File upload section
    st.sidebar.subheader("📁 Upload Files")
    
    # Option to choose between file upload or directory
    input_method = st.sidebar.radio(
        "Input Method",
        ["📤 Upload Files", "📂 Use Directory"],
        help="Choose to upload files directly or use a directory path"
        )
    
    uploaded_resumes = None
    uploaded_jds = None
    data_dir = None
    
    if input_method == "📤 Upload Files":
        # File uploaders for resumes and job descriptions
        uploaded_resumes = st.sidebar.file_uploader(
            "Upload Resumes",
            type=['pdf', 'docx', 'doc', 'txt'],
            accept_multiple_files=True,
            help="Upload one or more resume files (PDF, DOCX, DOC, or TXT)"
        )
        
        uploaded_jds = st.sidebar.file_uploader(
            "Upload Job Descriptions",
            type=['pdf', 'docx', 'doc', 'txt'],
            accept_multiple_files=True,
            help="Upload one or more job description files (PDF, DOCX, DOC, or TXT)"
        )
        
        # Store uploaded files in session state
        if uploaded_resumes:
            st.session_state.uploaded_resumes = uploaded_resumes
        if uploaded_jds:
            st.session_state.uploaded_jds = uploaded_jds
    else:
        # Directory input (original method)
        data_dir = st.sidebar.text_input(
            "Data Directory Path",
            value="/Users/junfeibai/Desktop/5560/test",
            help="Path to directory containing resumes and job descriptions"
        )
    
    # Mode selection
    matching_mode = st.sidebar.selectbox(
        "🎯 Matching Mode",
        ["🧠 Semantic (Sentence-BERT) Mode", "🚀 Improved Similarity Mode"],
        help="Choose between semantic matching or improved similarity with preprocessing"
    )
    
    # Show mode indicator
    if "Semantic" in matching_mode:
        st.sidebar.success("🧠 **Semantic Mode Active**")
        st.sidebar.info("**Features:** Sentence-BERT embeddings, section breakdown, professional summaries")
        if not SENTENCE_TRANSFORMERS_AVAILABLE:
            st.sidebar.error("⚠️ SentenceTransformers not installed. Please run: `pip install sentence-transformers`")
    elif "Improved Similarity" in matching_mode:
        st.sidebar.success("🚀 **Improved Similarity Mode Active**")
        st.sidebar.info("**Features:** Text preprocessing, weighted embeddings, improved scaling")
        if not SENTENCE_TRANSFORMERS_AVAILABLE:
            st.sidebar.error("⚠️ SentenceTransformers not installed. Please run: `pip install sentence-transformers`")
        
        # Scaling mode controls for improved similarity
        st.sidebar.subheader("📊 Scaling Controls")
        scaling_mode = st.sidebar.selectbox(
            "Scaling Method",
            ["Improved (Better Discrimination)", "Original (Sigmoid)", "Linear"],
            help="Choose how similarity scores are scaled for display"
        )
        
        if scaling_mode == "Original (Sigmoid)":
            st.sidebar.warning("⚠️ Original sigmoid scaling may compress high scores")
        elif scaling_mode == "Improved (Better Discrimination)":
            st.sidebar.success("✅ Improved scaling preserves differences in high scores")
        else:
            st.sidebar.info("ℹ️ Linear scaling shows raw cosine similarity")
        
        # Store scaling mode in session state
        st.session_state.scaling_mode = scaling_mode
    
    # Semantic mode controls
    if "Semantic" in matching_mode:
        st.sidebar.subheader("🎛️ Weight Controls")
        
        # Initialize slider values in session state if not present (matching improved algorithm defaults)
        if "education_weight_slider" not in st.session_state:
            st.session_state.education_weight_slider = 0.1
        if "skills_weight_slider" not in st.session_state:
            st.session_state.skills_weight_slider = 0.4
        if "experience_weight_slider" not in st.session_state:
            st.session_state.experience_weight_slider = 0.4
        if "summary_weight_slider" not in st.session_state:
            st.session_state.summary_weight_slider = 0.2
        
        # Get raw slider values (using keys - values are stored in session state automatically)
        raw_education_weight = st.sidebar.slider(
            "Education Weight",
            min_value=0.0,
            max_value=1.0,
            value=st.session_state.education_weight_slider,
            step=0.05,
            key="education_weight_slider",
            help="Weight for education section matching"
        )
        raw_skills_weight = st.sidebar.slider(
            "Skills Weight",
            min_value=0.0,
            max_value=1.0,
            value=st.session_state.skills_weight_slider,
            step=0.05,
            key="skills_weight_slider",
            help="Weight for skills section matching"
        )
        raw_experience_weight = st.sidebar.slider(
            "Experience Weight",
            min_value=0.0,
            max_value=1.0,
            value=st.session_state.experience_weight_slider,
            step=0.05,
            key="experience_weight_slider",
            help="Weight for experience section matching"
        )
        raw_summary_weight = st.sidebar.slider(
            "Summary Weight",
            min_value=0.0,
            max_value=1.0,
            value=st.session_state.summary_weight_slider,
            step=0.05,
            key="summary_weight_slider",
            help="Weight for summary section matching"
        )
        
        # Normalize weights for computation (calculate immediately from current slider values)
        # These values are read directly from sliders, so they update on every rerun
        total_weight = raw_education_weight + raw_skills_weight + raw_experience_weight + raw_summary_weight
        
        if total_weight > 0:
            education_weight = raw_education_weight / total_weight
            skills_weight = raw_skills_weight / total_weight
            experience_weight = raw_experience_weight / total_weight
            summary_weight = raw_summary_weight / total_weight
        else:
            # Fallback to equal weights if total is 0
            education_weight = skills_weight = experience_weight = summary_weight = 0.25
        
        # Store normalized weights in session state for use in similarity computation
        st.session_state.normalized_weights = {
            'education': education_weight,
            'skills': skills_weight,
            'experience': experience_weight,
            'summary': summary_weight
        }
        
        # Show current normalized weights - this section updates automatically when sliders change
        # because Streamlit reruns the script on slider interaction
        st.sidebar.markdown("**Current Weights:**")
        
        # Display each weight on its own line for clarity
        # Using f-strings ensures values are recalculated on every rerun
        st.sidebar.text(f"Education: {education_weight:.2f}")
        st.sidebar.text(f"Skills: {skills_weight:.2f}")
        st.sidebar.text(f"Experience: {experience_weight:.2f}")
        st.sidebar.text(f"Summary: {summary_weight:.2f}")
        
        # Show total
        total_normalized = education_weight + skills_weight + experience_weight + summary_weight
        st.sidebar.caption(f"Total: {total_normalized:.2f}")
        
        # Update session state weights for compatibility with other parts of the code
        current_raw_weights = {
            'education': raw_education_weight,
            'skills': raw_skills_weight,
            'experience': raw_experience_weight,
            'summary': raw_summary_weight
        }
        
        # Check if weights have changed (compare with previous values)
        previous_weights = st.session_state.get('weights')
        weights_changed = previous_weights is None or previous_weights != current_raw_weights
        st.session_state.weights = current_raw_weights
        if weights_changed:
            # Clear similarity data to force recomputation
            if 'similarity_matrix' in st.session_state:
                del st.session_state.similarity_matrix
            if 'section_matrices' in st.session_state:
                del st.session_state.section_matrices
            if 'similarities' in st.session_state:
                del st.session_state.similarities
            st.session_state.weights_changed = True
            st.sidebar.success("🔄 Weights updated! Similarity will be recomputed.")
    
    # Weighted embedding mode controls
    elif "Weighted Embedding" in matching_mode:
        st.sidebar.subheader("🎯 Weighted Embedding Controls")
        st.sidebar.info("**High-level alignment focus:** Summary (50%) + Skills (30%) + Experience (20%)")
        
        summary_weight = st.sidebar.slider(
            "Summary Weight",
            min_value=0.0,
            max_value=1.0,
            value=0.5,
            step=0.05,
            help="Weight for summary section (high-level alignment terms like 'CTR', 'MLflow')"
        )
        skills_weight = st.sidebar.slider(
            "Skills Weight",
            min_value=0.0,
            max_value=1.0,
            value=0.3,
            step=0.05,
            help="Weight for skills section (technical alignment)"
        )
        experience_weight = st.sidebar.slider(
            "Experience Weight",
            min_value=0.0,
            max_value=1.0,
            value=0.2,
            step=0.05,
            help="Weight for experience section (detailed alignment)"
        )
        
        # Normalize weights
        total_weight = summary_weight + skills_weight + experience_weight
        if total_weight > 0:
            summary_weight /= total_weight
            skills_weight /= total_weight
            experience_weight /= total_weight
    
    top_k = st.sidebar.slider(
        "Top Matches to Display",
        min_value=1,
        max_value=10,
        value=3,
        help="Number of top matches to show for each job description"
    )
    
    # Auto-refresh option
    auto_refresh = st.sidebar.checkbox(
        "🔄 Auto-refresh on file changes",
        value=False,
        help="Automatically reload documents when files change"
    )
    
    # Manual refresh button
    if st.sidebar.button("🔄 Refresh Documents", help="Manually reload all documents"):
        if 'matcher' in st.session_state:
            del st.session_state.matcher
        if 'similarity_matrix' in st.session_state:
            del st.session_state.similarity_matrix
        if 'section_matrices' in st.session_state:
            del st.session_state.section_matrices
        if 'similarities' in st.session_state:
            del st.session_state.similarities
        st.rerun()
    
    # File change detection
    if auto_refresh and 'matcher' in st.session_state:
        current_matcher = st.session_state.matcher
        current_info = current_matcher.get_directory_info()
        
        # Store directory info in session state for comparison
        if 'last_directory_info' not in st.session_state:
            st.session_state.last_directory_info = current_info
        
        # Check if files have changed
        if current_info != st.session_state.last_directory_info:
            st.info("🔄 Files have changed! Reloading documents...")
            del st.session_state.matcher
            if 'similarity_matrix' in st.session_state:
                del st.session_state.similarity_matrix
            st.session_state.last_directory_info = current_info
            st.rerun()
    
    # Enhanced Parser Status
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🔍 Enhanced Parser Status")
    
    if ENHANCED_PARSER_AVAILABLE:
        st.sidebar.success("✅ Enhanced Parser Available")
        st.sidebar.write("**Features:**")
        st.sidebar.write("• spaCy NER for entity recognition")
        st.sidebar.write("• pyresparser for structured extraction")
        st.sidebar.write("• Fallback regex patterns")
        st.sidebar.write("• Confidence scoring")
    else:
        st.sidebar.error("❌ Enhanced Parser Not Available")
        st.sidebar.write("**To enable:**")
        st.sidebar.write("```bash")
        st.sidebar.write("pip install spacy pyresparser nltk")
        st.sidebar.write("python -m spacy download en_core_web_sm")
        st.sidebar.write("```")
    
    # Auto-load uploaded files if they exist and haven't been loaded yet
    if input_method == "📤 Upload Files":
        uploaded_resumes_list = st.session_state.get('uploaded_resumes', None)
        uploaded_jds_list = st.session_state.get('uploaded_jds', None)
        
        # Check if files have changed (compare file names)
        files_changed = False
        if uploaded_resumes_list:
            current_resume_names = [f.name for f in uploaded_resumes_list] if uploaded_resumes_list else []
            previous_resume_names = [f.name for f in st.session_state.get('last_uploaded_resumes', [])] if st.session_state.get('last_uploaded_resumes') else []
            if current_resume_names != previous_resume_names:
                files_changed = True
                st.session_state.last_uploaded_resumes = uploaded_resumes_list
        
        if uploaded_jds_list:
            current_jd_names = [f.name for f in uploaded_jds_list] if uploaded_jds_list else []
            previous_jd_names = [f.name for f in st.session_state.get('last_uploaded_jds', [])] if st.session_state.get('last_uploaded_jds') else []
            if current_jd_names != previous_jd_names:
                files_changed = True
                st.session_state.last_uploaded_jds = uploaded_jds_list
        
        # Check if files are uploaded and need to be loaded
        files_need_loading = (
            (uploaded_resumes_list or uploaded_jds_list) and  # Files exist
            (
                files_changed or  # Files changed
                'matcher' not in st.session_state or  # No matcher yet
                'matching_mode' not in st.session_state or  # No mode set
                st.session_state.matching_mode != matching_mode  # Mode changed
            )
        )
        
        if files_need_loading:
            with st.spinner("🔄 Auto-loading uploaded files..."):
                if "Semantic" in matching_mode or "Improved Similarity" in matching_mode:
                    matcher = ResumeSemanticMatcher(None, use_onet_taxonomy=st.session_state.get('use_onet_taxonomy', False))
                else:
                    matcher = ResumeJDMatcher(None)
                
                matcher.load_documents_from_uploads(uploaded_resumes_list, uploaded_jds_list)
                
                if matcher.resumes or matcher.job_descriptions:
                    # Clear any existing similarity data
                    if 'similarity_matrix' in st.session_state:
                        del st.session_state.similarity_matrix
                    if 'section_matrices' in st.session_state:
                        del st.session_state.section_matrices
                    if 'similarities' in st.session_state:
                        del st.session_state.similarities
                    
                    st.session_state.matcher = matcher
                    st.session_state.matching_mode = matching_mode
                    st.success("✅ Files automatically loaded!")
                    st.rerun()
                elif uploaded_resumes_list or uploaded_jds_list:
                    st.warning("⚠️ Files uploaded but could not be processed. Please check file formats.")
    
    # Initialize matcher based on mode
    # Check if we need to recreate matcher due to mode change
    matcher_needs_recreation = (
        'matcher' not in st.session_state or 
        'matching_mode' not in st.session_state or 
        st.session_state.matching_mode != matching_mode
    )
    
    if st.sidebar.button("🔄 Load Documents") or matcher_needs_recreation:
        with st.spinner("Loading documents..."):
            if input_method == "📤 Upload Files":
                # Load from uploaded files
                uploaded_resumes_list = st.session_state.get('uploaded_resumes', None)
                uploaded_jds_list = st.session_state.get('uploaded_jds', None)
                
                if not uploaded_resumes_list and not uploaded_jds_list:
                    st.error("❌ Please upload at least one resume or job description file!")
                    return
                
                if "Semantic" in matching_mode or "Improved Similarity" in matching_mode:
                    matcher = ResumeSemanticMatcher(None, use_onet_taxonomy=st.session_state.get('use_onet_taxonomy', False))  # No directory needed
                else:
                    matcher = ResumeJDMatcher(None)  # No directory needed
                
                # Load documents from uploaded files
                matcher.load_documents_from_uploads(uploaded_resumes_list, uploaded_jds_list)
            else:
                # Load from directory (original method)
                if not data_dir:
                    st.error("❌ Please specify a directory path!")
                    return
                
                if "Semantic" in matching_mode or "Improved Similarity" in matching_mode:
                    matcher = ResumeSemanticMatcher(data_dir, use_onet_taxonomy=st.session_state.get('use_onet_taxonomy', False))
                else:
                    matcher = ResumeJDMatcher(data_dir)
                
                matcher.load_documents()
            
            # Check if documents were loaded successfully
            if matcher.resumes or matcher.job_descriptions:
                # Clear any existing similarity data to prevent dimension mismatches
                if 'similarity_matrix' in st.session_state:
                    del st.session_state.similarity_matrix
                if 'section_matrices' in st.session_state:
                    del st.session_state.section_matrices
                if 'similarities' in st.session_state:
                    del st.session_state.similarities
                
                st.session_state.matcher = matcher
                st.session_state.matching_mode = matching_mode
                st.success("✅ Documents loaded successfully!")
            else:
                st.error("❌ No documents found or failed to load!")
                return
    
    if 'matcher' not in st.session_state:
        st.info("👆 Please click 'Load Documents' to start")
        return
    
    matcher = st.session_state.matcher
    
    def build_robert_context(matcher_obj, candidate_name: Optional[str] = None, jd_name: Optional[str] = None) -> str:
        """Create a concise context string for Robert from selected resume/JD."""
        context_parts: List[str] = []
        
        def normalize_value(value: Any) -> str:
            if isinstance(value, str):
                return value.strip()
            if isinstance(value, list):
                return "; ".join(str(item).strip() for item in value if str(item).strip())
            if isinstance(value, dict):
                return "; ".join(f"{k}: {v}" for k, v in value.items() if v)
            return str(value)
        
        if candidate_name and candidate_name in matcher_obj.resumes:
            context_parts.append(f"Candidate: {candidate_name}")
            try:
                sections = matcher_obj.extract_sections(matcher_obj.resumes[candidate_name])
            except Exception:
                sections = {}
            for key, label, limit in [
                ('summary', 'Summary', 400),
                ('skills', 'Skills', 300),
                ('experience', 'Experience', 400),
                ('education', 'Education', 300)
            ]:
                value = sections.get(key)
                if not value:
                    continue
                normalized = normalize_value(value)
                if normalized:
                    context_parts.append(f"{label}: {normalized[:limit]}")
        
        if jd_name and jd_name in matcher_obj.job_descriptions:
            jd_text = matcher_obj.job_descriptions[jd_name]
            jd_preview = jd_text[:500] if jd_text else ""
            if jd_preview:
                context_parts.append(f"Job Description ({jd_name}): {jd_preview}")
        
        return "\n".join(context_parts).strip()
    
    def render_skill_insights(cache_prefix: str, skills_text: str, resume_text: str):
        """Render AI-generated skill taxonomy insights if available."""
        if not skills_text:
            return
        if not (st.session_state.get('use_llm', False) and LLM_AVAILABLE and is_llm_available()):
            return
        cache_key = f"skill_insights_{cache_prefix}"
        if cache_key not in st.session_state:
            try:
                llm_client = get_llm_client()
                st.session_state[cache_key] = llm_client.generate_skill_taxonomy(skills_text, resume_text)
            except Exception as exc:
                st.session_state[cache_key] = {"error": str(exc)}
        insights = st.session_state.get(cache_key)
        if not insights:
            return
        st.caption("AI Skill Insights")
        if isinstance(insights, dict):
            if "error" in insights:
                st.info(f"AI insight unavailable: {insights['error']}")
                return
            for category, values in insights.items():
                if not values:
                    continue
                label = category.replace('_', ' ').title()
                if isinstance(values, list):
                    st.write(f"**{label}:** {', '.join(values)}")
                else:
                    st.write(f"**{label}:** {values}")
        else:
            st.write(insights)
    
    def render_version_comparison(section_key_prefix: str):
        """Allow users to compare multiple resume versions for the same candidate."""
        if not matcher or not getattr(matcher, 'extracted_names', None):
            return
        version_groups: Dict[str, List[str]] = {}
        for std_name, base_name in matcher.extracted_names.items():
            key = base_name or std_name
            version_groups.setdefault(key, []).append(std_name)
        version_groups = {k: sorted(v) for k, v in version_groups.items() if len(v) > 1}
        if not version_groups:
            return
        st.markdown("### 🆚 Resume Version Comparison")
        base_candidate = st.selectbox(
            "Select candidate:",
            sorted(version_groups.keys()),
            key=f"version_candidate_{section_key_prefix}"
        )
        versions = version_groups.get(base_candidate, [])
        if len(versions) < 2:
            st.info("Need at least two versions to compare.")
            return
        col_a, col_b = st.columns(2)
        with col_a:
            version_a = st.selectbox(
                "Version A",
                versions,
                key=f"version_a_{section_key_prefix}"
            )
        with col_b:
            version_b_options = [v for v in versions if v != version_a]
            version_b = st.selectbox(
                "Version B",
                version_b_options or versions,
                key=f"version_b_{section_key_prefix}"
            )
        if version_a == version_b:
            st.warning("Please pick two different versions.")
            return
        if not (st.session_state.get('use_llm', False) and LLM_AVAILABLE and is_llm_available()):
            st.info("Enable AI enhancements and provide an API key to generate version differences.")
            return
        diff_cache_key = f"resume_diff_{version_a}_{version_b}"
        if st.button("Generate AI Version Diff", key=f"diff_btn_{section_key_prefix}"):
            try:
                llm_client = get_llm_client()
                text_a = matcher.resumes.get(version_a, "")
                text_b = matcher.resumes.get(version_b, "")
                st.session_state[diff_cache_key] = llm_client.compare_resume_versions(
                    text_a,
                    text_b,
                    version_a,
                    version_b
                )
            except Exception as exc:
                st.error(f"Unable to generate version diff: {exc}")
        diff_summary = st.session_state.get(diff_cache_key)
        if diff_summary:
            st.info(diff_summary)
    
    def render_robert_assistant():
        """Sidebar interface for the Robert AI helper."""
        st.sidebar.markdown("### 💬 Ask Robert")
        if 'robert_history' not in st.session_state:
            st.session_state.robert_history = []
        if not matcher:
            st.sidebar.info("Load documents to chat with Robert.")
            return
        if not (use_llm and LLM_AVAILABLE and is_llm_available()):
            st.sidebar.info("Enable AI enhancements and enter a Gemini API key to chat with Robert.")
            return
        
        candidate_options = ["(none)"] + matcher.candidate_names
        jd_options = ["(none)"] + matcher.jd_names
        
        selected_candidate_label = st.sidebar.selectbox(
            "Candidate (optional)",
            candidate_options,
            index=0,
            key="robert_candidate_select"
        )
        selected_candidate = selected_candidate_label if selected_candidate_label != "(none)" else None
        
        selected_jd_label = st.sidebar.selectbox(
            "Job Description (optional)",
            jd_options,
            index=0,
            key="robert_jd_select"
        )
        selected_jd = selected_jd_label if selected_jd_label != "(none)" else None
        
        if 'robert_question_input' not in st.session_state:
            st.session_state.robert_question_input = ""
        question = st.sidebar.text_area(
            "Your question for Robert",
            key="robert_question_input",
            height=100,
            placeholder="e.g., Why was the summary empty for Chloe?"
        )
        
        col_send, col_clear = st.sidebar.columns(2)
        with col_send:
            send_clicked = st.button("Ask Robert", key="robert_send_btn")
        with col_clear:
            clear_clicked = st.button("Clear Chat", key="robert_clear_btn")
        
        if clear_clicked:
            st.session_state.robert_history = []
            st.session_state.pop('robert_question_input', None)
            cache_keys = [key for key in st.session_state.keys() if key.startswith("robert_reply_")]
            for cache_key in cache_keys:
                del st.session_state[cache_key]
            st.rerun()
        
        if send_clicked:
            if not question.strip():
                st.sidebar.warning("Please enter a question for Robert.")
            else:
                try:
                    llm_client = get_llm_client()
                    context_text = build_robert_context(matcher, selected_candidate, selected_jd)
                    history = st.session_state.robert_history + [{"role": "user", "content": question.strip()}]
                    cache_key = None
                    if selected_candidate or selected_jd:
                        cache_base = f"{selected_candidate or 'none'}_{selected_jd or 'none'}_{len(history)}"
                        cache_key = f"robert_reply_{hashlib.sha256(cache_base.encode()).hexdigest()[:12]}"
                    response = llm_client.chat_with_robert(history, context=context_text, cache_key=cache_key)
                    history.append({"role": "assistant", "content": response})
                    st.session_state.robert_history = history
                    st.session_state.pop('robert_question_input', None)
                    st.rerun()
                except Exception as exc:
                    st.sidebar.error(f"Robert is unavailable: {exc}")
        
        if st.session_state.robert_history:
            st.sidebar.markdown("**Recent Replies:**")
            for entry in st.session_state.robert_history[-6:]:
                label = "You" if entry.get("role") == "user" else "Robert"
                message = entry.get("content", "")
                if message:
                    st.sidebar.write(f"**{label}:** {message}")
    
    render_robert_assistant()
    
    # File monitoring section
    st.subheader("📁 Directory Monitoring")
    
    # Show current directory status
    dir_info = matcher.get_directory_info()
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Resume Files", len(dir_info["resume_files"]))
    
    with col2:
        st.metric("Job Description Files", len(dir_info["jd_files"]))
    
    with col3:
        st.metric("Total Files", dir_info["total_files"])
    
    # Show file details
    if dir_info["resume_files"] or dir_info["jd_files"]:
        with st.expander("📋 File Details", expanded=False):
            if dir_info["resume_files"]:
                st.write("**Resume Files:**")
                for file_info in dir_info["resume_files"]:
                    size_info = f" ({file_info.get('size', 'N/A')} bytes)" if 'size' in file_info else ""
                    st.write(f"• {file_info['name']}{size_info}")
            
            if dir_info["jd_files"]:
                st.write("**Job Description Files:**")
                for file_info in dir_info["jd_files"]:
                    size_info = f" ({file_info.get('size', 'N/A')} bytes)" if 'size' in file_info else ""
                    st.write(f"• {file_info['name']}{size_info}")
    
    # Main content
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📄 Loaded Documents")
        
        st.write("**Resumes:**")
        for name in matcher.candidate_names:
            st.write(f"• {name}")
        
        st.write("**Job Descriptions:**")
        for name in matcher.jd_names:
            st.write(f"• {name}")
    
    with col2:
        st.subheader("⚙️ Processing Options")
        
        # Check if we need to recompute due to weight changes
        should_recompute = (
            st.button("🔍 Compute Similarity", type="primary") or 
            st.session_state.get('weights_changed', False)
        )
        
        if should_recompute:
            if "Semantic" in matching_mode:
                with st.spinner("Computing semantic similarity with Sentence-BERT..."):
                    result = matcher.compute_semantic_similarity(
                        education_weight=education_weight,
                        skills_weight=skills_weight,
                        experience_weight=experience_weight,
                        summary_weight=summary_weight
                    )
                    
                    if result is not None:
                        similarity_matrix, section_matrices, similarities = result
                        st.session_state.similarity_matrix = similarity_matrix
                        st.session_state.section_matrices = section_matrices
                        st.session_state.similarities = similarities
                        st.session_state.weights_changed = False  # Reset the flag
                        st.success("✅ Semantic similarity computation completed!")
                    else:
                        st.error("❌ Failed to compute semantic similarity!")
            elif "Improved Similarity" in matching_mode:
                with st.spinner("Computing improved similarity with preprocessing..."):
                    result = matcher.compute_improved_similarity_matrix()
                    
                    if result is not None:
                        similarity_matrix, similarities = result
                        st.session_state.similarity_matrix = similarity_matrix
                        st.session_state.similarities = similarities
                        st.session_state.weights_changed = False  # Reset the flag
                        st.success("✅ Improved similarity computation completed!")
                    else:
                        st.error("❌ Failed to compute improved similarity!")
            else:
                st.error("❌ Invalid matching mode selected!")
    
    # Results section
    if 'similarity_matrix' in st.session_state and st.session_state.similarity_matrix is not None:
        st.markdown("---")
        
        # Mode indicator
        current_mode = st.session_state.get('matching_mode', 'TF-IDF')
        if "Semantic" in current_mode:
            st.subheader("🧠 Semantic Results Dashboard")
        else:
            st.subheader("📊 TF-IDF Results Dashboard")
        
        similarity_matrix = st.session_state.similarity_matrix
        
        # Export functionality and cache management
        col1, col2, col3 = st.columns([2, 1, 1])
        with col2:
            if st.button("📥 Export Results", help="Download results as CSV"):
                # Create exportable dataframe
                export_data = []
                for jd_idx, jd_name in enumerate(matcher.jd_names):
                    for resume_idx, resume_name in enumerate(matcher.candidate_names):
                        row = {
                            'Job Description': jd_name,
                            'Candidate': resume_name,
                            'Total Similarity': similarity_matrix[resume_idx, jd_idx],
                            'Match Percentage': f"{similarity_matrix[resume_idx, jd_idx] * 100:.1f}%"
                        }
                        
                        # Add section scores for semantic mode
                        if "Semantic" in current_mode and 'section_matrices' in st.session_state:
                            section_matrices = st.session_state.section_matrices
                            for section in ['education', 'skills', 'experience']:
                                if section in section_matrices:
                                    row[f'{section.title()} Similarity'] = section_matrices[section][resume_idx, jd_idx]
                        
                        # Add explanation for semantic mode
                        if "Semantic" in current_mode and 'similarities' in st.session_state:
                            similarities = st.session_state.similarities
                            if jd_name in similarities and resume_name in similarities[jd_name]:
                                explanation = matcher.generate_explanation(resume_name, jd_name, similarities)
                                row['Explanation'] = explanation
                        
                        export_data.append(row)
                
                export_df = pd.DataFrame(export_data)
                csv = export_df.to_csv(index=False)
                st.download_button(
                    label="Download CSV",
                    data=csv,
                    file_name=f"resume_matching_results_{current_mode.lower().replace(' ', '_')}.csv",
                    mime="text/csv"
                )
        
        with col3:
            if st.button("🔄 Clear Cache", help="Clear cached results to force recalculation"):
                # Clear all caches
                cache_keys_to_clear = [key for key in st.session_state.keys() if 'cache' in key.lower() or 'embedding' in key.lower()]
                for key in cache_keys_to_clear:
                    del st.session_state[key]
                st.success("✅ Cache cleared! Results will be recalculated.")
                st.rerun()
        
        # Create tabs for different views
        if "Semantic" in current_mode:
            tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
                "🔥 Overall Match Score", 
                "📈 Top Candidates", 
                "📊 Detailed Analysis", 
                "🎯 Component Analysis (Education/Skills/Experience)", 
                "👤 Candidate Profiles", 
                "📄 Job Requirements", 
                "✨ AI-Generated Summaries"
            ])
        elif "Improved Similarity" in current_mode:
            tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
                "🔥 Overall Match Score", 
                "📈 Top Candidates", 
                "📊 Detailed Analysis", 
                "🚀 Similarity Breakdown", 
                "👤 Candidate Profiles", 
                "📄 Job Requirements", 
                "✨ AI-Generated Summaries"
            ])
        else:
            st.error("❌ Invalid mode configuration!")
        
        with tab1:
            st.subheader("🔥 Overall Match Score Heatmap")
            st.markdown("**Visual representation of how well each candidate matches each job description**")
            
            # Debug: Check dimensions
            print(f"DEBUG: Similarity matrix shape: {similarity_matrix.shape}")
            print(f"DEBUG: Candidate names count: {len(matcher.candidate_names)}")
            print(f"DEBUG: JD names count: {len(matcher.jd_names)}")
            
            # Ensure matrix dimensions match candidate names
            if similarity_matrix.shape[0] != len(matcher.candidate_names):
                st.warning(f"⚠️ Matrix shape mismatch: {similarity_matrix.shape[0]} rows vs {len(matcher.candidate_names)} candidates")
                # Create a new matrix with correct dimensions
                new_matrix = np.zeros((len(matcher.candidate_names), len(matcher.jd_names)))
                # Copy existing data if possible
                min_rows = min(similarity_matrix.shape[0], len(matcher.candidate_names))
                min_cols = min(similarity_matrix.shape[1], len(matcher.jd_names))
                new_matrix[:min_rows, :min_cols] = similarity_matrix[:min_rows, :min_cols]
                similarity_matrix = new_matrix
                st.info(f"✅ Adjusted matrix to shape: {similarity_matrix.shape}")
            
            if similarity_matrix.shape[1] != len(matcher.jd_names):
                st.warning(f"⚠️ Matrix shape mismatch: {similarity_matrix.shape[1]} cols vs {len(matcher.jd_names)} JDs")
                # Create a new matrix with correct dimensions
                new_matrix = np.zeros((len(matcher.candidate_names), len(matcher.jd_names)))
                # Copy existing data if possible
                min_rows = min(similarity_matrix.shape[0], len(matcher.candidate_names))
                min_cols = min(similarity_matrix.shape[1], len(matcher.jd_names))
                new_matrix[:min_rows, :min_cols] = similarity_matrix[:min_rows, :min_cols]
                similarity_matrix = new_matrix
                st.info(f"✅ Adjusted matrix to shape: {similarity_matrix.shape}")
            
            fig, df = create_heatmap(similarity_matrix, matcher.candidate_names, matcher.jd_names)
            st.plotly_chart(fig, use_container_width=True)
            
            # Display raw similarity matrix with better formatting
            st.subheader("📊 Detailed Score Matrix")
            st.markdown("**Complete similarity scores for all candidate-job pairs**")
            
            # Color-code the dataframe
            def color_similarity(val):
                if isinstance(val, (int, float)):
                    if val >= 0.7:
                        return 'background-color: #90EE90'  # Light green
                    elif val >= 0.4:
                        return 'background-color: #FFE4B5'  # Light yellow
                    else:
                        return 'background-color: #FFB6C1'  # Light red
                return ''
            
            styled_df = df.style.applymap(color_similarity)
            st.dataframe(styled_df, use_container_width=True)
            
            # Show statistics
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Max Score", f"{df.max().max():.3f}")
            with col2:
                st.metric("Avg Score", f"{df.mean().mean():.3f}")
            with col3:
                non_zero = (df > 0).sum().sum()
                total = df.size
                st.metric("Non-Zero Matches", f"{non_zero}/{total}")
            with col4:
                high_matches = (df >= 0.7).sum().sum()
                st.metric("Strong Matches (≥70%)", high_matches)
        
        with tab2:
            st.subheader(f"📈 Top {top_k} Candidates per Job")
            st.markdown("**Ranked list of best-matching candidates for each position**")
            top_matches = matcher.get_top_matches(top_k)
            
            if not top_matches.empty:
                # Display as cards
                for jd in top_matches['Job Description'].unique():
                    st.write(f"**{jd}**")
                    jd_matches = top_matches[top_matches['Job Description'] == jd]
                    
                    for _, match in jd_matches.iterrows():
                        score = match['Similarity Score']
                        color = "🟢" if score > 0.7 else "🟡" if score > 0.4 else "🔴"
                        
                        st.write(f"{color} **{match['Candidate']}** - {match['Match Percentage']}")
                    
                    st.write("---")
            else:
                st.warning("No matches found!")
        
        with tab3:
            st.subheader("📊 Detailed Analysis")
            st.markdown("**Complete similarity analysis with all scores and explanations**")
            
            # Summary statistics
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("Average Similarity", f"{similarity_matrix.mean():.3f}")
            
            with col2:
                st.metric("Highest Match", f"{similarity_matrix.max():.3f}")
            
            with col3:
                st.metric("Lowest Match", f"{similarity_matrix.min():.3f}")
            
            # Detailed table
            st.subheader("All Similarity Scores")
            
            # Ensure unique names to avoid non-unique index error
            unique_candidate_names = []
            unique_jd_names = []
            
            # Handle duplicate candidate names
            candidate_name_counts = {}
            for name in matcher.candidate_names:
                if name in candidate_name_counts:
                    candidate_name_counts[name] += 1
                    unique_candidate_names.append(f"{name} ({candidate_name_counts[name]})")
                else:
                    candidate_name_counts[name] = 1
                    unique_candidate_names.append(name)
            
            # Handle duplicate JD names
            jd_name_counts = {}
            for name in matcher.jd_names:
                if name in jd_name_counts:
                    jd_name_counts[name] += 1
                    unique_jd_names.append(f"{name} ({jd_name_counts[name]})")
                else:
                    jd_name_counts[name] = 1
                    unique_jd_names.append(name)
            
            detailed_df = pd.DataFrame(
                similarity_matrix,
                index=unique_candidate_names,
                columns=unique_jd_names
            )
            
            # Add color coding
            def color_similarity(val):
                if val > 0.7:
                    return 'background-color: #d4edda'  # Green
                elif val > 0.4:
                    return 'background-color: #fff3cd'  # Yellow
                else:
                    return 'background-color: #f8d7da'  # Red
            
            styled_df = detailed_df.style.applymap(color_similarity)
            st.dataframe(styled_df, use_container_width=True)
        
        # Section breakdown tab (only for semantic mode)
        if "Semantic" in current_mode and 'section_matrices' in st.session_state:
            with tab4:
                st.subheader("🎯 Component Analysis: Education, Skills & Experience")
                st.markdown("**Breakdown of match scores by component to understand what drives each match**")
                
                section_matrices = st.session_state.section_matrices
                
                # Better section names
                section_display_names = {
                    'education': '🎓 Academic Qualifications & Education',
                    'skills': '💻 Technical Skills & Competencies',
                    'experience': '💼 Work Experience & Career History',
                    'summary': '📝 Professional Summary & Overview'
                }
                
                # Create section heatmaps with better naming
                for section_name, section_matrix in section_matrices.items():
                    if section_name in section_display_names and section_matrix.max() > 0:
                        display_name = section_display_names[section_name]
                        st.markdown(f"### {display_name}")
                        
                        # Show API enhancement status if available and enabled
                        use_llm_enabled = st.session_state.get('use_llm', False)
                        if use_llm_enabled and LLM_AVAILABLE and is_llm_available():
                            st.caption("🤖 Enhanced with AI-powered extraction")
                        
                        fig, _ = create_heatmap(section_matrix, matcher.candidate_names, matcher.jd_names)
                        fig.update_layout(
                            title=f"{display_name} - Similarity Scores",
                            height=400
                        )
                        st.plotly_chart(fig, use_container_width=True)
                        
                        # Show statistics
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Max Score", f"{section_matrix.max():.3f}")
                        with col2:
                            st.metric("Avg Score", f"{section_matrix.mean():.3f}")
                        with col3:
                            non_zero = np.count_nonzero(section_matrix)
                            total = section_matrix.size
                            st.metric("Non-Zero Matches", f"{non_zero}/{total}")
                        
                        st.write("---")
                
                # Detailed section scores table with better naming
                st.subheader("📊 Component Score Details")
                st.markdown("**Detailed breakdown showing how each component contributes to the overall match**")
                
                section_data = []
                for jd_idx, jd_name in enumerate(matcher.jd_names):
                    for resume_idx, resume_name in enumerate(matcher.candidate_names):
                        row = {
                            'Job Description': jd_name,
                            'Candidate': resume_name,
                            'Overall Match Score': f"{similarity_matrix[resume_idx, jd_idx]:.3f}",
                            'Match Percentage': f"{similarity_matrix[resume_idx, jd_idx] * 100:.1f}%"
                        }
                        
                        # Use better section names
                        section_display_map = {
                            'education': '🎓 Education',
                            'skills': '💻 Skills',
                            'experience': '💼 Experience',
                            'summary': '📝 Summary'
                        }
                        
                        for section_name, section_matrix in section_matrices.items():
                            if section_name in section_display_map:
                                score = section_matrix[resume_idx, jd_idx]
                                row[section_display_map[section_name]] = f"{score:.3f}"
                        
                        # Add explanation
                        if 'similarities' in st.session_state:
                            similarities = st.session_state.similarities
                            if jd_name in similarities and resume_name in similarities[jd_name]:
                                explanation = matcher.generate_explanation(resume_name, jd_name, similarities)
                                row['AI Explanation'] = explanation
                        
                        section_data.append(row)
                
                section_df = pd.DataFrame(section_data)
                st.dataframe(section_df, use_container_width=True, height=400)
        
        # Graded breakdown tab (only for graded scoring mode)
        if "Graded Scoring" in current_mode and 'similarities' in st.session_state:
            with tab4:
                st.subheader("⭐ Graded Scoring Breakdown")
                
                similarities = st.session_state.similarities
                
                # Select candidate and JD for detailed breakdown
                col1, col2 = st.columns(2)
                with col1:
                    selected_candidate = st.selectbox(
                        "Select Candidate:",
                        matcher.candidate_names,
                        key="graded_candidate"
                    )
                with col2:
                    selected_jd = st.selectbox(
                        "Select Job Description:",
                        matcher.jd_names,
                        key="graded_jd"
                    )
                
                if selected_jd in similarities and selected_candidate in similarities[selected_jd]:
                    match_data = similarities[selected_jd][selected_candidate]
                    
                    # Display score breakdown
                    st.subheader("📊 Score Breakdown")
                    
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric(
                            "Required Skills",
                            f"{match_data['required_skills']:.2f}",
                            help="Binary gates for must-have skills"
                        )
                    with col2:
                        st.metric(
                            "Preferred Skills",
                            f"{match_data['preferred_skills']:.2f}",
                            help="Weighted bonuses for nice-to-have skills"
                        )
                    with col3:
                        st.metric(
                            "YOE Score",
                            f"{match_data['yoe_score']:.2f}",
                            help="Years of experience scaling"
                        )
                    with col4:
                        st.metric(
                            "Total Score",
                            f"{match_data['total']:.2f}",
                            help="Composite score: 60% required + 30% preferred + 10% YOE"
                        )
                    
                    # Display detailed breakdown
                    st.subheader("🔍 Detailed Breakdown")
                    
                    for item in match_data['breakdown']:
                        if item.startswith("✅"):
                            st.success(item)
                        elif item.startswith("❌"):
                            st.error(item)
                        elif item.startswith("🎯"):
                            st.info(item)
                        elif item.startswith("⚪"):
                            st.warning(item)
                        elif item.startswith("📅"):
                            st.info(item)
                        else:
                            st.write(item)
                    
                    # Display formula explanation
                    st.subheader("🧮 Scoring Formula")
                    st.info("""
                    **Composite Score Formula:**
                    - Required Skills: 60% weight (binary gates)
                    - Preferred Skills: 30% weight (weighted bonuses)
                    - Years of Experience: 10% weight (linear scaling)
                    
                    **Skill Similarity Examples:**
                    - Kafka ↔ Flink: 0.8 similarity
                    - PyTorch ↔ TensorFlow: 0.8 similarity
                    - MLflow ↔ Weights & Biases: 0.7 similarity
                    """)
        
                else:
                    st.warning("No similarity data available for the selected candidate and job description.")
        
        # Similarity breakdown tab (only for improved similarity mode)
        if "Improved Similarity" in current_mode and 'similarities' in st.session_state:
            with tab4:
                st.subheader("🚀 Improved Similarity Breakdown")
                
                similarities = st.session_state.similarities
                
                # Select candidate and JD for detailed breakdown
                col1, col2 = st.columns(2)
                with col1:
                    selected_candidate = st.selectbox(
                        "Select Candidate:",
                        matcher.candidate_names,
                        key="improved_candidate"
                    )
                with col2:
                    selected_jd = st.selectbox(
                        "Select Job Description:",
                        matcher.jd_names,
                        key="improved_jd"
                    )
                
                if selected_jd in similarities and selected_candidate in similarities[selected_jd]:
                    match_data = similarities[selected_jd][selected_candidate]
                    
                    # Display overall metrics
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Raw Similarity", f"{match_data['raw_similarity']:.4f}")
                    with col2:
                        st.metric("Display Score", f"{match_data['display_score']:.1f}%")
                    with col3:
                        st.metric("Match Quality", 
                                "Excellent" if match_data['display_score'] > 80 else
                                "Good" if match_data['display_score'] > 60 else
                                "Fair" if match_data['display_score'] > 40 else "Poor")
                    
                    # Section breakdown
                    st.subheader("📊 Section-by-Section Breakdown")
                    breakdown = match_data['breakdown']
                    
                    breakdown_data = []
                    for section, score in breakdown.items():
                        breakdown_data.append({
                            'Section': section.title(),
                            'Raw Similarity': f"{score:.4f}",
                            'Display Score': f"{score * 100:.1f}%",
                            'Quality': "High" if score > 0.6 else "Medium" if score > 0.3 else "Low"
                        })
                    
                    breakdown_df = pd.DataFrame(breakdown_data)
                    st.dataframe(breakdown_df, use_container_width=True)
                    
                    # Visual breakdown
                    st.subheader("📈 Similarity Visualization")
                    fig = px.bar(
                        breakdown_df, 
                        x='Section', 
                        y=[float(score) for score in breakdown_df['Raw Similarity']],
                        title=f"Section Similarity Scores: {selected_candidate} vs {selected_jd}",
                        color=[float(score) for score in breakdown_df['Raw Similarity']],
                        color_continuous_scale='RdYlGn'
                    )
                    fig.update_layout(yaxis_title="Similarity Score", showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Processing details
                    st.subheader("🔧 Processing Details")
                    st.info("""
                    **Improved Similarity Processing:**
                    - ✅ Text preprocessing removes noise (URLs, phone numbers, emails)
                    - ✅ Weighted embeddings emphasize skills (40%) and experience (40%)
                    - ✅ Sigmoid scaling improves score distribution
                    - ✅ Section-by-section analysis for interpretability
                    """)
                else:
                    st.warning("No similarity data available for the selected candidate and job description.")
        
        # Resume Details tab
        if "Semantic" in current_mode:
            with tab5:
                st.subheader("👤 Resume Details Dashboard")
                
                # Candidate selection
                selected_candidate = st.selectbox(
                    "Select a candidate to view details:",
                    matcher.candidate_names,
                    key="resume_selector_semantic"
                )
                
                if selected_candidate and selected_candidate in matcher.resumes:
                    resume_text = matcher.resumes[selected_candidate]
                    
                    # Use proper section extraction instead of basic extract_resume_summary
                    if hasattr(matcher, 'extract_sections'):
                        sections = matcher.extract_sections(resume_text)
                    else:
                        # Fallback to basic extraction
                        sections = {
                            'education': '',
                            'skills': '',
                            'experience': '',
                            'summary': '',
                            'raw_text': resume_text
                        }
                    
                    # Try to get API-generated profile if enabled
                    api_profile = None
                    use_llm_enabled = st.session_state.get('use_llm', False)
                    if use_llm_enabled and LLM_AVAILABLE and is_llm_available():
                        try:
                            llm_client = get_llm_client()
                            with st.spinner("🤖 Generating AI-enhanced profile..."):
                                api_profile = llm_client.generate_candidate_profile(resume_text)
                        except Exception as e:
                            print(f"DEBUG: API profile generation failed: {e}")
                    
                    # Extract candidate name
                    candidate_name = matcher.extract_candidate_name(resume_text)
                    
                    # Display resume details with clearer layout and source indicators
                    contact_col, summary_col = st.columns([1, 2])
                    
                    with contact_col:
                        st.markdown("### 📋 Contact Information")
                        if api_profile and api_profile.get('contact_information'):
                            st.caption("Source: Gemini API")
                            contact_info = api_profile['contact_information']
                            if 'Name:' in contact_info:
                                name_match = re.search(r'Name:\s*([^\n,]+)', contact_info)
                                if name_match:
                                    st.write(f"**Name:** {name_match.group(1).strip()}")
                            else:
                                st.write(f"**Name:** {candidate_name}")
                            if 'Email:' in contact_info:
                                email_match = re.search(r'Email:\s*([^\n,]+)', contact_info)
                                if email_match:
                                    st.write(f"**Email:** {email_match.group(1).strip()}")
                            if 'Phone:' in contact_info:
                                phone_match = re.search(r'Phone:\s*([^\n,]+)', contact_info)
                                if phone_match:
                                    st.write(f"**Phone:** {phone_match.group(1).strip()}")
                            if 'Location:' in contact_info:
                                location_match = re.search(r'Location:\s*([^\n,]+)', contact_info)
                                if location_match:
                                    st.write(f"**Location:** {location_match.group(1).strip()}")
                        else:
                            st.caption("Source: Resume text")
                            st.write(f"**Name:** {candidate_name}")
                            email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', resume_text)
                            if email_match:
                                st.write(f"**Email:** {email_match.group(0)}")
                            phone_match = re.search(r'\+?1?[-.\s]?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})', resume_text)
                            if phone_match:
                                phone = f"({phone_match.group(1)}) {phone_match.group(2)}-{phone_match.group(3)}"
                                st.write(f"**Phone:** {phone}")
                    
                    with summary_col:
                        st.markdown("### 🎯 Professional Summary")
                        summary_text = ""
                        summary_source = ""
                        if api_profile and api_profile.get('professional_summary'):
                            summary_text = api_profile['professional_summary'].strip()
                            summary_source = "Gemini API"
                        elif sections.get('summary'):
                            summary_text = sections['summary'].strip()
                            summary_source = "Resume summary section"
                        if (not summary_text or len(summary_text) < 80) and hasattr(matcher, 'generate_professional_summary'):
                            generated_summary = matcher.generate_professional_summary(sections)
                            if generated_summary:
                                summary_text = generated_summary.strip()
                                summary_source = "Rule-based generator"
                        if summary_text:
                            summary_text = re.sub(r'^[•\-\*\d+\.\)]\s*', '', summary_text, flags=re.MULTILINE)
                            st.caption(f"Source: {summary_source or 'Resume text'}")
                            st.write(summary_text)
                        else:
                            st.caption("Source: Resume text")
                            st.info("No summary available")
                    
                    st.divider()
                    
                    skills_col, experience_col = st.columns([1, 2])
                    
                    with skills_col:
                        st.markdown("### 🛠️ Skills")
                        skills_text = ""
                        skills_source = ""
                        if api_profile and api_profile.get('skills'):
                            skills_text = api_profile['skills'].strip()
                            skills_source = "Gemini API"
                        elif sections.get('skills'):
                            skills_text = sections['skills'].strip()
                            skills_source = "Resume skills section"
                        elif SKILLS_DATABASE_AVAILABLE:
                            extracted_skills = extract_skills_from_text(resume_text)
                            if extracted_skills:
                                skills_text = ", ".join(extracted_skills[:30])
                                skills_source = "Skills database fallback"
                        if skills_text:
                            st.caption(f"Source: {skills_source or 'Resume text'}")
                            if ',' in skills_text:
                                skills_list = [s.strip() for s in skills_text.split(',') if s.strip()]
                                st.write(", ".join(skills_list))
                            else:
                                st.write(skills_text)
                        else:
                            st.caption("Source: Resume text")
                            st.info("No skills information available")
                        render_skill_insights(f"standard_{selected_candidate}", skills_text, resume_text)
                        render_skill_insights(f"semantic_{selected_candidate}", skills_text, resume_text)
                    
                        st.markdown("### 🎓 Education")
                        edu_text = ""
                        edu_source = ""
                        if api_profile and api_profile.get('education'):
                            edu_text = api_profile['education'].strip()
                            edu_source = "Gemini API"
                        elif sections.get('education'):
                            edu_text = sections['education'].strip()
                            edu_source = "Resume education section"
                        if not edu_text:
                            fallback_edu = matcher._extract_education_fallback(resume_text)
                            if fallback_edu:
                                edu_text = fallback_edu
                                edu_source = "Heuristic fallback"
                        if edu_text:
                            st.caption(f"Source: {edu_source or 'Resume text'}")
                            structured_entries = matcher._structure_education_entries(edu_text)
                            if structured_entries:
                                for entry in structured_entries:
                                    parts = []
                                    if entry['school']:
                                        parts.append(f"**{entry['school']}**")
                                    degree_major = entry['degree']
                                    if entry['major']:
                                        degree_major = f"{degree_major or ''}{' in ' if degree_major else ''}{entry['major']}"
                                    if degree_major:
                                        parts.append(degree_major.strip())
                                    if entry['year']:
                                        parts.append(f"({entry['year']})")
                                    st.write(" ".join(part for part in parts if part))
                            else:
                                cleaned_edu = re.sub(r'\s+', ' ', edu_text)
                                st.write(cleaned_edu)
                        else:
                            st.caption("Source: Resume text")
                            st.info("No education information available")
                    
                    with experience_col:
                        st.markdown("### 💼 Work Experience")
                        exp_text = ""
                        exp_source = ""
                        if api_profile and api_profile.get('work_experience'):
                            exp_text = api_profile['work_experience'].strip()
                            exp_source = "Gemini API"
                        elif sections.get('experience'):
                            exp_text = sections['experience'].strip()
                            exp_source = "Resume experience section"
                        if exp_text:
                            exp_text = re.sub(r'^[•\-\*\d+\.\)]\s*', '', exp_text, flags=re.MULTILINE)
                            exp_text = re.sub(r'\s+', ' ', exp_text)
                            st.caption(f"Source: {exp_source or 'Resume text'}")
                            if len(exp_text) > 500:
                                st.write(exp_text[:500] + "…")
                                with st.expander("Show full experience"):
                                    st.write(exp_text)
                            else:
                                st.write(exp_text)
                        else:
                            st.caption("Source: Resume text")
                            st.info("No work experience information available")
                    
                    # Similarity scores for this candidate
                    if 'similarity_matrix' in st.session_state and st.session_state.similarity_matrix is not None:
                        st.markdown("### 📊 Match Scores")
                        candidate_idx = matcher.candidate_names.index(selected_candidate)
                        candidate_scores = similarity_matrix[candidate_idx]
                        
                        scores_df = pd.DataFrame({
                            'Job Description': matcher.jd_names,
                            'Similarity Score': candidate_scores,
                            'Match Percentage': [f"{score * 100:.1f}%" for score in candidate_scores]
                        }).sort_values('Similarity Score', ascending=False)
                        
                        st.dataframe(scores_df, use_container_width=True)
                    
                    render_version_comparison("standard")
                    
                    render_version_comparison("semantic")
        else:
            with tab4:
                st.subheader("👤 Resume Details Dashboard")
                
                # Candidate selection
                selected_candidate = st.selectbox(
                    "Select a candidate to view details:",
                    matcher.candidate_names,
                    key="resume_selector"
                )
                
                if selected_candidate and selected_candidate in matcher.resumes:
                    resume_text = matcher.resumes[selected_candidate]
                    
                    # Use proper section extraction instead of basic extract_resume_summary
                    if hasattr(matcher, 'extract_sections'):
                        sections = matcher.extract_sections(resume_text)
                    else:
                        # Fallback to basic extraction
                        sections = {
                            'education': '',
                            'skills': '',
                            'experience': '',
                            'summary': '',
                            'raw_text': resume_text
                        }
                    
                    api_profile = None
                    use_llm_enabled = st.session_state.get('use_llm', False)
                    if use_llm_enabled and LLM_AVAILABLE and is_llm_available():
                        try:
                            llm_client = get_llm_client()
                            with st.spinner("🤖 Generating AI-enhanced profile..."):
                                api_profile = llm_client.generate_candidate_profile(resume_text)
                        except Exception as e:
                            print(f"DEBUG: API profile generation failed: {e}")
                    
                    # Extract candidate name
                    candidate_name = matcher.extract_candidate_name(resume_text)
                    
                    # Display resume details with enhanced layout
                    contact_col, summary_col = st.columns([1, 2])
                    
                    with contact_col:
                        st.markdown("### 📋 Contact Information")
                        if api_profile and api_profile.get('contact_information'):
                            st.caption("Source: Gemini API")
                            contact_info = api_profile['contact_information']
                            if 'Name:' in contact_info:
                                name_match = re.search(r'Name:\s*([^\n,]+)', contact_info)
                                if name_match:
                                    st.write(f"**Name:** {name_match.group(1).strip()}")
                            else:
                                st.write(f"**Name:** {candidate_name}")
                            if 'Email:' in contact_info:
                                email_match = re.search(r'Email:\s*([^\n,]+)', contact_info)
                                if email_match:
                                    st.write(f"**Email:** {email_match.group(1).strip()}")
                            if 'Phone:' in contact_info:
                                phone_match = re.search(r'Phone:\s*([^\n,]+)', contact_info)
                                if phone_match:
                                    st.write(f"**Phone:** {phone_match.group(1).strip()}")
                            if 'Location:' in contact_info:
                                location_match = re.search(r'Location:\s*([^\n,]+)', contact_info)
                                if location_match:
                                    st.write(f"**Location:** {location_match.group(1).strip()}")
                        else:
                            st.caption("Source: Resume text")
                            st.write(f"**Name:** {candidate_name}")
                            email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', resume_text)
                            if email_match:
                                st.write(f"**Email:** {email_match.group(0)}")
                            phone_match = re.search(r'\+?1?[-.\s]?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})', resume_text)
                            if phone_match:
                                phone = f"({phone_match.group(1)}) {phone_match.group(2)}-{phone_match.group(3)}"
                                st.write(f"**Phone:** {phone}")
                    
                    with summary_col:
                        st.markdown("### 🎯 Professional Summary")
                        summary_text = ""
                        summary_source = ""
                        if api_profile and api_profile.get('professional_summary'):
                            summary_text = api_profile['professional_summary'].strip()
                            summary_source = "Gemini API"
                        elif sections.get('summary'):
                            summary_text = sections['summary'].strip()
                            summary_source = "Resume summary section"
                        if summary_text:
                            summary_text = re.sub(r'^[•\-\*\d+\.\)]\s*', '', summary_text, flags=re.MULTILINE)
                            st.caption(f"Source: {summary_source or 'Resume text'}")
                            st.write(summary_text)
                        else:
                            st.caption("Source: Resume text")
                            st.info("No summary available")
                    
                    st.divider()
                    
                    skills_col, experience_col = st.columns([1, 2])
                    
                    with skills_col:
                        st.markdown("### 🛠️ Skills")
                        skills_text = ""
                        skills_source = ""
                        if api_profile and api_profile.get('skills'):
                            skills_text = api_profile['skills'].strip()
                            skills_source = "Gemini API"
                        elif sections.get('skills'):
                            skills_text = sections['skills'].strip()
                            skills_source = "Resume skills section"
                        elif SKILLS_DATABASE_AVAILABLE:
                            extracted_skills = extract_skills_from_text(resume_text)
                            if extracted_skills:
                                skills_text = ", ".join(extracted_skills[:30])
                                skills_source = "Skills database fallback"
                        if skills_text:
                            st.caption(f"Source: {skills_source or 'Resume text'}")
                            if ',' in skills_text:
                                skills_list = [s.strip() for s in skills_text.split(',') if s.strip()]
                                st.write(", ".join(skills_list))
                            else:
                                st.write(skills_text)
                        else:
                            st.caption("Source: Resume text")
                            st.info("No skills information available")
                        render_skill_insights(f"standard_{selected_candidate}", skills_text, resume_text)
                        render_skill_insights(f"semantic_{selected_candidate}", skills_text, resume_text)
                    
                        st.markdown("### 🎓 Education")
                        edu_text = ""
                        edu_source = ""
                        if api_profile and api_profile.get('education'):
                            edu_text = api_profile['education'].strip()
                            edu_source = "Gemini API"
                        elif sections.get('education'):
                            edu_text = sections['education'].strip()
                            edu_source = "Resume education section"
                        if not edu_text:
                            fallback_edu = matcher._extract_education_fallback(resume_text)
                            if fallback_edu:
                                edu_text = fallback_edu
                                edu_source = "Heuristic fallback"
                        if edu_text:
                            st.caption(f"Source: {edu_source or 'Resume text'}")
                            structured_entries = matcher._structure_education_entries(edu_text)
                            if structured_entries:
                                for entry in structured_entries:
                                    parts = []
                                    if entry['school']:
                                        parts.append(f"**{entry['school']}**")
                                    degree_major = entry['degree']
                                    if entry['major']:
                                        degree_major = f"{degree_major or ''}{' in ' if degree_major else ''}{entry['major']}"
                                    if degree_major:
                                        parts.append(degree_major.strip())
                                    if entry['year']:
                                        parts.append(f"({entry['year']})")
                                    st.write(" ".join(part for part in parts if part))
                        else:
                            st.caption("Source: Resume text")
                            st.info("No education information available")
                    
                    with experience_col:
                        st.markdown("### 💼 Work Experience")
                        exp_text = ""
                        exp_source = ""
                        if api_profile and api_profile.get('work_experience'):
                            exp_text = api_profile['work_experience'].strip()
                            exp_source = "Gemini API"
                        elif sections.get('experience'):
                            exp_text = sections['experience'].strip()
                            exp_source = "Resume experience section"
                        if exp_text:
                            exp_text = re.sub(r'^[•\-\*\d+\.\)]\s*', '', exp_text, flags=re.MULTILINE)
                            exp_text = re.sub(r'\s+', ' ', exp_text)
                            st.caption(f"Source: {exp_source or 'Resume text'}")
                            if len(exp_text) > 500:
                                st.write(exp_text[:500] + "…")
                                with st.expander("Show full experience"):
                                    st.write(exp_text)
                            else:
                                st.write(exp_text)
                        else:
                            st.caption("Source: Resume text")
                            st.info("No work experience information available")
                    
                    # Similarity scores for this candidate
                    st.markdown("### 📊 Match Scores")
                    candidate_idx = matcher.candidate_names.index(selected_candidate)
                    candidate_scores = similarity_matrix[candidate_idx]
                    
                    scores_df = pd.DataFrame({
                        'Job Description': matcher.jd_names,
                        'Similarity Score': candidate_scores,
                        'Match Percentage': [f"{score * 100:.1f}%" for score in candidate_scores]
                    }).sort_values('Similarity Score', ascending=False)
                    
                    st.dataframe(scores_df, use_container_width=True)
        
        # Job Description tab
        with tab6:
            st.subheader("📄 Job Description Dashboard")
            st.markdown("*View and analyze job descriptions with summaries and requirements*")
            
            # Job description selection
            selected_jd = st.selectbox(
                "Select a job description to view:",
                matcher.jd_names,
                key="jd_selector"
            )
            
            if selected_jd and selected_jd in matcher.job_descriptions:
                jd_text = matcher.job_descriptions[selected_jd]
                
                # Display job description summary
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("### 📋 Job Information")
                    # Extract basic job info
                    lines = jd_text.split('\n')
                    title = ""
                    company = ""
                    location = ""
                    
                    for line in lines[:10]:  # Check first 10 lines
                        line = line.strip()
                        if not title and any(word in line.lower() for word in ['engineer', 'analyst', 'manager', 'developer', 'scientist']):
                            title = line
                        elif not company and any(word in line.lower() for word in ['inc', 'corp', 'llc', 'company', 'technologies']):
                            company = line
                        elif not location and any(word in line.lower() for word in ['remote', 'hybrid', 'office', 'location']):
                            location = line
                    
                    if title:
                        st.write(f"**Title:** {title}")
                    if company:
                        st.write(f"**Company:** {company}")
                    if location:
                        st.write(f"**Location:** {location}")
                    
                    # Show word count and key metrics
                    word_count = len(jd_text.split())
                    st.write(f"**Word Count:** {word_count}")
                    
                    # Extract key skills mentioned
                    skill_keywords = ['python', 'sql', 'java', 'javascript', 'react', 'angular', 'docker', 'kubernetes', 'aws', 'azure', 'machine learning', 'data science', 'analytics']
                    mentioned_skills = [skill for skill in skill_keywords if skill in jd_text.lower()]
                    
                    if mentioned_skills:
                        st.write(f"**Key Skills Mentioned:** {', '.join(mentioned_skills[:8])}")
                
                with col2:
                    st.markdown("### 📝 Job Description Preview")
                    # Show first 300 characters
                    preview_text = jd_text[:300] + "..." if len(jd_text) > 300 else jd_text
                    st.text_area("", preview_text, height=200, disabled=True)
                
                # Full job description
                st.markdown("### 📄 Full Job Description")
                with st.expander("View Complete Job Description", expanded=False):
                    st.text(jd_text)
                
                # Requirements extraction with structured categorization
                st.markdown("### ✅ Key Requirements")
                
                # Extract structured requirements
                structured_reqs = extract_structured_requirements(jd_text)
                
                # Display Required section with summary
                if structured_reqs['required']['summary']:
                    st.markdown("#### 🔴 Required (Must Have)")
                    # Display as a clean paragraph
                    required_summary = structured_reqs['required']['summary']
                    # Remove markdown formatting for cleaner display, or keep it for structure
                    st.markdown(required_summary)
                elif structured_reqs['required'].get('yoe') or structured_reqs['required'].get('skills_summary') or structured_reqs['required'].get('education_summary'):
                    st.markdown("#### 🔴 Required (Must Have)")
                    summary_parts = []
                    if structured_reqs['required'].get('yoe'):
                        summary_parts.append(f"**Years of Experience:** {structured_reqs['required']['yoe']}")
                    if structured_reqs['required'].get('skills_summary'):
                        summary_parts.append(f"**Skills & Experience:** {structured_reqs['required']['skills_summary']}")
                    if structured_reqs['required'].get('education_summary'):
                        summary_parts.append(f"**Education:** {structured_reqs['required']['education_summary']}")
                    if summary_parts:
                        st.markdown("\n\n".join(summary_parts))
                else:
                    # Try to extract from full text if no structured data
                    st.markdown("#### 🔴 Required (Must Have)")
                    st.info("No required requirements explicitly found. Check the full job description above.")
                
                # Display Better to Have section with summary
                if structured_reqs['preferred']['summary']:
                    st.markdown("#### 🟡 Better to Have (Preferred)")
                    # Display as a clean paragraph
                    preferred_summary = structured_reqs['preferred']['summary']
                    st.markdown(preferred_summary)
                elif structured_reqs['preferred'].get('yoe') or structured_reqs['preferred'].get('skills_summary') or structured_reqs['preferred'].get('education_summary'):
                    st.markdown("#### 🟡 Better to Have (Preferred)")
                    summary_parts = []
                    if structured_reqs['preferred'].get('yoe'):
                        summary_parts.append(f"**Years of Experience:** {structured_reqs['preferred']['yoe']}")
                    if structured_reqs['preferred'].get('skills_summary'):
                        summary_parts.append(f"**Skills & Experience:** {structured_reqs['preferred']['skills_summary']}")
                    if structured_reqs['preferred'].get('education_summary'):
                        summary_parts.append(f"**Education:** {structured_reqs['preferred']['education_summary']}")
                    if summary_parts:
                        st.markdown("\n\n".join(summary_parts))
                else:
                    # Only show if there's actually preferred content
                    pass
                
                # Top candidates for this job
                if 'similarity_matrix' in st.session_state:
                    st.markdown("### 🏆 Top Candidates for This Job")
                    jd_idx = matcher.jd_names.index(selected_jd)
                    jd_scores = similarity_matrix[:, jd_idx]
                    
                    # Get top 5 candidates
                    top_indices = np.argsort(jd_scores)[::-1][:5]
                    
                    top_candidates_df = pd.DataFrame({
                        'Rank': range(1, len(top_indices) + 1),
                        'Candidate': [matcher.candidate_names[i] for i in top_indices],
                        'Similarity Score': jd_scores[top_indices],
                        'Match Percentage': [f"{score * 100:.1f}%" for score in jd_scores[top_indices]]
                    })
                    
                    st.dataframe(top_candidates_df, use_container_width=True)
                    
                    # Show best match details
                    if len(top_indices) > 0:
                        best_candidate = matcher.candidate_names[top_indices[0]]
                        best_score = jd_scores[top_indices[0]]
                        
                        st.success(f"🎯 **Best Match:** {best_candidate} with {best_score * 100:.1f}% similarity")
            else:
                st.info("Please select a job description to view its details.")
        
        # Professional Summaries tab (available in both modes)
        with tab7:
            st.subheader("✨ AI-Generated Professional Summaries")
            st.markdown("*Intelligent summaries generated from experience and skills analysis*")
            
            # Check if API is enabled
            use_llm_enabled = st.session_state.get('use_llm', False)
            
            # Generate summaries for all candidates
            summaries_data = []
            
            for candidate_name in matcher.candidate_names:
                # Get the resume text
                resume_text = matcher.resumes.get(candidate_name, "")
                if not resume_text:
                    continue
                
                # Extract sections using the matcher's method
                if hasattr(matcher, 'extract_sections'):
                    resume_sections = matcher.extract_sections(resume_text)
                    # Ensure raw_text is included (extract_sections should add it, but ensure it's there)
                    if 'raw_text' not in resume_sections:
                        resume_sections['raw_text'] = resume_text
                else:
                    # Fallback to basic section extraction
                    resume_sections = {
                        'summary': '',
                        'experience': '',
                        'skills': '',
                        'education': '',
                        'raw_text': resume_text  # Always include raw text for fallback
                    }
                
                # Generate professional summary
                if hasattr(matcher, 'generate_professional_summary'):
                    professional_summary = matcher.generate_professional_summary(resume_sections)
                else:
                    professional_summary = "Summary generation not available in this mode."
                
                # Get API-generated summary if enabled (replace "Original Summary" with "API Generated Summary")
                use_llm_enabled = st.session_state.get('use_llm', False)
                api_generated_summary = None
                if use_llm_enabled and LLM_AVAILABLE and is_llm_available():
                    try:
                        llm_client = get_llm_client()
                        # Generate summary using API - use full text without truncation
                        experience_text = resume_sections.get('experience', '') or resume_sections.get('raw_text', '')
                        skills_text = resume_sections.get('skills', '')
                        education_text = resume_sections.get('education', '')
                        raw_text_full = resume_sections.get('raw_text', '')
                        
                        api_generated_summary = llm_client.generate_professional_summary(
                            experience=experience_text,
                            skills=skills_text,
                            education=education_text,
                            raw_text=raw_text_full if raw_text_full else None
                        )
                    except Exception as e:
                        print(f"DEBUG: API summary generation failed: {e}")
                
                summaries_data.append({
                    'Candidate': candidate_name,
                    'AI-Generated Summary': professional_summary,
                    'API Generated Summary': api_generated_summary if api_generated_summary else 'Not available (API disabled or failed)',
                    'Experience Length': len(resume_sections.get('experience', '')),
                    'Skills Count': len(resume_sections.get('skills', '').split(',')) if resume_sections.get('skills') else 0
                })
            
            if summaries_data:
                summaries_df = pd.DataFrame(summaries_data)
                
                # Display metrics
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Total Candidates", len(summaries_df))
                with col2:
                    avg_length = summaries_df['AI-Generated Summary'].str.len().mean()
                    st.metric("Avg Summary Length", f"{avg_length:.0f} chars")
                with col3:
                    summaries_with_experience = len(summaries_df[summaries_df['Experience Length'] > 0])
                    st.metric("With Experience Data", f"{summaries_with_experience}/{len(summaries_df)}")
                
                # Display summaries
                st.subheader("📋 Generated Summaries")
                
                for idx, row in summaries_df.iterrows():
                    with st.expander(f"👤 {row['Candidate']}", expanded=False):
                        st.markdown("**🤖 AI-Generated Summary:**")
                        # Use text_area instead of st.info to show full text without truncation
                        summary_text = row['AI-Generated Summary']
                        # Calculate height based on text length (min 150, max 400)
                        text_height = min(max(150, len(summary_text) // 3), 400)
                        st.text_area(
                            "AI-Generated Summary",
                            value=summary_text,
                            height=text_height,
                            disabled=True,
                            label_visibility="collapsed",
                            key=f"ai_summary_{idx}"
                        )
                        
                        # Show API Generated Summary instead of Original Summary
                        if 'API Generated Summary' in row and row['API Generated Summary'] != 'Not available (API disabled or failed)':
                            st.markdown("**🤖 API Generated Summary:**")
                            api_summary_text = str(row['API Generated Summary'])  # Ensure it's a string
                            # Calculate height based on text length, with higher max for longer summaries
                            # Use more generous height calculation to ensure full text is visible
                            text_height = min(max(300, len(api_summary_text) // 1.5), 1200)
                            st.text_area(
                                "API Generated Summary",
                                value=api_summary_text,
                                height=int(text_height),
                                disabled=True,
                                label_visibility="collapsed",
                                key=f"api_summary_{idx}",
                                max_chars=None  # No character limit
                            )
                        elif use_llm_enabled:
                            st.info("ℹ️ API summary generation is enabled but failed. Check API key and connection.")
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            st.metric("Experience Length", f"{row['Experience Length']} chars")
                        with col2:
                            st.metric("Skills Count", row['Skills Count'])
                
                # Download option
                csv_data = summaries_df.to_csv(index=False)
                st.download_button(
                    label="📥 Download Professional Summaries as CSV",
                    data=csv_data,
                    file_name="professional_summaries.csv",
                    mime="text/csv"
                )
            else:
                st.warning("No candidate data available for summary generation.")
    
    # Footer
    st.markdown("---")
    st.markdown("**💡 Tips:**")
    st.markdown("""
    - **Green (0.7+)**: Excellent match
    - **Yellow (0.4-0.7)**: Good match  
    - **Red (<0.4)**: Poor match
    - Higher scores indicate better alignment between candidate skills and job requirements
    """)

if __name__ == "__main__":
    main()
